# -*- coding: utf-8 -*-
"""Génère le Guide Complet d'Utilisation MUKMAP (PDF + Word) — édition MUKESHABA.

Usage :
    python manage.py generer_guide
"""
import json
import os

from django.conf import settings
from django.core.management.base import BaseCommand

from cartographie.guide_contenu import APROPOS, CHAPITRES, FOOTER, META


class Command(BaseCommand):
    help = "Génère le Guide Complet d'Utilisation MUKMAP en PDF et Word (édition MUKESHABA)."

    def handle(self, *args, **options):
        dossier = os.path.join(settings.BASE_DIR, 'cartographie', 'static', 'docs')
        os.makedirs(dossier, exist_ok=True)
        base = os.path.join(dossier, 'MUKMAP_Guide_Complet_Utilisateur')
        docx = self._docx(base + '.docx')
        pdf = self._pdf(base + '.pdf')
        infos = {
            'app': META['app_nom'],
            'version': META['app_version'],
            'date': META['date_guide'],
            'editeur': META['societe'],
            'chapitres': len(CHAPITRES),
            'pdf': os.path.basename(pdf),
            'docx': os.path.basename(docx),
        }
        with open(os.path.join(dossier, 'infos.json'), 'w', encoding='utf-8') as f:
            json.dump(infos, f, ensure_ascii=False, indent=2)
        self.stdout.write(self.style.SUCCESS("Guide généré : %s" % pdf))
        self.stdout.write(self.style.SUCCESS("Guide généré : %s" % docx))

    # ============================== DOCX ==============================
    def _docx(self, chemin):
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document()
        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.2)
            section.right_margin = Cm(2.2)

        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

        for nom, taille, couleur in (('Heading 1', 16, '1F4E79'), ('Heading 2', 13, '2E74B5')):
            st = doc.styles[nom]
            st.font.name = 'Calibri'
            st.font.size = Pt(taille)
            st.font.color.rgb = RGBColor.from_string(couleur)
            st.font.bold = True

        for section in doc.sections:
            footer = section.footer
            p = footer.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(FOOTER + "   |   Page ")
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            fld = OxmlElement('w:fldSimple')
            fld.set(qn('w:instr'), 'PAGE')
            p._p.append(fld)

        # Page de couverture
        for _ in range(6):
            doc.add_paragraph()
        t = doc.add_paragraph()
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = t.add_run("GUIDE COMPLET D'UTILISATION")
        r.font.size = Pt(34)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        t2 = doc.add_paragraph()
        t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = t2.add_run("%s %s" % (META['app_nom'], META['app_version']))
        r2.font.size = Pt(44)
        r2.font.bold = True
        r2.font.color.rgb = RGBColor(0xC0, 0x50, 0x4D)
        t3 = doc.add_paragraph()
        t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3 = t3.add_run("Manuel officiel — Propriété de " + META['societe'])
        r3.font.size = Pt(16)
        t4 = doc.add_paragraph()
        t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r4 = t4.add_run(META['slogan'])
        r4.font.size = Pt(12)
        r4.font.italic = True
        r4.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        for _ in range(3):
            doc.add_paragraph()
        t5 = doc.add_paragraph()
        t5.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r5 = t5.add_run("Édition %s — Version %s" % (META['date_guide'], META['app_version']))
        r5.font.size = Pt(12)
        doc.add_page_break()

        # Table des matières (champ Word mis à jour à l'ouverture)
        h = doc.add_paragraph()
        rh = h.add_run("Table des matières")
        rh.font.size = Pt(18)
        rh.font.bold = True
        rh.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        p = doc.add_paragraph()
        run = p.add_run()
        fld_begin = OxmlElement('w:fldChar')
        fld_begin.set(qn('w:fldCharType'), 'begin')
        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = 'TOC \\o "1-2" \\h \\z \\u'
        fld_sep = OxmlElement('w:fldChar')
        fld_sep.set(qn('w:fldCharType'), 'separate')
        txt = OxmlElement('w:t')
        txt.text = "La table des matières se met à jour automatiquement à l'ouverture du document (Word)."
        fld_end = OxmlElement('w:fldChar')
        fld_end.set(qn('w:fldCharType'), 'end')
        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_sep)
        run._r.append(txt)
        run._r.append(fld_end)
        settings_elem = doc.settings.element
        upd = OxmlElement('w:updateFields')
        upd.set(qn('w:val'), 'true')
        settings_elem.append(upd)
        doc.add_page_break()

        def ajouter_bloc(bloc):
            typ = bloc[0]
            val = bloc[1]
            if typ == "p":
                doc.add_paragraph(val)
            elif typ == "b":
                for item in val:
                    doc.add_paragraph(item, style='List Bullet')
            elif typ == "n":
                pa = doc.add_paragraph()
                rn = pa.add_run("ℹ " + val)
                rn.font.italic = True
                rn.font.size = Pt(10.5)
                rn.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            elif typ == "c":
                pa = doc.add_paragraph()
                pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
                rc = pa.add_run("[" + val + "]")
                rc.font.italic = True
                rc.font.size = Pt(10)
                rc.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            elif typ == "t":
                colonnes = val["colonnes"]
                lignes = val["lignes"]
                table = doc.add_table(rows=1 + len(lignes), cols=len(colonnes))
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                for j, c in enumerate(colonnes):
                    cell = table.cell(0, j)
                    cell.text = ""
                    pr = cell.paragraphs[0]
                    rcell = pr.add_run(c)
                    rcell.font.bold = True
                    rcell.font.size = Pt(10)
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:fill'), '1F4E79')
                    cell._tc.get_or_add_tcPr().append(shd)
                    rcell.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                for i, ligne in enumerate(lignes):
                    for j, val_cell in enumerate(ligne):
                        cell = table.cell(i + 1, j)
                        cell.text = ""
                        rcell = cell.paragraphs[0].add_run(val_cell)
                        rcell.font.size = Pt(9.5)
                        if i % 2 == 1:
                            shd = OxmlElement('w:shd')
                            shd.set(qn('w:val'), 'clear')
                            shd.set(qn('w:fill'), 'EAF1F8')
                            cell._tc.get_or_add_tcPr().append(shd)
                doc.add_paragraph()

        for chap in CHAPITRES:
            doc.add_heading(chap["titre"], level=1)
            for section in chap["sections"]:
                doc.add_heading(section["titre"], level=2)
                for bloc in section["blocs"]:
                    ajouter_bloc(bloc)
            doc.add_page_break()

        doc.add_page_break()
        doc.add_heading(APROPOS["titre"], level=1)
        for bloc in APROPOS["blocs"]:
            ajouter_bloc(bloc)
        doc.save(chemin)
        return chemin

    # ============================== PDF ==============================
    def _pdf(self, chemin):
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import cm
        from reportlab.lib.colors import HexColor
        from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_LEFT
        from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, PageBreak,
                                        Table, TableStyle, KeepTogether)
        from reportlab.platypus.tableofcontents import TableOfContents
        from reportlab.lib.styles import ParagraphStyle

        BLEU = HexColor('#1F4E79')
        BLEU2 = HexColor('#2E74B5')
        ROUGE = HexColor('#C0504D')
        GRIS = HexColor('#666666')
        FOND_TAB = HexColor('#EAF1F8')

        styles = {
            'Chapitre': ParagraphStyle('Chapitre', fontName='Helvetica-Bold', fontSize=15,
                                       leading=19, textColor=BLEU, spaceBefore=18,
                                       spaceAfter=10, keepWithNext=1),
            'Section': ParagraphStyle('Section', fontName='Helvetica-Bold', fontSize=12,
                                      leading=15, textColor=BLEU2, spaceBefore=12,
                                      spaceAfter=6, keepWithNext=1),
            'Corps': ParagraphStyle('Corps', fontName='Helvetica', fontSize=10, leading=14,
                                    alignment=TA_JUSTIFY, spaceAfter=6),
            'Puce': ParagraphStyle('Puce', fontName='Helvetica', fontSize=10, leading=14,
                                   leftIndent=14, bulletIndent=2, spaceAfter=3,
                                   alignment=TA_LEFT),
            'Note': ParagraphStyle('Note', fontName='Helvetica-Oblique', fontSize=9.5,
                                   leading=13, textColor=BLEU, spaceBefore=4, spaceAfter=8),
            'Capture': ParagraphStyle('Capture', fontName='Helvetica-Oblique', fontSize=9.5,
                                      leading=13, textColor=GRIS, alignment=TA_CENTER,
                                      spaceBefore=4, spaceAfter=10),
            'TabCell': ParagraphStyle('TabCell', fontName='Helvetica', fontSize=8.8,
                                      leading=11),
            'TabHead': ParagraphStyle('TabHead', fontName='Helvetica-Bold', fontSize=8.8,
                                      leading=11, textColor=HexColor('#FFFFFF')),
            'TOC': ParagraphStyle('TOC', fontName='Helvetica-Bold', fontSize=18, leading=22,
                                  textColor=BLEU, spaceAfter=14),
        }

        def pied_page(canvas, doc):
            canvas.saveState()
            if doc.page > 1:
                canvas.setFont('Helvetica', 8)
                canvas.setFillColor(GRIS)
                canvas.drawCentredString(A4[0] / 2.0, 1.15 * cm, FOOTER)
                canvas.drawCentredString(A4[0] / 2.0, 0.75 * cm, "Page %d" % doc.page)
            canvas.restoreState()

        class DocAvecTOC(SimpleDocTemplate):
            def afterFlowable(self, flowable):
                if isinstance(flowable, Paragraph):
                    nom = flowable.style.name
                    if nom == 'Chapitre':
                        self.notify('TOCEntry', (0, flowable.getPlainText(), self.page))
                    elif nom == 'Section':
                        self.notify('TOCEntry', (1, flowable.getPlainText(), self.page))

        doc = DocAvecTOC(chemin, pagesize=A4,
                         leftMargin=2.2 * cm, rightMargin=2.2 * cm,
                         topMargin=2 * cm, bottomMargin=2.2 * cm,
                         title="Guide complet d'utilisation — %s" % META['app_nom'],
                         author=META['societe'])
        story = []

        # Couverture
        story.append(Spacer(1, 4.2 * cm))
        story.append(Paragraph("GUIDE COMPLET D'UTILISATION",
                               ParagraphStyle('Cov1', fontName='Helvetica-Bold', fontSize=32,
                                              leading=38, textColor=BLEU, alignment=TA_CENTER)))
        story.append(Spacer(1, 0.8 * cm))
        story.append(Paragraph("%s %s" % (META['app_nom'], META['app_version']),
                               ParagraphStyle('Cov2', fontName='Helvetica-Bold', fontSize=44,
                                              leading=52, textColor=ROUGE, alignment=TA_CENTER)))
        story.append(Spacer(1, 0.6 * cm))
        story.append(Paragraph("Manuel officiel — Propriété de " + META['societe'],
                               ParagraphStyle('Cov3', fontName='Helvetica', fontSize=16,
                                              leading=22, alignment=TA_CENTER)))
        story.append(Spacer(1, 0.4 * cm))
        story.append(Paragraph(META['slogan'],
                               ParagraphStyle('Cov4', fontName='Helvetica-Oblique', fontSize=12,
                                              leading=16, textColor=GRIS, alignment=TA_CENTER)))
        story.append(Spacer(1, 3.2 * cm))
        story.append(Paragraph("Édition %s — Version %s" % (META['date_guide'], META['app_version']),
                               ParagraphStyle('Cov5', fontName='Helvetica', fontSize=12,
                                              leading=16, alignment=TA_CENTER)))
        story.append(PageBreak())

        # Table des matières
        story.append(Paragraph("Table des matières", styles['TOC']))
        toc = TableOfContents()
        toc.levelStyles = [
            ParagraphStyle('TOC0', fontName='Helvetica-Bold', fontSize=10.5, leading=16),
            ParagraphStyle('TOC1', fontName='Helvetica', fontSize=9.5, leading=13,
                           leftIndent=16),
        ]
        toc.dotsMinLevel = 0
        story.append(toc)
        story.append(PageBreak())

        def ajouter_bloc(story, bloc):
            typ = bloc[0]
            val = bloc[1]
            if typ == "p":
                story.append(Paragraph(val.replace("«", "&#171;").replace("»", "&#187;"),
                                       styles['Corps']))
            elif typ == "b":
                for item in val:
                    story.append(Paragraph("•  " + item, styles['Puce']))
            elif typ == "n":
                story.append(Paragraph("ℹ " + val, styles['Note']))
            elif typ == "c":
                story.append(Paragraph("[" + val + "]", styles['Capture']))
            elif typ == "t":
                colonnes = val["colonnes"]
                lignes = val["lignes"]
                data = [[Paragraph(c, styles['TabHead']) for c in colonnes]]
                for ligne in lignes:
                    data.append([Paragraph(c, styles['TabCell']) for c in ligne])
                table = Table(data, repeatRows=1, hAlign='CENTER')
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), BLEU),
                    ('GRID', (0, 0), (-1, -1), 0.5, HexColor('#9DB6CE')),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('LEFTPADDING', (0, 0), (-1, -1), 5),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 5),
                    ('TOPPADDING', (0, 0), (-1, -1), 3),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
                ]))
                for i in range(1, len(data)):
                    if i % 2 == 1:
                        table.setStyle(TableStyle(
                            [('BACKGROUND', (0, i), (-1, i), FOND_TAB)]))
                story.append(Spacer(1, 3))
                story.append(table)
                story.append(Spacer(1, 7))

        for chap in CHAPITRES:
            story.append(KeepTogether([Paragraph(chap["titre"], styles['Chapitre'])]))
            for section in chap["sections"]:
                story.append(KeepTogether([Paragraph(section["titre"], styles['Section'])]))
                for bloc in section["blocs"]:
                    ajouter_bloc(story, bloc)

        story.append(PageBreak())
        story.append(Paragraph(APROPOS["titre"], styles['Chapitre']))
        for bloc in APROPOS["blocs"]:
            ajouter_bloc(story, bloc)

        doc.multiBuild(story, onFirstPage=pied_page, onLaterPages=pied_page)
        return chemin