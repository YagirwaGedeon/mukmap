import json
import re
import base64
import xml.etree.ElementTree as ET
from xml.sax.saxutils import escape as _xml_escape
import zipfile
import io
import os
import math
import csv
from datetime import date, timedelta

from django.shortcuts import render, redirect, get_object_or_404
from django.urls import reverse
from django.contrib import messages
from django.contrib.auth import authenticate, login, logout, update_session_auth_hash
from django.contrib.auth.decorators import login_required, user_passes_test
from django.contrib.auth.models import User
from django.http import JsonResponse, HttpResponse
from django.conf import settings
from django.db import models as db_models
from django.db.models import Q
from django.db.models.functions import TruncMonth
from django.template.loader import render_to_string
from django.utils import timezone

from .models import (
    PointGeographique, Projet, Activite, ActiviteModele, PhotoActivite,
    ProfilAgent, ZoneSecurite, Itineraire,
    CoucheGeometrie, Geometrie, JournalAudit, MediaPoint,
    CodeAccesAvance, PreferenceUtilisateur, FondCartePersonnalise, CoucheWMS, ImageAerienne,
    SessionTravail, MeteoActivite
)
from .i18n import langue_active

# ─── HELPERS ───────────────────────────────────────────────────


def _xml_safe(texte):
    if not texte:
        return ''
    return _xml_escape(str(texte))


def _float_ou_nul(valeur):
    """Float si convertible, sinon None (précision GPS facultative)."""
    if valeur is None or valeur == '':
        return None
    try:
        return float(valeur)
    except (TypeError, ValueError):
        return None


def _valider_coordonnees_wgs84(lat, lon):
    """Valide des coordonnées WGS84. Retourne (lat, lon) flottants ou (None, None)."""
    lat = _float_ou_nul(lat)
    lon = _float_ou_nul(lon)
    if lat is None or lon is None:
        return None, None
    if not (-90 <= lat <= 90 and -180 <= lon <= 180):
        return None, None
    return lat, lon


def _est_admin(user):
    return user.is_authenticated and user.is_superuser


def _json_script(donnees):
    """JSON sûr pour injection inline dans un bloc <script> (échappe <, >, & et U+2028/29).

    json.dumps n'échappe pas ces caractères : une valeur utilisateur contenant
    « </script> » briserait le bloc et permettrait du XSS stocké.
    """
    return json.dumps(donnees, ensure_ascii=False) \
        .replace('<', '\\u003c').replace('>', '\\u003e') \
        .replace('&', '\\u0026') \
        .replace('\u2028', '\\u2028').replace('\u2029', '\\u2029')


def _est_admin_principal(user):
    """L'administrateur principal (YAGIRWA) dispose du contrôle total (Full Control)."""
    return user.is_authenticated and user.username == 'YAGIRWA'


def _projet_actif(request):
    """Projet sélectionné en session (obligatoire pour toute création de donnée)."""
    pid = request.session.get('projet_actif_id')
    if not pid:
        return None
    try:
        return Projet.objects.filter(pk=int(pid), statut='actif').first()
    except (ValueError, TypeError):
        return None


def _activite_actuelle(request):
    """Nom de l'activité en cours + éventuellement l'id d'une Activite existante."""
    return {
        'nom': request.session.get('activite_actuelle_nom', ''),
        'id': request.session.get('activite_actuelle_id') or None,
    }


def _audit(request, action, details=''):
    if request is None:
        return
    if request.user.is_authenticated:
        ip = request.META.get('REMOTE_ADDR', '')
        JournalAudit.objects.create(
            utilisateur=request.user,
            action=action,
            adresse_ip=ip,
            details=details,
        )


def _distance_haversine(lat1, lon1, lat2, lon2):
    R = 6371000
    phi1, phi2 = math.radians(lat1), math.radians(lat2)
    dphi = math.radians(lat2 - lat1)
    dlambda = math.radians(lon2 - lon1)
    a = math.sin(dphi/2)**2 + math.cos(phi1)*math.cos(phi2)*math.sin(dlambda/2)**2
    return R * 2 * math.atan2(math.sqrt(a), math.sqrt(1-a))


def _point_dans_zone(lat, lon, zone):
    if zone.type_geometrie == 'Point':
        coords = zone.coordonnees
        dist = _distance_haversine(lat, lon, coords[1], coords[0])
        return dist <= (zone.rayon if zone.rayon > 0 else 100)
    return False


def _analyser_itineraire(coords_list):
    zones = ZoneSecurite.objects.all()
    resultats = []
    alertes = []
    for zone in zones:
        touche = False
        for coord in coords_list:
            if _point_dans_zone(coord[1], coord[0], zone):
                touche = True
                break
        if touche:
            resultats.append({
                'zone_id': zone.pk,
                'zone_nom': zone.nom,
                'statut': zone.statut,
                'couleur': zone.couleur(),
                'motif': zone.motif,
            })
            if zone.statut == 'dangereuse':
                alertes.append(f"ATTENTION : Vous traversez la zone dangereuse « {zone.nom} » — {zone.motif}")
    return resultats, alertes


# ─── AUTH ──────────────────────────────────────────────────────


_NB_TENTATIVES_MAX = 5
_DELAI_BLOCAGE_S = 900  # 15 minutes


def _cle_blocage(identifiant, ip):
    return 'mukmap:blocage:{0}:{1}'.format((identifiant or '').strip().lower(), ip or '')


def _est_bloque(identifiant, ip):
    from django.core.cache import cache
    return cache.get(_cle_blocage(identifiant, ip)) or 0


def _incrementer_echec(identifiant, ip):
    from django.core.cache import cache
    cle = _cle_blocage(identifiant, ip)
    compteur = (cache.get(cle) or 0) + 1
    cache.set(cle, compteur, _DELAI_BLOCAGE_S)
    return compteur


def _reinitialiser_echecs(identifiant, ip):
    from django.core.cache import cache
    cache.delete(_cle_blocage(identifiant, ip))


def connexion(request):
    if request.method == "POST":
        username = request.POST.get('username')
        password = request.POST.get('password')
        type_login = request.POST.get('type', 'agent')
        remember = request.POST.get('remember') == 'on'
        ip = request.META.get('REMOTE_ADDR', '')

        if _est_bloque(username, ip) >= _NB_TENTATIVES_MAX:
            messages.error(request, "Trop de tentatives échouées. Réessayez dans 15 minutes.")
            return render(request, 'cartographie/connexion.html')

        user = authenticate(request, username=username, password=password)

        if user is None:
            _incrementer_echec(username, ip)
            messages.error(request, "Identifiants invalides.")
            return render(request, 'cartographie/connexion.html')

        _reinitialiser_echecs(username, ip)

        if password in settings.PASSWORDS_DEFAUT_SUPERADMIN:
            _audit(request, "Connexion refusée (mot de passe par défaut)",
                   f"Compte « {username} »")
            messages.error(request, "Votre mot de passe est encore celui par défaut. "
                                    "Contactez l'administrateur pour le réinitialiser.")
            return render(request, 'cartographie/connexion.html')

        if hasattr(user, 'profil') and user.profil.est_bloque:
            messages.error(request, "Votre compte a été bloqué par un administrateur.")
            return render(request, 'cartographie/connexion.html')

        if type_login == 'admin' and not user.is_superuser:
            messages.error(request, "Ce compte n'est pas un administrateur.")
            return render(request, 'cartographie/connexion.html')

        if type_login == 'agent' and user.is_superuser:
            messages.error(request, "Utilisez l'onglet « Administration » pour vous connecter.")
            return render(request, 'cartographie/connexion.html')

        login(request, user)
        if remember:
            request.session.set_expiry(60 * 60 * 24 * 30)
        else:
            request.session.set_expiry(0)
        _audit(request, "Connexion")
        session_travail = SessionTravail.objects.create(
            utilisateur=user,
            debut=timezone.now(),
        )
        request.session['session_travail_id'] = session_travail.pk

        if user.is_superuser:
            return redirect('index_cartographie')

        if hasattr(user, 'profil'):
            return redirect('index_cartographie')
        return redirect('profil_creer')

    return render(request, 'cartographie/connexion.html')


def deconnexion(request):
    if request.method != "POST":
        return redirect('connexion')
    observations = request.POST.get('observations', '').strip()[:1000]
    session_travail = None
    sid = request.session.get('session_travail_id')
    if sid:
        try:
            session_travail = SessionTravail.objects.get(pk=sid)
        except (SessionTravail.DoesNotExist, ValueError):
            session_travail = None
    if session_travail is None and request.user.is_authenticated:
        session_travail = (SessionTravail.objects
                           .filter(utilisateur=request.user, fin__isnull=True)
                           .order_by('-debut').first())
    if session_travail is not None:
        session_travail.fin = timezone.now()
        if observations:
            session_travail.observations = observations
        session_travail.save()
    _audit(request, "Déconnexion", details=observations)
    logout(request)
    return redirect('connexion')


@login_required
def profil_creer(request):
    if hasattr(request.user, 'profil'):
        return redirect('index_cartographie')
    if request.method == "POST":
        telephone = request.POST.get('telephone')
        fonction = request.POST.get('fonction')
        motif = request.POST.get('motif_mission', '')
        if not all([telephone, fonction]):
            messages.error(request, "Tous les champs obligatoires doivent être remplis.")
            return render(request, 'cartographie/profil_form.html')
        ProfilAgent.objects.create(
            utilisateur=request.user,
            telephone=telephone,
            fonction=fonction,
            motif_mission=motif,
        )
        _audit(request, "Création de profil agent")
        messages.success(request, "Profil créé avec succès.")
        return redirect('index_cartographie')
    return render(request, 'cartographie/profil_form.html')


@login_required
def profil_edit(request):
    profil = get_object_or_404(ProfilAgent, utilisateur=request.user)
    if request.method == "POST":
        profil.telephone = request.POST.get('telephone', profil.telephone)
        profil.fonction = request.POST.get('fonction', profil.fonction)
        profil.motif_mission = request.POST.get('motif_mission', '')
        if request.FILES.get('photo'):
            profil.photo = request.FILES['photo']
        profil.save()
        _audit(request, "Modification de profil")
        messages.success(request, "Profil mis à jour.")
        return redirect('index_cartographie')
    return render(request, 'cartographie/profil_form.html', {'profil': profil, 'edition': True})


# ─── PAGE PRINCIPALE (CARTE) ───────────────────────────────────


def _extraire_exif(fichier):
    """Retourne (date_prise, latitude, longitude) extraits des EXIF d'une photo."""
    date_prise = latitude = longitude = None
    try:
        import io as _io
        from PIL import Image, ExifTags
        contenu = fichier.read()
        fichier.seek(0)
        img = Image.open(_io.BytesIO(contenu))
        exif = img._getexif() or {}
        tags = {ExifTags.TAGS.get(k, k): v for k, v in exif.items()}
        if tags.get('DateTimeOriginal'):
            try:
                from datetime import datetime as _dt
                date_prise = _dt.strptime(tags['DateTimeOriginal'], '%Y:%m:%d %H:%M:%S')
            except (ValueError, TypeError):
                pass
        gps = tags.get('GPSInfo')
        if gps:
            gps_tags = {ExifTags.GPSTAGS.get(k, k): v for k, v in gps.items()}

            def _deg(v):
                d, m, s = v
                return float(d) + float(m) / 60.0 + float(s) / 3600.0
            try:
                lat = _deg(gps_tags['GPSLatitude'])
                if gps_tags.get('GPSLatitudeRef') in ('S', 's'):
                    lat = -lat
                lng = _deg(gps_tags['GPSLongitude'])
                if gps_tags.get('GPSLongitudeRef') in ('W', 'w'):
                    lng = -lng
                latitude, longitude = round(lat, 6), round(lng, 6)
            except (KeyError, ValueError, TypeError):
                pass
        img.close()
    except Exception:
        pass
    fichier.seek(0)
    return date_prise, latitude, longitude


def _creer_medias_point(point, fichiers, utilisateur=None, commentaire='', date_prise_defaut=None):
    from django.conf import settings as _settings
    taille_max = getattr(_settings, 'MEDIA_TAILLE_MAX_OCTETS', 25 * 1024 * 1024)
    for f in fichiers or []:
        ext = (f.name or '').rsplit('.', 1)[-1].lower() if '.' in (f.name or '') else ''
        if ext in ('png', 'jpg', 'jpeg', 'gif', 'webp', 'bmp'):
            type_media = 'photo'
        elif ext in ('mp4', 'webm', 'mov', 'avi', 'mkv'):
            type_media = 'video'
        elif ext == 'pdf':
            type_media = 'pdf'
        elif ext in ('mp3', 'wav', 'ogg', 'm4a', 'aac'):
            type_media = 'audio'
        else:
            type_media = 'photo'
        if f.size and f.size > taille_max:
            continue
        if type_media == 'photo':
            try:
                from PIL import Image as _ImgVerif
                f.seek(0)
                _image = _ImgVerif.open(f)
                _image.verify()
                f.seek(0)
                _image = _ImgVerif.open(f)
                _image.load()
                f.seek(0)
            except Exception:
                continue
        date_prise, latitude, longitude = None, None, None
        if type_media == 'photo':
            date_prise, latitude, longitude = _extraire_exif(f)
        MediaPoint.objects.create(
            point=point, type=type_media, fichier=f,
            date_prise=date_prise or date_prise_defaut,
            latitude=latitude, longitude=longitude,
            utilisateur=utilisateur, commentaire=commentaire,
        )


def _serialiser_points_carte(points):
    """Sérialise une queryset de PointGeographique au format attendu par la carte."""
    return [
        {
            "id": p.pk, "nom": p.nom, "description": p.description,
            "latitude": p.latitude, "longitude": p.longitude,
            "photo": p.photo.url if p.photo else '',
            "precision_gps_m": p.precision_gps_m,
            "categorie": p.categorie, "statut": p.statut,
            "province": p.province, "commune": p.commune, "quartier": p.quartier,
            "projet": p.projet.nom if p.projet else '',
            "projet_id": p.projet_id,
            "donnees": p.donnees or {},
            "source_fichier": p.source_fichier,
            "source_format": p.source_format,
            "medias": [{"url": m.fichier.url, "type": m.type,
                        "date_prise": m.date_prise.strftime('%d/%m/%Y %H:%M') if m.date_prise else '',
                        "date_upload": m.date_upload.strftime('%d/%m/%Y %H:%M'),
                        "latitude": m.latitude, "longitude": m.longitude,
                        "utilisateur": (m.utilisateur.get_full_name() or m.utilisateur.username) if m.utilisateur else '',
                        "commentaire": m.commentaire} for m in p.medias.all()],
            "auteur": p.auteur.get_full_name() or p.auteur.username if p.auteur else 'Anonyme',
            "date": p.date_creation.strftime('%d/%m/%Y %H:%M'),
        }
        for p in points
    ]


def points_donnees(request):
    """JSON des points pour recharger la carte après un import AJAX."""
    qs = PointGeographique.objects.select_related('auteur', 'projet').prefetch_related('medias') \
        .filter(supprime=False)
    projet_actif = _projet_actif(request)
    if projet_actif is not None:
        qs = qs.filter(projet=projet_actif)
    return JsonResponse(_serialiser_points_carte(qs), safe=False)


def index_cartographie(request):
    est_invite = not request.user.is_authenticated
    projet_actif = _projet_actif(request)
    activite_actuelle = _activite_actuelle(request)

    if request.method == "POST":
        if est_invite:
            messages.error(request, "Veuillez vous connecter pour enregistrer des données.")
            return redirect('connexion')
        if projet_actif is None:
            messages.error(request, "Veuillez sélectionner un projet avant de commencer.")
            return redirect('index_cartographie')
        nom = request.POST.get('nom')
        latitude = request.POST.get('latitude')
        longitude = request.POST.get('longitude')
        if not all([nom, latitude, longitude]):
            messages.error(request, "Nom, latitude et longitude requis.")
            return redirect('index_cartographie')
        lat, lng = _valider_coordonnees_wgs84(latitude, longitude)
        if lat is None:
            messages.error(request, "Coordonnées invalides (hors plage WGS84).")
            return redirect('index_cartographie')
        activite_id = None
        if activite_actuelle.get('id'):
            activite_id = activite_actuelle['id']
        point = PointGeographique.objects.create(
            nom=nom[:200],
            description=request.POST.get('description', ''),
            latitude=lat,
            longitude=lng,
            precision_gps_m=_float_ou_nul(request.POST.get('precision_gps_m')),
            categorie=request.POST.get('categorie', 'autre'),
            statut=request.POST.get('statut', 'actif'),
            province=request.POST.get('province', ''),
            commune=request.POST.get('commune', ''),
            quartier=request.POST.get('quartier', ''),
            projet=projet_actif,
            activite_id=activite_id,
            auteur=request.user,
        )
        if request.FILES.get('photo'):
            point.photo = request.FILES['photo']
            point.save()
        _creer_medias_point(point, request.FILES.getlist('medias'), utilisateur=request.user, commentaire=request.POST.get('commentaire_medias', ''), date_prise_defaut=timezone.now())
        _audit(request, "Création de point", f"Point #{point.pk} - {nom} ({projet_actif.nom})")
        messages.success(request, f"Point « {nom} » enregistré.")
        return redirect('index_cartographie')

    points_qs = PointGeographique.objects.select_related('auteur', 'projet').prefetch_related('medias') \
        .filter(supprime=False)
    activites_qs = Activite.objects.select_related('projet', 'agent').prefetch_related('photos')
    couches_qs = CoucheGeometrie.objects.all()
    zones_qs = ZoneSecurite.objects.all()

    if projet_actif is not None:
        activites_qs = activites_qs.filter(projet=projet_actif)
        couches_qs = couches_qs.filter(projet=projet_actif)
        zones_qs = zones_qs.filter(projet=projet_actif)
    elif est_invite:
        zones_qs = zones_qs.none()
        couches_qs = couches_qs.none()

    points = points_qs.all()
    points_liste = _serialiser_points_carte(points)

    activites = activites_qs.all()
    activites_json = [
        {
            "id": a.pk, "projet": a.projet.nom,
            "nom_activite": a.nom_activite or '',
            "rapport": a.rapport[:100], "observations": a.observations,
            "beneficiaires": a.nombre_beneficiaires,
            "latitude": a.latitude, "longitude": a.longitude,
            "date": a.date_creation.strftime('%d/%m/%Y %H:%M'),
            "photos": [p.image.url for p in a.photos.all()],
            "agent": a.agent.get_full_name() or a.agent.username if a.agent else '',
            "zone_visitee": a.zone_visitee,
            "niveau_securite": a.niveau_securite,
        }
        for a in activites
    ]

    couches = couches_qs.all()
    projets = Projet.objects.filter(statut='actif').order_by('nom')
    zones = zones_qs.all()

    agents_data = [
        {
            "nom": p.utilisateur.get_full_name() or p.utilisateur.username,
            "telephone": p.telephone,
            "fonction": p.fonction,
            "latitude": p.latitude,
            "longitude": p.longitude,
            "photo": p.photo.url if p.photo else '',
        }
        for p in ProfilAgent.objects.select_related('utilisateur').all()
        if p.latitude and p.longitude
    ]

    zones_json = [
        {
            "id": z.pk, "nom": z.nom, "statut": z.statut,
            "couleur": z.couleur(), "motif": z.motif,
            "type": z.type_geometrie, "coordonnees": z.coordonnees,
            "rayon": z.rayon, "auteur": z.auteur.username if z.auteur else '',
            "date": z.date_declaration.strftime('%d/%m/%Y'),
        }
        for z in zones
    ]

    ctx = {
        'points_json': _json_script(points_liste),
        'activites_json': _json_script(activites_json),
        'agents_json': _json_script(agents_data),
        'couches': couches,
        'projets': projets,
        'zones_json': _json_script(zones_json),
        'est_invite': est_invite,
        'est_admin': _est_admin(request.user),
        'projet_actif': projet_actif,
        'activite_actuelle': activite_actuelle,
        'activites_recentes': activites_qs.exclude(nom_activite='').values('nom_activite').distinct()[:8],
        'session_projet_id': request.session.get('projet_actif_id') or 0,
        'rapport_import': _json_script(request.session.pop('rapport_import', None)),
        'couche_importee': request.GET.get('importe', ''),
        'est_admin_principal': _est_admin_principal(request.user),
        'mode_json': _json_script(_etat_mode(request)),
    }
    return render(request, 'cartographie/carte.html', ctx)


# ─── API GÉOMÉTRIES ────────────────────────────────────────────


def geometrie_donnees(request):
    est_invite = not request.user.is_authenticated
    if est_invite:
        return JsonResponse([], safe=False)
    projet_actif = _projet_actif(request)
    couches_qs = CoucheGeometrie.objects.all()
    if projet_actif is not None:
        couches_qs = couches_qs.filter(projet=projet_actif)
    donnees = []
    for couche in couches_qs:
        features = []
        for geom in couche.geometries.all():
            features.append({
                "type": "Feature",
                "geometry": {"type": geom.type, "coordinates": geom.coordonnees},
                "properties": {**geom.proprietes, "couche_nom": couche.nom, "couche_id": couche.pk, "couleur": couche.style_couleur, "geom_id": geom.pk},
            })
        donnees.append({
            "id": couche.pk, "nom": couche.nom, "type": couche.type_geometrie,
            "couleur": couche.style_couleur,
            "style_options": couche.style_options or {},
            "fichier_source": couche.fichier_source or '',
            "date_import": couche.date_import.isoformat() if couche.date_import else '',
            "nb_geometries": couche.geometries.count(),
            "geojson": {"type": "FeatureCollection", "features": features},
        })
    return JsonResponse(donnees, safe=False)


# ─── DASHBOARD ─────────────────────────────────────────────────


@login_required
def dashboard(request):
    projet_id = request.GET.get('projet')
    agent_id = request.GET.get('agent')
    projet_actif = _projet_actif(request)
    est_admin = _est_admin(request.user)

    if not projet_id and (projet_actif is not None or not est_admin):
        projet_id = projet_actif.pk if projet_actif is not None else ''

    activites = Activite.objects.select_related('projet', 'agent').prefetch_related('photos')
    points_qs = PointGeographique.objects.filter(supprime=False)
    zones_qs = ZoneSecurite.objects.all()
    if projet_id:
        activites = activites.filter(projet_id=projet_id)
        points_qs = points_qs.filter(projet_id=projet_id)
        zones_qs = zones_qs.filter(projet_id=projet_id)
    if agent_id:
        activites = activites.filter(agent_id=agent_id)

    total_activites = activites.count()
    total_beneficiaires = activites.aggregate(s=db_models.Sum('nombre_beneficiaires'))['s'] or 0
    zones_dangereuses = zones_qs.filter(statut='dangereuse').count()
    zones_securisees = zones_qs.filter(statut='securisee').count()
    zones_indisponibles = zones_qs.filter(statut='indisponible').count()
    total_zones = zones_dangereuses + zones_securisees + zones_indisponibles

    total_points = points_qs.count()
    points_categories = [
        {'nom': dict(PointGeographique.CATEGORIE_CHOICES).get(c['categorie'], c['categorie']),
         'valeur': c['total']}
        for c in points_qs.values('categorie')
        .annotate(total=db_models.Count('id')).order_by('-total')
    ]
    points_par_statut = [
        {'nom': dict(PointGeographique.STATUT_CHOICES).get(s['statut'], s['statut']),
         'valeur': s['total']}
        for s in points_qs.values('statut')
        .annotate(total=db_models.Count('id')).order_by('-total')
    ]
    points_par_province = [
        {'nom': p['province'] or 'Non renseignée', 'valeur': p['total']}
        for p in points_qs.values('province')
        .annotate(total=db_models.Count('id')).order_by('-total')[:10]
    ]
    points_par_mois = [
        {'mois': m['mois'].strftime('%m/%Y'), 'valeur': m['total']}
        for m in reversed(list(points_qs.annotate(mois=TruncMonth('date_creation'))
        .values('mois').annotate(total=db_models.Count('id')).order_by('-mois')[:12]))
    ]
    activites_par_projet = [
        {'nom': a['projet__nom'], 'valeur': a['total']}
        for a in activites.values('projet__nom')
        .annotate(total=db_models.Count('id')).order_by('-total')[:10]
    ]
    benef_par_mois = [
        {'mois': b['mois'].strftime('%m/%Y'), 'valeur': b['total']}
        for b in reversed(list(activites.annotate(mois=TruncMonth('date_creation'))
        .values('mois').annotate(total=db_models.Sum('nombre_beneficiaires'))
        .order_by('-mois')[:12]))
    ]

    ctx = {
        'activites': activites,
        'projets': Projet.objects.filter(statut='actif'),
        'agents': ProfilAgent.objects.select_related('utilisateur').all(),
        'projet_selectionne': int(projet_id) if projet_id else None,
        'agent_selectionne': int(agent_id) if agent_id else None,
        'projet_actif': projet_actif,
        'activite_actuelle': _activite_actuelle(request),
        'maintenant': timezone.now(),
        'est_admin': est_admin,
        'total_activites': total_activites,
        'total_beneficiaires': total_beneficiaires,
        'zones_dangereuses': zones_dangereuses,
        'zones_securisees': zones_securisees,
        'zones_indisponibles': zones_indisponibles,
        'total_zones': total_zones,
        'zones': zones_qs.select_related('auteur').order_by('-date_declaration')[:8],
        'total_points': total_points,
        'points_categories_json': _json_script(points_categories),
        'points_par_statut_json': _json_script(points_par_statut),
        'points_par_province_json': _json_script(points_par_province),
        'points_par_mois_json': _json_script(points_par_mois),
        'activites_par_projet_json': _json_script(activites_par_projet),
        'benef_par_mois_json': _json_script(benef_par_mois),
    }
    _audit(request, "Consultation du tableau de bord")
    return render(request, 'cartographie/dashboard.html', ctx)


# ─── QUALITÉ DES DONNÉES ───────────────────────────────────────


@login_required
def tableau_qualite(request):
    """Contrôle de qualité automatique : tableau des entités et des règles."""
    from . import qualite
    from .models import (OuvrageHydraulique, ReleveSource, ReleveConsommation,
                         TraceAdduction, ProjetAdduction)

    module_f = (request.GET.get('module') or '').strip()
    regle_f = (request.GET.get('regle') or '').strip()
    gravite_f = (request.GET.get('gravite') or '').strip()
    texte_q = (request.GET.get('q') or '').strip().lower()
    try:
        page = max(1, int(request.GET.get('page') or 1))
    except (TypeError, ValueError):
        page = 1
    par_page = 100

    # 1) Points du module générique (hors corbeille)
    points = list(PointGeographique.objects.filter(supprime=False)
                  .select_related('projet', 'auteur'))
    res_points = qualite.evaluer_points(points)

    # 2) Ouvrages hydrauliques (+ débits présents dans les relevés)
    ouvrages = list(OuvrageHydraulique.objects.select_related('projet').all())
    releve_debits = {}
    for reg_id in ReleveSource.objects.exclude(debit_mesure=None) \
            .values_list('ouvrage_id', flat=True):
        releve_debits[reg_id] = True
    for reg_id in ReleveConsommation.objects.exclude(debit_estime=None) \
            .values_list('ouvrage_id', flat=True):
        releve_debits[reg_id] = True
    res_ouv = qualite.evaluer_ouvrages(ouvrages, releve_debits)

    # 3) Tracés de conduites (rattachement aux extrémités)
    traces = list(TraceAdduction.objects.all())
    ouvrages_par_proj = {}
    for o in ouvrages:
        ouvrages_par_proj.setdefault(o.projet_id, []).append(o)
    res_traces = qualite.evaluer_traces(traces, ouvrages_par_proj)

    # Consolidation en lignes
    def etiquettes(codes):
        return [{'code': c,
                 'cle': qualite.REGLES_PAR_CODE.get(c, {}).get('cle', c),
                 'gravite': qualite.REGLES_PAR_CODE.get(c, {}).get('gravite',
                                                                   qualite.GRAVITE_OK)}
                for c in sorted(codes)]

    lignes = []
    for p, codes, _gr in res_points:
        lignes.append({
            'module': 'points', 'module_cle': 'q_module_points',
            'sigle': 'P', 'id': p.pk, 'nom': p.nom,
            'type_leg': dict(PointGeographique.CATEGORIE_CHOICES).get(p.categorie, p.categorie),
            'projet': p.projet.nom if p.projet else '',
            'lat': p.latitude, 'lon': p.longitude,
            'codes': etiquettes(codes),
            'gravite': qualite.gravite_statut(codes),
        })
    for o, codes, _gr in res_ouv:
        lignes.append({
            'module': 'adduction', 'module_cle': 'q_module_adduction',
            'sigle': 'A', 'id': o.pk, 'nom': o.nom,
            'type_leg': dict(OuvrageHydraulique.TYPE_CHOICES).get(o.type, o.type),
            'projet': o.projet.nom if o.projet else '',
            'lat': o.latitude, 'lon': o.longitude,
            'codes': etiquettes(codes),
            'gravite': qualite.gravite_statut(codes),
        })
    for t, codes, _gr in res_traces:
        projet_nom = ''
        for o in ouvrages:
            if o.projet_id == t.projet_id:
                projet_nom = o.projet.nom if o.projet else ''
                break
        lignes.append({
            'module': 'traces', 'module_cle': 'q_module_traces',
            'sigle': 'T', 'id': t.pk, 'nom': t.nom or 'Tracé #' + str(t.pk),
            'type_leg': 'Tracé de conduite',
            'projet': projet_nom,
            'lat': None, 'lon': None,
            'codes': etiquettes(codes),
            'gravite': qualite.gravite_statut(codes),
        })

    # Totaux globaux (avant filtres)
    totaux = {qualite.GRAVITE_OK: 0, qualite.GRAVITE_A_VERIFIER: 0,
              qualite.GRAVITE_ERREUR: 0}
    for l in lignes:
        totaux[l['gravite']] = totaux.get(l['gravite'], 0) + 1

    # Statistiques par règle (entités touchées), sur toutes les lignes
    regles_stats = []
    for code, gravite, cle in qualite.REGLES:
        nb = sum(1 for l in lignes if any(c['code'] == code for c in l['codes']))
        if nb:
            regles_stats.append({'code': code, 'gravite': gravite,
                                 'cle': cle, 'nb': nb})
    regles_stats.sort(key=lambda r: (-qualite.POIDS[r['gravite']], -r['nb']))

    # Filtres
    lignes_f = lignes
    if module_f:
        lignes_f = [l for l in lignes_f if l['module'] == module_f]
    if gravite_f:
        lignes_f = [l for l in lignes_f if l['gravite'] == gravite_f]
    if regle_f:
        lignes_f = [l for l in lignes_f
                    if any(c['code'] == regle_f for c in l['codes'])]
    if texte_q:
        lignes_f = [l for l in lignes_f
                    if texte_q in (l['nom'] or '').lower()
                    or texte_q in (l['projet'] or '').lower()]
    lignes_f.sort(key=lambda l: (-qualite.POIDS.get(l['gravite'], 0),
                                 (l['nom'] or '').lower()))

    if request.GET.get('export') == 'csv':
        reponse = HttpResponse(content_type='text/csv; charset=utf-8')
        reponse['Content-Disposition'] = 'attachment; filename="mukmap_qualite_donnees.csv"'
        w = csv.writer(reponse)
        en_tete = ['module', 'type', 'nom', 'projet', 'latitude', 'longitude',
                   'statut', 'regles']
        w.writerow(en_tete)
        for l in lignes_f:
            w.writerow([l['module'], l['type_leg'], l['nom'], l['projet'],
                        l['lat'], l['lon'], l['gravite'],
                        ' | '.join(c['code'] for c in l['codes'])])
        return reponse

    total_filtre = len(lignes_f)
    debut = (page - 1) * par_page
    lignes_page = lignes_f[debut:debut + par_page]
    pages = max(1, math.ceil(total_filtre / par_page)) if total_filtre else 1

    ctx = {
        'lignes': lignes_page,
        'totaux': totaux,
        'regles_stats': regles_stats,
        'total_filtre': total_filtre,
        'page': page, 'pages': pages,
        'pages_range': range(max(1, page - 4), min(pages, page + 4) + 1),
        'module_filtre': module_f, 'regle_filtre': regle_f,
        'gravite_filtre': gravite_f, 'recherche': texte_q,
        'module_choices': [('points', 'Points'), ('adduction', 'Adduction'),
                           ('traces', 'Tracés')],
    }
    _audit(request, "Consultation du tableau de qualité des données")
    return render(request, 'cartographie/qualite_tableau.html', ctx)


# ─── TABLEAU DE BORD DU PROJET D'ADDUCTION D'EAU ──────────────


@login_required
def adduction_dashboard(request):
    """Tableau de bord du projet d'adduction d'eau : indicateurs clés
    (sources, villages, bénéficiaires, bornes-fontaines, réservoirs,
    conduites, longueurs, distances, photos, vérifications) + graphiques."""
    from .models import (OuvrageHydraulique, TraceAdduction, ProjetAdduction,
                         ReleveSource, ReleveVillage, ReleveConsommation,
                         ReleveRepere, ReleveReservoir)
    from . import qualite

    projets = list(ProjetAdduction.objects.all().order_by('nom'))
    pid = (request.GET.get('projet') or '').strip()
    projet = None
    if pid:
        projet = next((p for p in projets if str(p.pk) == pid), None)
    if projet is None and projets:
        projet = projets[0]

    # Stub si aucun projet : page vide mais accessible
    if projet is None:
        ctx = {
            'projets': projets, 'projet': None,
            'projet_selectionne': None,
            'aucun_projet': True,
            'kpis': {}, 'chart_types': '[]', 'chart_villages': '[]',
            'chart_conduites': '[]', 'chart_statuts': '[]',
            'chart_mois': '[]', 'chart_qualite': '[]',
        }
        _audit(request, "Consultation du tableau de bord adduction (aucun projet)")
        return render(request, 'cartographie/adduction_dashboard.html', ctx)

    ouvrages = list(projet.ouvrages.all())
    traces = list(projet.tracs.all())

    # ── Indicateurs clés ───────────────────────────────────────
    nb_sources = sum(1 for o in ouvrages if o.type == 'source')
    nb_villages = sum(1 for o in ouvrages if o.type == 'village')
    nb_bornes = sum(1 for o in ouvrages if o.type == 'borne')
    nb_reservoirs = sum(1 for o in ouvrages if o.type == 'reservoir')
    nb_conduites = len(traces)
    nb_points = len(ouvrages)

    longueur_m = sum((t.longueur_m or 0) for t in traces)
    longueur_totale = round(longueur_m / 1000.0, 2)

    releve_villages = {rv.ouvrage_id: rv for rv in ReleveVillage.objects.filter(ouvrage__projet=projet)}
    nb_beneficiaires = 0
    for o in ouvrages:
        rv = releve_villages.get(o.pk)
        if o.type == 'village' and rv and rv.beneficiaires_estimes:
            nb_beneficiaires += rv.beneficiaires_estimes
        else:
            nb_beneficiaires += o.beneficiaires or 0

    # Photos : photo principale + photos des formulaires spécialisés
    nb_photos = sum(1 for o in ouvrages if o.photo)
    for rc in ReleveConsommation.objects.filter(ouvrage__projet=projet):
        nb_photos += len(rc.photos or [])
    for rrv in ReleveReservoir.objects.filter(ouvrage__projet=projet):
        nb_photos += len(rrv.photos or [])
    for rr in ReleveRepere.objects.filter(ouvrage__projet=projet):
        if rr.photo:
            nb_photos += 1

    # Distance moyenne source → village (mesurée ou haversine au plus proche)
    releve_sources = {rs.ouvrage_id: rs for rs in ReleveSource.objects.filter(ouvrage__projet=projet)}
    sources = [o for o in ouvrages if o.type in ('source', 'captage')]
    villages = [o for o in ouvrages if o.type == 'village']
    distances = []
    for s in sources:
        rs = releve_sources.get(s.pk)
        d = None
        if rs and rs.distance_village_m:
            d = rs.distance_village_m
        elif villages:
            distances_v = []
            for v in villages:
                if not qualite._coordonnees_valides(v.latitude, v.longitude):
                    continue
                dv = qualite._distance_m(s.latitude, s.longitude, v.latitude, v.longitude)
                if dv is not None:
                    distances_v.append(dv)
            if distances_v:
                d = min(distances_v)
        if d is not None and d > 0:
            distances.append(d)
    distance_moyenne = round(sum(distances) / len(distances), 0) if distances else None

    # Points nécessitant une vérification (module qualité)
    releve_debits = set()
    for reg_id in ReleveSource.objects.exclude(debit_mesure=None).values_list('ouvrage_id', flat=True):
        releve_debits.add(reg_id)
    for reg_id in ReleveConsommation.objects.exclude(debit_estime=None).values_list('ouvrage_id', flat=True):
        releve_debits.add(reg_id)
    res_ouv = qualite.evaluer_ouvrages(ouvrages, {i: True for i in releve_debits})
    nb_verifier = sum(1 for _o, codes, _g in res_ouv
                      if qualite.gravite_statut(codes) != qualite.GRAVITE_OK)
    res_traces = qualite.evaluer_traces(traces, {projet.pk: ouvrages})
    nb_verifier += sum(1 for _t, codes, _g in res_traces
                       if qualite.gravite_statut(codes) != qualite.GRAVITE_OK)

    kpis = {
        'sources': nb_sources, 'villages': nb_villages,
        'beneficiaires': nb_beneficiaires, 'bornes': nb_bornes,
        'reservoirs': nb_reservoirs, 'conduites': nb_conduites,
        'longueur': longueur_totale, 'distance_moyenne': distance_moyenne,
        'points': nb_points, 'photos': nb_photos, 'verifier': nb_verifier,
    }

    # ── Données des graphiques ──────────────────────────────────
    type_labels = dict(OuvrageHydraulique.TYPE_CHOICES) if OuvrageHydraulique else {}
    par_type = {}
    for o in ouvrages:
        par_type[o.type] = par_type.get(o.type, 0) + 1
    types_chart = [{'nom': type_labels.get(k, k), 'valeur': v}
                   for k, v in sorted(par_type.items(), key=lambda x: -x[1])]

    statut_labels = dict(OuvrageHydraulique.STATUT_CHOICES) if OuvrageHydraulique else {}
    par_statut = {}
    for o in ouvrages:
        par_statut[o.statut] = par_statut.get(o.statut, 0) + 1
    statuts_chart = [{'nom': statut_labels.get(k, k or '—'), 'valeur': v}
                     for k, v in sorted(par_statut.items(), key=lambda x: -x[1])]

    villages_chart = []
    for o in sorted(ouvrages, key=lambda x: -(x.beneficiaires or 0))[:10]:
        if o.type != 'village':
            continue
        rv = releve_villages.get(o.pk)
        benef = (rv.beneficiaires_estimes if rv and rv.beneficiaires_estimes
                 else o.beneficiaires or 0)
        villages_chart.append({'nom': o.nom, 'valeur': benef})
    if len(villages_chart) < 1:
        villages_chart = [{'nom': o.nom, 'valeur': o.beneficiaires or 0}
                          for o in sorted([x for x in ouvrages if x.type == 'village'],
                                          key=lambda x: -(x.beneficiaires or 0))[:10]]

    conduites_chart = [{'nom': (t.nom or 'Tracé #' + str(t.pk)),
                        'valeur': round(t.longueur_m or 0, 1)}
                       for t in sorted(traces, key=lambda x: -(x.longueur_m or 0))[:10]]

    mois_chart = []
    for m in reversed(list(projet.ouvrages.annotate(mois=TruncMonth('date_releve'))
                           .values('mois').annotate(total=db_models.Count('id'))
                           .order_by('-mois')[:12])):
        mois_chart.append({'mois': m['mois'].strftime('%m/%Y'), 'valeur': m['total']})

    nb_ok = sum(1 for _o, codes, _g in res_ouv if qualite.gravite_statut(codes) == qualite.GRAVITE_OK)
    nb_av = sum(1 for _o, codes, _g in res_ouv if qualite.gravite_statut(codes) == qualite.GRAVITE_A_VERIFIER)
    nb_err = sum(1 for _o, codes, _g in res_ouv if qualite.gravite_statut(codes) == qualite.GRAVITE_ERREUR)
    qualite_chart = [
        {'nom': 'Conforme', 'valeur': nb_ok, 'couleur': '#22c55e'},
        {'nom': 'À vérifier', 'valeur': nb_av, 'couleur': '#f59e0b'},
        {'nom': 'Erreur', 'valeur': nb_err, 'couleur': '#ef4444'},
    ]

    ctx = {
        'projets': projets, 'projet': projet,
        'projet_selectionne': projet.pk,
        'aucun_projet': False,
        'maintenant': timezone.now(),
        'kpis': kpis,
        'types_chart_json': _json_script(types_chart),
        'villages_chart_json': _json_script(villages_chart),
        'conduites_chart_json': _json_script(conduites_chart),
        'statuts_chart_json': _json_script(statuts_chart),
        'mois_chart_json': _json_script(mois_chart),
        'qualite_chart_json': _json_script(qualite_chart),
    }
    _audit(request, "Consultation du tableau de bord adduction",
           f"Projet #{projet.pk} - {projet.nom}")
    return render(request, 'cartographie/adduction_dashboard.html', ctx)


# ─── ACTIVITÉS ─────────────────────────────────────────────────


def _extraire_meteo_post(request):
    """Lit les champs météo cachés du formulaire d'activité (remplis par le widget).
    Retourne un dict normalisé (mêmes clés que meteo.recuperer_meteo) ou None."""
    nom_champs = ('meteo_latitude', 'meteo_longitude', 'meteo_temperature',
                  'meteo_conditions', 'meteo_code', 'meteo_humidite', 'meteo_vent_kmh',
                  'meteo_vent_direction', 'meteo_vent_direction_deg', 'meteo_proba_pluie',
                  'meteo_lever', 'meteo_coucher', 'meteo_localisation', 'meteo_source',
                  'meteo_horodatage')
    if not any(request.POST.get(c) for c in nom_champs):
        return None

    def _f(v):
        try:
            return float(v)
        except (TypeError, ValueError):
            return None

    def _i(v):
        try:
            return int(float(v))
        except (TypeError, ValueError):
            return None

    from .meteo import _valider_coordonnees
    lat, lon = _valider_coordonnees(
        request.POST.get('meteo_latitude'), request.POST.get('meteo_longitude'))
    source = (request.POST.get('meteo_source') or 'temps_reel')[:20]
    if source not in ('temps_reel', 'cache', 'synchronise'):
        source = 'temps_reel'
    return {
        'lat': lat, 'lon': lon,
        'temperature': _f(request.POST.get('meteo_temperature')),
        'conditions': (request.POST.get('meteo_conditions') or '')[:120],
        'code': _i(request.POST.get('meteo_code')),
        'humidite': _i(request.POST.get('meteo_humidite')),
        'vent_kmh': _f(request.POST.get('meteo_vent_kmh')),
        'vent_direction': (request.POST.get('meteo_vent_direction') or '')[:40],
        'vent_direction_deg': _i(request.POST.get('meteo_vent_direction_deg')),
        'proba_pluie': _i(request.POST.get('meteo_proba_pluie')),
        'lever_soleil': request.POST.get('meteo_lever') or None,
        'coucher_soleil': request.POST.get('meteo_coucher') or None,
        'localisation': (request.POST.get('meteo_localisation') or '')[:200],
        'source': source,
        'horodatage': request.POST.get('meteo_horodatage') or None,
    }


def _enregistrer_meteo_activite(activite, donnees):
    """Persiste le snapshot météo d'une activité (OneToOne + update_or_create
    = protection anti-doublons ; ne lève jamais d'exception métier)."""
    from django.utils.dateparse import parse_datetime
    lever = parse_datetime(donnees.get('lever_soleil') or '')
    coucher = parse_datetime(donnees.get('coucher_soleil') or '')
    horo = parse_datetime(donnees.get('horodatage') or '')
    MeteoActivite.objects.update_or_create(
        activite=activite,
        defaults={
            'latitude': donnees.get('lat'),
            'longitude': donnees.get('lon'),
            'localisation': (donnees.get('localisation') or '')[:200],
            'temperature_c': donnees.get('temperature'),
            'conditions': (donnees.get('conditions') or '')[:120],
            'code_conditions': donnees.get('code'),
            'humidite': donnees.get('humidite'),
            'vent_kmh': donnees.get('vent_kmh'),
            'vent_direction': (donnees.get('vent_direction') or '')[:40],
            'vent_direction_deg': donnees.get('vent_direction_deg'),
            'proba_pluie': donnees.get('proba_pluie'),
            'lever_soleil': lever,
            'coucher_soleil': coucher,
            'donnees_disponibles': donnees.get('temperature') is not None,
            'source': donnees.get('source', 'temps_reel'),
            'horodatage_meteo': horo or timezone.now(),
        })


@login_required
def activite_create(request):
    if request.method == "POST":
        projet_id = request.POST.get('projet')
        if not projet_id:
            projet_actif = _projet_actif(request)
            projet_id = projet_actif.pk if projet_actif else ''
        rapport = request.POST.get('rapport')
        nom_activite = request.POST.get('nom_activite', '').strip()
        description = request.POST.get('description', '')
        objectif = request.POST.get('objectif', '')
        resultats = request.POST.get('resultats', '')
        difficultes = request.POST.get('difficultes', '')
        recommandations = request.POST.get('recommandations', '')
        observations = request.POST.get('observations', '')
        nb = request.POST.get('nombre_beneficiaires', 0)
        hommes = request.POST.get('hommes', '')
        femmes = request.POST.get('femmes', '')
        enfants = request.POST.get('enfants', '')
        menages = request.POST.get('menages', '')
        latitude = request.POST.get('latitude')
        longitude = request.POST.get('longitude')
        zone_visitee = request.POST.get('zone_visitee', '')
        agent_id = request.POST.get('agent')
        niveau_securite = ''

        if zone_visitee:
            zones = ZoneSecurite.objects.all()
            for z in zones:
                try:
                    lat, lon = float(latitude), float(longitude)
                    if _point_dans_zone(lat, lon, z):
                        niveau_securite = z.get_statut_display()
                        break
                except (ValueError, TypeError):
                    pass

        if not all([projet_id, rapport, latitude, longitude]):
            messages.error(request, "Tous les champs obligatoires doivent être remplis.")
            return redirect('activite_create')

        agent = None
        if _est_admin(request.user) and agent_id:
            try:
                agent = User.objects.get(pk=int(agent_id))
            except (User.DoesNotExist, ValueError):
                pass
        if agent is None:
            agent = request.user if request.user.is_authenticated else None

        try:
            activite = Activite.objects.create(
                projet_id=int(projet_id),
                agent=agent,
                nom_activite=nom_activite,
                description=description,
                objectif=objectif,
                resultats=resultats,
                difficultes=difficultes,
                recommandations=recommandations,
                rapport=rapport,
                observations=observations,
                nombre_beneficiaires=int(nb) if nb else 0,
                hommes=int(hommes) if str(hommes).isdigit() else None,
                femmes=int(femmes) if str(femmes).isdigit() else None,
                enfants=int(enfants) if str(enfants).isdigit() else None,
                menages=int(menages) if str(menages).isdigit() else None,
                latitude=float(latitude),
                longitude=float(longitude),
                zone_visitee=zone_visitee,
                niveau_securite=niveau_securite,
            )
        except (ValueError, TypeError):
            messages.error(request, "Erreur dans les données saisies.")
            return redirect('activite_create')

        photos = request.FILES.getlist('photos')
        for photo in photos:
            PhotoActivite.objects.create(activite=activite, image=photo)

        # ── Snapshot météo (historisation ; ne bloque jamais la création) ──
        try:
            meteo_donnees = _extraire_meteo_post(request)
            if meteo_donnees is None:
                from .meteo import recuperer_meteo
                meteo_donnees = recuperer_meteo(activite.latitude, activite.longitude)
            if meteo_donnees:
                _enregistrer_meteo_activite(activite, meteo_donnees)
        except Exception:
            pass

        _audit(request, "Création d'activité", f"Activité #{activite.pk} assignée à {agent.username if agent else 'N/A'}")
        messages.success(request, "Activité encodée avec succès.")
        return redirect('dashboard')

    agents_disponibles = User.objects.filter(is_superuser=False, profil__isnull=False)
    return render(request, 'cartographie/activite_form.html', {
        'projets': Projet.objects.all(),
        'agents_list': agents_disponibles,
        'est_admin': _est_admin(request.user),
    })


@login_required
def activite_detail(request, pk):
    activite = get_object_or_404(Activite.objects.select_related('projet', 'agent').prefetch_related('photos'), pk=pk)
    return render(request, 'cartographie/activite_detail.html', {'activite': activite})


@login_required
def activite_delete(request, pk):
    activite = get_object_or_404(Activite, pk=pk)
    if request.method == "POST":
        _audit(request, "Suppression d'activité", f"Activité #{pk}")
        activite.delete()
        messages.success(request, "Activité supprimée.")
        return redirect('dashboard')
    return render(request, 'cartographie/activite_confirm_delete.html', {'activite': activite})


# ─── PROJETS ───────────────────────────────────────────────────


@user_passes_test(_est_admin)
def projet_list(request):
    projets = Projet.objects.annotate(
        nb_activites=db_models.Count('activites'),
        nb_points=db_models.Count('points'),
    ).order_by('-date_creation')
    return render(request, 'cartographie/projet_list.html', {
        'projets': projets,
        'activites_modeles': ActiviteModele.objects.select_related('projet').order_by('projet__nom', 'nom'),
    })


@login_required
def projet_create(request):
    if request.method == "POST":
        nom = request.POST.get('nom')
        but = request.POST.get('but', '')
        description = request.POST.get('description', '')
        if nom:
            p = Projet.objects.create(nom=nom, but=but, description=description, cree_par=request.user)
            _audit(request, "Création de projet", f"Projet #{p.pk} - {nom}")
            messages.success(request, f"Projet « {nom} » créé.")
            return redirect('projet_list')
        messages.error(request, "Le nom du projet est obligatoire.")
    return render(request, 'cartographie/projet_form.html')


@user_passes_test(_est_admin)
def projet_edit(request, pk):
    projet = get_object_or_404(Projet, pk=pk)
    if request.method == "POST":
        nom = request.POST.get('nom')
        if not nom:
            messages.error(request, "Le nom du projet est obligatoire.")
            return redirect('projet_edit', pk)
        projet.nom = nom
        projet.but = request.POST.get('but', '')
        projet.description = request.POST.get('description', '')
        projet.save()
        _audit(request, "Modification de projet", f"Projet #{pk} - {nom}")
        messages.success(request, f"Projet « {nom} » mis à jour.")
        return redirect('projet_list')
    return render(request, 'cartographie/projet_form.html', {'projet': projet, 'edition': True})


@user_passes_test(_est_admin)
def projet_archive(request, pk):
    projet = get_object_or_404(Projet, pk=pk)
    if request.method == "POST":
        projet.statut = 'archive' if projet.statut != 'archive' else 'actif'
        projet.save()
        etat = 'archivé' if projet.statut == 'archive' else 'réactivé'
        _audit(request, f"Projet {etat}", f"Projet #{pk} - {projet.nom}")
        messages.success(request, f"Projet « {projet.nom} » {etat}.")
    return redirect('projet_list')


@user_passes_test(_est_admin)
def activite_modele_create(request, pk):
    projet = get_object_or_404(Projet, pk=pk)
    if request.method == "POST":
        nom = request.POST.get('nom', '').strip()
        if nom:
            ActiviteModele.objects.create(projet=projet, nom=nom[:255], cree_par=request.user)
            _audit(request, "Création d'activité modèle", f"« {nom} » — {projet.nom}")
            messages.success(request, f"Activité modèle « {nom} » créée.")
        else:
            messages.error(request, "Le nom de l'activité modèle est obligatoire.")
    return redirect('projet_list')


@user_passes_test(_est_admin)
def activite_modele_delete(request, pk):
    am = get_object_or_404(ActiviteModele, pk=pk)
    if request.method == "POST":
        _audit(request, "Suppression d'activité modèle", f"« {am.nom} »")
        am.delete()
        messages.success(request, "Activité modèle supprimée.")
    return redirect('projet_list')


def changer_langue(request):
    """Change la langue pour toute l'application (session + cookie) puis revient."""
    from .i18n import LANGUES
    langue = request.POST.get('langue') or request.GET.get('langue') or ''
    if langue not in dict(LANGUES):
        langue = 'fr'
    request.session['langue'] = langue
    referer = request.META.get('HTTP_REFERER') or ''
    if not referer.startswith(request.build_absolute_uri('/')):
        referer = '/'
    response = redirect(referer)
    response.set_cookie('mukmap_langue', langue, max_age=60 * 60 * 24 * 365, samesite='Lax')
    return response


def selection_projet(request):
    """Enregistre le projet + l'activité en cours dans la session (obligatoire avant collecte)."""
    if request.method != "POST":
        return redirect('index_cartographie')
    projet_id = request.POST.get('projet_id')
    nom_activite = request.POST.get('nom_activite', '').strip()
    activite_id = request.POST.get('activite_id', '').strip()

    projet = None
    if projet_id:
        projet = Projet.objects.filter(pk=int(projet_id), statut='actif').first()
    if projet is None:
        messages.error(request, "Veuillez sélectionner un projet valide.")
        return redirect('index_cartographie')
    if not nom_activite:
        messages.error(request, "Veuillez saisir le nom de l'activité à réaliser.")
        return redirect('index_cartographie')

    request.session['projet_actif_id'] = projet.pk
    request.session['activite_actuelle_nom'] = nom_activite[:255]
    if activite_id and activite_id.isdigit():
        request.session['activite_actuelle_id'] = int(activite_id)
    else:
        request.session.pop('activite_actuelle_id', None)
    sid = request.session.get('session_travail_id')
    if sid:
        try:
            session_travail = SessionTravail.objects.get(pk=sid)
            session_travail.projet = projet
            session_travail.activite_nom = nom_activite[:255]
            session_travail.save()
        except SessionTravail.DoesNotExist:
            pass
    _audit(request, "Sélection de projet", f"Projet {projet.nom} — Activité : {nom_activite[:120]}")
    messages.success(request, f"Projet « {projet.nom} » — Activité « {nom_activite} » enregistrée.")
    return redirect('index_cartographie')


def api_projets(request):
    if not request.user.is_authenticated:
        return JsonResponse({'erreur': 'Authentification requise.'}, status=403)
    projets = Projet.objects.filter(statut='actif').values('id', 'nom', 'description')
    return JsonResponse(list(projets), safe=False)


def api_activites_suggestions(request):
    """Suggestions d'activités : activités modèles de l'admin + historique du projet."""
    if not request.user.is_authenticated:
        return JsonResponse({'erreur': 'Authentification requise.'}, status=403)
    try:
        projet_id = int(request.GET.get('projet') or 0)
    except (ValueError, TypeError):
        projet_id = 0
    q = (request.GET.get('q') or '').strip().lower()

    modeles = ActiviteModele.objects.filter(projet_id=projet_id).values_list('nom', flat=True)
    historiques = Activite.objects.filter(projet_id=projet_id)
    if q:
        historiques = historiques.filter(nom_activite__icontains=q)
    historiques = list(historiques.exclude(nom_activite='').order_by('nom_activite').distinct().values_list('nom_activite', flat=True))

    noms = list(modeles) + [h for h in historiques if h not in modeles]
    if q:
        noms = [n for n in noms if q in n.lower()]
    return JsonResponse(noms[:25], safe=False)


# ─── IMPORT FICHIER (POINTS: GeoJSON/KML) ─────────────────────


@login_required
def importer_fichier(request):
    if request.method != "POST":
        return redirect('index_cartographie')
    est_ajax = request.POST.get('ajax') == '1' or request.GET.get('ajax') == '1'
    fichier = request.FILES.get('fichier_import')
    if not fichier:
        if est_ajax:
            return JsonResponse({'ok': False, 'erreur': "Aucun fichier sélectionné."}, status=400)
        messages.error(request, "Aucun fichier sélectionné.")
        return redirect('index_cartographie')

    nom_fichier = fichier.name.lower()
    contenu = fichier.read()

    try:
        if nom_fichier.endswith('.geojson'):
            points = _parser_geojson(contenu)
        elif nom_fichier.endswith('.kml'):
            points = _parser_kml(contenu)
        else:
            if est_ajax:
                return JsonResponse({'ok': False, 'erreur': "Format non supporté. Utilisez .geojson ou .kml."}, status=400)
            messages.error(request, "Format non supporté. Utilisez .geojson ou .kml.")
            return redirect('index_cartographie')
    except Exception:
        if est_ajax:
            return JsonResponse({'ok': False, 'erreur': "Fichier invalide ou corrompu."}, status=400)
        messages.error(request, "Fichier invalide ou corrompu.")
        return redirect('index_cartographie')

    if not points:
        if est_ajax:
            return JsonResponse({'ok': False, 'erreur': "Aucun point valide trouvé dans le fichier."}, status=400)
        messages.warning(request, "Aucun point valide trouvé.")
        return redirect('index_cartographie')

    inserer = 0
    for p in points:
        try:
            proprietes = p.get('proprietes') or {}
            lat, lng = _valider_coordonnees_wgs84(p.get('latitude'), p.get('longitude'))
            if lat is None:
                continue
            source_format = nom_fichier.rsplit('.', 1)[-1].upper() if '.' in nom_fichier else ''
            PointGeographique.objects.create(
                nom=(p.get('nom') or 'Sans nom')[:200],
                description=p.get('description', ''),
                latitude=lat,
                longitude=lng,
                donnees=proprietes,
                source_fichier=fichier.name,
                source_format=source_format,
                auteur=request.user if request.user.is_authenticated else None,
            )
            inserer += 1
        except (ValueError, TypeError, KeyError):
            continue

    _audit(request, "Import fichier", f"{inserer} points depuis {nom_fichier}")
    if est_ajax:
        return JsonResponse({'ok': True, 'importes': inserer, 'fichier': fichier.name})
    messages.success(request, f"{inserer} point(s) importé(s) avec succès.")
    return redirect('index_cartographie')


def _parser_geojson(contenu):
    data = json.loads(contenu)
    features = data.get('features', [])
    resultats = []
    for feature in features:
        geom = feature.get('geometry', {})
        if geom.get('type') != 'Point':
            continue
        coords = geom.get('coordinates', [])
        if len(coords) < 2:
            continue
        props = feature.get('properties', {}) or {}
        proprietes = {}
        for k, v in props.items():
            if v is None:
                v = ''
            if isinstance(v, (dict, list)):
                try:
                    v = json.dumps(v, ensure_ascii=False)
                except (TypeError, ValueError):
                    v = str(v)
            proprietes[str(k)] = str(v)
        resultats.append({
            'nom': props.get('nom') or props.get('name') or props.get('title'),
            'description': props.get('description') or props.get('desc') or '',
            'longitude': coords[0], 'latitude': coords[1],
            'proprietes': proprietes,
        })
    return resultats


def _parser_kml(contenu):
    root = ET.fromstring(contenu)
    ns = {'kml': 'http://www.opengis.net/kml/2.2'}
    resultats = []
    placemarks = root.findall('.//kml:Placemark', ns) or root.findall('.//Placemark')
    for pm in placemarks:
        nom_el = pm.find('kml:name', ns)
        if nom_el is None:
            nom_el = pm.find('name')
        desc_el = pm.find('kml:description', ns)
        if desc_el is None:
            desc_el = pm.find('description')
        coord_el = pm.find('.//kml:coordinates', ns)
        if coord_el is None:
            coord_el = pm.find('.//coordinates')
        nom = nom_el.text.strip() if nom_el is not None and nom_el.text else None
        description = desc_el.text.strip() if desc_el is not None and desc_el.text else ''
        proprietes = {}
        if nom:
            proprietes['Nom'] = nom
        if description:
            proprietes['Description'] = description
        ext_el = pm.find('kml:ExtendedData', ns)
        if ext_el is None:
            ext_el = pm.find('ExtendedData')
        if ext_el is not None:
            for data_el in ext_el.findall('kml:Data', ns) + ext_el.findall('Data'):
                k = data_el.get('name')
                val_el = data_el.find('kml:value', ns)
                if val_el is None:
                    val_el = data_el.find('value')
                if k and val_el is not None and val_el.text is not None:
                    proprietes[k] = val_el.text.strip()
            for sd_el in ext_el.findall('kml:SimpleData', ns) + ext_el.findall('SimpleData'):
                k = sd_el.get('name')
                if k and sd_el.text is not None:
                    proprietes[k] = sd_el.text.strip()
        if coord_el is None or not coord_el.text:
            continue
        parts = coord_el.text.strip().split(',')
        if len(parts) < 2:
            continue
        try:
            longitude, latitude = float(parts[0]), float(parts[1])
        except ValueError:
            continue
        resultats.append({'nom': nom, 'description': description, 'longitude': longitude, 'latitude': latitude, 'proprietes': proprietes})
    return resultats


# ─── IMPORT EXCEL → KML ────────────────────────────────────────


@login_required
def importer_excel(request):
    if request.method != "POST":
        return redirect('index_cartographie')

    fichier = request.FILES.get('fichier_excel')
    if not fichier or not fichier.name.lower().endswith('.xlsx'):
        messages.error(request, "Veuillez fournir un fichier Excel (.xlsx).")
        return redirect('index_cartographie')

    import openpyxl

    try:
        wb = openpyxl.load_workbook(io.BytesIO(fichier.read()))
        ws = wb.active
        rows = list(ws.iter_rows(values_only=True))
    except Exception:
        messages.error(request, "Impossible de lire le fichier Excel.")
        return redirect('index_cartographie')

    if not rows or len(rows) < 2:
        messages.error(request, "Le fichier est vide.")
        return redirect('index_cartographie')

    headers = [str(h).lower().strip() if h else '' for h in rows[0]]
    idx_nom = next((i for i, h in enumerate(headers) if 'nom' in h), None)
    idx_lat = next((i for i, h in enumerate(headers) if 'lat' in h), None)
    idx_lng = next((i for i, h in enumerate(headers) if 'lon' in h or 'lng' in h), None)
    idx_desc = next((i for i, h in enumerate(headers) if 'desc' in h), None)

    if idx_lat is None or idx_lng is None:
        messages.error(request, "Colonnes Latitude/Longitude introuvables.")
        return redirect('index_cartographie')

    nom_couche = f"Excel_{date.today().strftime('%Y%m%d_%H%M%S')}"
    couche = CoucheGeometrie.objects.create(
        nom=nom_couche, type_geometrie='point', fichier_source=fichier.name,
        projet=_projet_actif(request),
    )

    kml_parts = [
        '<?xml version="1.0" encoding="UTF-8"?>',
        '<kml xmlns="http://www.opengis.net/kml/2.2"><Document>',
        f'<name>{nom_couche}</name>',
    ]

    points_crees = 0
    for row in rows[1:]:
        if not row or not any(row):
            continue
        try:
            lat = float(row[idx_lat]) if row[idx_lat] is not None else None
            lng = float(row[idx_lng]) if row[idx_lng] is not None else None
        except (ValueError, TypeError):
            continue
        if lat is None or lng is None:
            continue

        nom = str(row[idx_nom])[:200] if idx_nom is not None and row[idx_nom] else 'Sans nom'
        desc = str(row[idx_desc]) if idx_desc is not None and row[idx_desc] else ''

        PointGeographique.objects.create(
            nom=nom, description=desc, latitude=lat, longitude=lng,
            projet=_projet_actif(request),
            auteur=request.user if request.user.is_authenticated else None,
        )

        geo = Geometrie.objects.create(
            couche=couche, type='Point',
            coordonnees=[lng, lat],
            proprietes={'nom': nom, 'description': desc},
        )

        kml_parts.append('<Placemark>')
        kml_parts.append(f'<name>{_xml_safe(nom)}</name>')
        kml_parts.append(f'<description>{_xml_safe(desc)}</description>')
        kml_parts.append(f'<Point><coordinates>{lng},{lat},0</coordinates></Point>')
        kml_parts.append('</Placemark>')
        points_crees += 1

    kml_parts.append('</Document></kml>')
    kml_content = '\n'.join(kml_parts)

    # Sauvegarder le KML sur le disque
    kml_filename = f"{nom_couche}.kml"
    kml_dir = os.path.join(settings.MEDIA_ROOT, 'kml_imports')
    os.makedirs(kml_dir, exist_ok=True)
    kml_path = os.path.join(kml_dir, kml_filename)
    with open(kml_path, 'w', encoding='utf-8') as f:
        f.write(kml_content)

    couche.fichier_kml = f'kml_imports/{kml_filename}'
    couche.save()

    _audit(request, "Import Excel", f"{points_crees} points depuis {fichier.name}")
    if points_crees > 0:
        messages.success(request, f"{points_crees} point(s) importés depuis Excel et convertis en KML.")
    else:
        messages.warning(request, "Aucun point valide trouvé dans le fichier.")
    return redirect('index_cartographie')


# ─── ZONES DE SÉCURITÉ ─────────────────────────────────────────


@user_passes_test(_est_admin)
def zone_list(request):
    statut = request.GET.get('statut', '')
    zones = ZoneSecurite.objects.select_related('auteur', 'modifie_par').all()
    projet_actif = _projet_actif(request)
    if projet_actif is not None:
        zones = zones.filter(projet=projet_actif)
    if statut:
        zones = zones.filter(statut=statut)
    return render(request, 'cartographie/zone_list.html', {
        'zones': zones, 'statut_filtre': statut,
    })


@user_passes_test(_est_admin)
def zone_create(request):
    if request.method == "POST":
        nom = request.POST.get('nom')
        statut = request.POST.get('statut')
        motif = request.POST.get('motif', '')
        rayon = request.POST.get('rayon', 0)
        type_geom = request.POST.get('type_geometrie', 'Point')
        coords_raw = request.POST.get('coordonnees', '[]')

        if not all([nom, statut, coords_raw]):
            messages.error(request, "Nom, statut et coordonnées requis.")
            return redirect('zone_create')

        try:
            coords = json.loads(coords_raw)
        except json.JSONDecodeError:
            messages.error(request, "Coordonnées invalides.")
            return redirect('zone_create')

        zone = ZoneSecurite.objects.create(
            nom=nom, statut=statut, motif=motif if statut == 'dangereuse' else '',
            type_geometrie=type_geom, coordonnees=coords,
            rayon=float(rayon), auteur=request.user,
            projet=_projet_actif(request),
        )
        _audit(request, "Déclaration de zone", f"Zone #{zone.pk} - {nom} ({statut})")
        messages.success(request, f"Zone « {nom} » déclarée comme {dict(ZoneSecurite.STATUT_CHOICES).get(statut)}.")
        return redirect('zone_list')

    return render(request, 'cartographie/zone_form.html')


@user_passes_test(_est_admin)
def zone_edit(request, pk):
    zone = get_object_or_404(ZoneSecurite, pk=pk)
    if request.method == "POST":
        zone.nom = request.POST.get('nom', zone.nom)
        zone.statut = request.POST.get('statut', zone.statut)
        zone.motif = request.POST.get('motif', '')
        zone.rayon = float(request.POST.get('rayon', zone.rayon))
        zone.modifie_par = request.user
        zone.save()
        _audit(request, "Modification de zone", f"Zone #{pk}")
        messages.success(request, "Zone mise à jour.")
        return redirect('zone_list')
    return render(request, 'cartographie/zone_form.html', {'zone': zone, 'edition': True})


@user_passes_test(_est_admin)
def zone_delete(request, pk):
    zone = get_object_or_404(ZoneSecurite, pk=pk)
    if request.method == "POST":
        _audit(request, "Suppression de zone", f"Zone #{pk} - {zone.nom}")
        zone.delete()
        messages.success(request, "Zone supprimée.")
        return redirect('zone_list')
    return render(request, 'cartographie/zone_confirm_delete.html', {'zone': zone})


# ─── ITINÉRAIRES ────────────────────────────────────────────────


@login_required
def itineraire_list(request):
    iti = Itineraire.objects.filter(utilisateur=request.user).all()
    return render(request, 'cartographie/itineraire_list.html', {'itineraire': iti})


@login_required
def itineraire_create(request):
    if request.method == "POST":
        nom = request.POST.get('nom', '')
        coords_raw = request.POST.get('coordonnees', '[]')
        if not nom:
            messages.error(request, "Nom de l'itinéraire requis.")
            return redirect('itineraire_create')
        try:
            coords = json.loads(coords_raw)
        except json.JSONDecodeError:
            messages.error(request, "Coordonnées invalides.")
            return redirect('itineraire_create')
        if len(coords) < 2:
            messages.error(request, "Tracez au moins 2 points sur la carte.")
            return redirect('itineraire_create')

        analyse, alertes = _analyser_itineraire(coords)
        alerte_msg = '\n'.join(alertes) if alertes else ''

        iti = Itineraire.objects.create(
            utilisateur=request.user, nom=nom,
            coordonnees=coords, analyse=analyse, alerte=alerte_msg,
        )
        _audit(request, "Création d'itinéraire", f"Itinéraire #{iti.pk} - {nom}")
        if alertes:
            messages.warning(request, alerte_msg)
        messages.success(request, f"Itinéraire « {nom} » enregistré.")
        return redirect('itineraire_list')

    return render(request, 'cartographie/itineraire_form.html')


@login_required
def itineraire_detail(request, pk):
    iti = get_object_or_404(Itineraire, pk=pk, utilisateur=request.user)
    return render(request, 'cartographie/itineraire_detail.html', {'iti': iti})


@login_required
def itineraire_delete(request, pk):
    iti = get_object_or_404(Itineraire, pk=pk, utilisateur=request.user)
    if request.method == "POST":
        _audit(request, "Suppression d'itinéraire", f"Itinéraire #{pk}")
        iti.delete()
        messages.success(request, "Itinéraire supprimé.")
        return redirect('itineraire_list')
    return render(request, 'cartographie/itineraire_confirm_delete.html', {'iti': iti})


# ─── JOURNAL D'AUDIT ────────────────────────────────────────────


@user_passes_test(_est_admin)
def audit_list(request):
    logs = JournalAudit.objects.select_related('utilisateur').all()
    user_id = request.GET.get('user')
    if user_id:
        logs = logs.filter(utilisateur_id=user_id)
    return render(request, 'cartographie/audit_list.html', {
        'logs': logs,
        'users': ProfilAgent.objects.select_related('utilisateur').all(),
        'user_filtre': int(user_id) if user_id else None,
    })


# ─── GESTION DES AGENTS (ADMIN) ─────────────────────────────────


@user_passes_test(_est_admin)
def agent_list(request):
    profils = ProfilAgent.objects.select_related('utilisateur').all()
    return render(request, 'cartographie/agent_list.html', {'profils': profils})


@user_passes_test(_est_admin)
def agent_create(request):
    if request.method == "POST":
        nom_complet = request.POST.get('nom_complet', '').strip()
        telephone = request.POST.get('telephone', '').strip()
        titre = request.POST.get('titre', '').strip()
        email = request.POST.get('email', '').strip()
        username = request.POST.get('username', '').strip()
        password = request.POST.get('password', '').strip()
        latitude = request.POST.get('latitude', '').strip()
        longitude = request.POST.get('longitude', '').strip()

        if not all([nom_complet, telephone, titre, email, username, password]):
            messages.error(request, "Tous les champs obligatoires doivent être remplis.")
            return render(request, 'cartographie/agent_form.html')

        from django.contrib.auth.password_validation import validate_password
        from django.core.exceptions import ValidationError as _ValErr
        try:
            validate_password(password, user=None)
        except _ValErr as e:
            messages.error(request, "Mot de passe trop faible : " + " ".join(e.messages))
            return render(request, 'cartographie/agent_form.html')

        if User.objects.filter(username=username).exists():
            messages.error(request, f"L'identifiant « {username} » est déjà utilisé.")
            return render(request, 'cartographie/agent_form.html')

        parts = nom_complet.split(None, 1)
        first_name = parts[0]
        last_name = parts[1] if len(parts) > 1 else ''

        user = User.objects.create_user(
            username=username,
            email=email,
            password=password,
            first_name=first_name,
            last_name=last_name,
        )

        lat_val = float(latitude) if latitude else None
        lng_val = float(longitude) if longitude else None

        ProfilAgent.objects.create(
            utilisateur=user,
            telephone=telephone,
            fonction=titre,
            latitude=lat_val,
            longitude=lng_val,
        )
        _audit(request, "Création de compte agent", f"Agent {nom_complet} ({username})")
        messages.success(
            request,
            f"Agent « {nom_complet} » créé avec succès. Identifiant : {username}"
        )
        return redirect('agent_list')

    return render(request, 'cartographie/agent_form.html')


@user_passes_test(_est_admin)
def agent_bloquer(request, pk):
    profil = get_object_or_404(ProfilAgent, pk=pk)
    profil.est_bloque = not profil.est_bloque
    profil.save()
    etat = "bloqué" if profil.est_bloque else "débloqué"
    _audit(request, f"Agent {etat}", f"Agent #{pk} - {profil.utilisateur.username}")
    messages.success(request, f"Agent {etat}.")
    return redirect('agent_list')


# ─── RAPPORTS ────────────────────────────────────────────────────


TOUTES_SECTIONS_RAPPORT = (
    'infos', 'stats', 'activites', 'beneficiaires', 'agents',
    'terrain', 'itineraire', 'points', 'zones', 'dangers',
    'photos', 'observations', 'recommandations', 'audits',
    'conditions_meteo',
)


def _filtres_rapport(request, user):
    """Analyse les filtres de l'assistant de rapport (période, projet, activités, agent, zone, sections)."""
    est_admin = _est_admin(user)
    aujourd = timezone.localdate()

    type_ = request.GET.get('type', 'mensuel')
    mapping = {
        'journalier': (aujourd, aujourd),
        'hebdomadaire': (aujourd - timedelta(days=7), aujourd),
        'mensuel': (aujourd - timedelta(days=30), aujourd),
        'annuel': (aujourd - timedelta(days=365), aujourd),
    }
    if type_ == 'personnalise':
        try:
            debut = date.fromisoformat(request.GET.get('date_debut', ''))
            fin = date.fromisoformat(request.GET.get('date_fin', ''))
        except ValueError:
            debut, fin = mapping['mensuel']
    else:
        debut, fin = mapping.get(type_, mapping['mensuel'])
    if fin < debut:
        debut, fin = fin, debut

    if est_admin:
        projets = Projet.objects.all().order_by('nom')
    else:
        ids_miens = set(Activite.objects.filter(agent=user).values_list('projet_id', flat=True))
        ids_miens |= set(Projet.objects.filter(cree_par=user).values_list('pk', flat=True))
        projets = Projet.objects.filter(pk__in=ids_miens).order_by('nom')

    pid = None
    val_projet = request.GET.get('projet', '') or ''
    if val_projet.isdigit():
        pid = int(val_projet)
        if not projets.filter(pk=pid).exists():
            pid = None
    projet = Projet.objects.filter(pk=pid).first() if pid else None

    qs_dispo = Activite.objects.filter(date_creation__date__gte=debut, date_creation__date__lte=fin)
    if pid:
        qs_dispo = qs_dispo.filter(projet_id=pid)
    if not est_admin:
        qs_dispo = qs_dispo.filter(agent=user)
    activites_disponibles = qs_dispo.select_related('projet').order_by('-date_creation')

    activite_ids = None
    val_act = request.GET.getlist('activites')
    if val_act and val_act != ['toutes']:
        ids = [int(x) for x in val_act if x.strip().isdigit()]
        if ids:
            dispo_ids = set(activites_disponibles.values_list('pk', flat=True))
            activite_ids = [i for i in ids if i in dispo_ids] or None

    agent_id = None
    if est_admin:
        val_agent = request.GET.get('agent', '') or ''
        if val_agent.isdigit():
            agent_id = int(val_agent)
    else:
        agent_id = user.pk

    zone = request.GET.get('zone', '') or ''
    mode_global = bool(request.GET.get('global')) and est_admin

    val_sec = request.GET.getlist('sections')
    if val_sec and val_sec != ['toutes']:
        sections = [s for s in val_sec if s in TOUTES_SECTIONS_RAPPORT] or list(TOUTES_SECTIONS_RAPPORT)
    else:
        sections = list(TOUTES_SECTIONS_RAPPORT)

    return {
        'type': type_, 'debut': debut, 'fin': fin,
        'projet': projet, 'projet_id': pid,
        'activite_ids': activite_ids, 'activites_disponibles': activites_disponibles,
        'agent_id': agent_id, 'zone': zone, 'mode_global': mode_global,
        'sections': sections, 'est_admin': est_admin, 'projets': projets,
    }


def _donnees_rapport_v2(f):
    """Construit le contexte complet d'un rapport à partir des filtres analysés."""
    debut, fin = f['debut'], f['fin']
    sections = set(f['sections'])

    qs = Activite.objects.filter(date_creation__date__gte=debut, date_creation__date__lte=fin)
    if f['projet_id']:
        qs = qs.filter(projet_id=f['projet_id'])
    if f['agent_id']:
        qs = qs.filter(agent_id=f['agent_id'])
    if f['zone']:
        qs = qs.filter(zone_visitee__icontains=f['zone'])
    if f['activite_ids']:
        qs = qs.filter(pk__in=f['activite_ids'])
    activites = list(qs.select_related('projet', 'agent').prefetch_related('photos').order_by('date_creation'))

    total = len(activites)
    agr = qs.aggregate(
        bene=db_models.Sum('nombre_beneficiaires'),
        hommes=db_models.Sum('hommes'),
        femmes=db_models.Sum('femmes'),
        enfants=db_models.Sum('enfants'),
        menages=db_models.Sum('menages'),
    )
    bene_total = agr['bene'] or 0
    tot_h = agr['hommes'] or 0
    tot_f = agr['femmes'] or 0
    tot_e = agr['enfants'] or 0
    tot_m = agr['menages'] or 0

    projet_ids = [a.projet_id for a in activites]
    user_ids = [a.agent_id for a in activites if a.agent_id]

    points_qs = PointGeographique.objects.filter(
        supprime=False, date_creation__date__gte=debut, date_creation__date__lte=fin)
    if projet_ids:
        points_qs = points_qs.filter(projet_id__in=projet_ids)
    if f['agent_id']:
        points_qs = points_qs.filter(auteur_id=f['agent_id'])
    points = list(points_qs.select_related('projet', 'auteur').order_by('-date_creation'))

    iti_qs = Itineraire.objects.filter(date_creation__date__gte=debut, date_creation__date__lte=fin)
    if projet_ids:
        iti_qs = iti_qs.filter(projet_id__in=projet_ids)
    if f['agent_id']:
        iti_qs = iti_qs.filter(utilisateur_id=f['agent_id'])
    itineraires = list(iti_qs.select_related('projet', 'utilisateur').order_by('-date_creation'))

    zones_qs = ZoneSecurite.objects.select_related('projet', 'auteur').filter(
        date_declaration__date__gte=debut, date_declaration__date__lte=fin)
    if f['projet_id']:
        zones_qs = zones_qs.filter(projet_id=f['projet_id'])
    elif projet_ids:
        zones_qs = zones_qs.filter(projet_id__in=projet_ids)
    zones = list(zones_qs.order_by('-date_declaration'))
    zd = sum(1 for z in zones if z.statut == 'dangereuse')
    zs = sum(1 for z in zones if z.statut == 'securisee')
    zi = sum(1 for z in zones if z.statut == 'indisponible')

    profils_qs = ProfilAgent.objects.select_related('utilisateur')
    if user_ids:
        profils_qs = profils_qs.filter(utilisateur_id__in=user_ids)
    profils = list(profils_qs.order_by('utilisateur__username'))

    ses_qs = SessionTravail.objects.select_related('utilisateur', 'projet').filter(
        debut__date__gte=debut, debut__date__lte=fin)
    if f['projet_id']:
        ses_qs = ses_qs.filter(projet_id=f['projet_id'])
    if f['agent_id']:
        ses_qs = ses_qs.filter(utilisateur_id=f['agent_id'])
    sessions = list(ses_qs.order_by('-debut'))

    audits = []
    if f['mode_global']:
        audits = list(JournalAudit.objects.filter(
            utilisateur__is_superuser=True, date__date__gte=debut, date__date__lte=fin)
            .select_related('utilisateur').order_by('-date'))

    par_projet = []
    for pid in sorted(set(projet_ids)):
        sub = [a for a in activites if a.projet_id == pid]
        par_projet.append({
            'projet': sub[0].projet,
            'count': len(sub),
            'benef': sum(a.nombre_beneficiaires or 0 for a in sub),
        })
    par_projet.sort(key=lambda x: x['count'], reverse=True)

    cat_noms = dict(PointGeographique.CATEGORIE_CHOICES)
    par_categorie = {}
    for p in points:
        par_categorie[cat_noms.get(p.categorie, p.categorie)] = par_categorie.get(cat_noms.get(p.categorie, p.categorie), 0) + 1

    nb_photos = sum(a.photos.count() for a in activites)

    meteo_par_activite = {}
    if 'conditions_meteo' in sections and activites:
        for m in MeteoActivite.objects.select_related('activite').filter(
                activite_id__in=[a.pk for a in activites]):
            meteo_par_activite[m.activite_id] = m

    graph_activites = {'labels': [x['projet'].nom for x in par_projet], 'values': [x['count'] for x in par_projet]}
    graph_benef = {'labels': [x['projet'].nom for x in par_projet], 'values': [x['benef'] for x in par_projet]}
    graph_zones = {'labels': ['dangereuse', 'securisee', 'indisponible'], 'values': [zd, zs, zi]}
    graph_categories = {'labels': list(par_categorie.keys()), 'values': list(par_categorie.values())}

    intitule_projet = f['projet'].nom if f['projet'] else None
    ref = (f['projet'].code or f['projet'].nom if f['projet'] else 'global')
    nom_fichier = f"rapport_{re.sub(r'[^A-Za-z0-9]+', '_', str(ref)).strip('_')}_{debut.strftime('%Y%m%d')}_{fin.strftime('%Y%m%d')}"

    return {
        'f': f, 'debut': debut, 'fin': fin, 'type': f['type'],
        'projet': f['projet'], 'projet_id': f['projet_id'], 'zone': f['zone'],
        'est_admin': f['est_admin'], 'mode_global': f['mode_global'],
        'sections': list(sections),
        'projets': f['projets'], 'activites_disponibles': f['activites_disponibles'],
        'activites': activites, 'total': total,
        'bene_total': bene_total, 'tot_h': tot_h, 'tot_f': tot_f, 'tot_e': tot_e, 'tot_m': tot_m,
        'zones': zones, 'zones_dangereuses': zd, 'zones_securisees': zs, 'zones_indisponibles': zi,
        'total_zones': len(zones),
        'points': points, 'total_points': len(points),
        'itineraires': itineraires, 'total_itineraires': len(itineraires),
        'sessions': sessions, 'total_sessions': len(sessions),
        'profils': profils, 'total_agents': len(profils),
        'audits': audits,
        'par_projet': par_projet, 'par_categorie': par_categorie,
        'nb_photos': nb_photos,
        'meteo_par_activite': meteo_par_activite,
        'graph_activites': _json_script(graph_activites),
        'graph_benef': _json_script(graph_benef),
        'graph_zones': _json_script(graph_zones),
        'graph_categories': _json_script(graph_categories),
        'nom_fichier': nom_fichier,
        'intitule_projet': intitule_projet,
    }


@login_required
def rapport_generer(request):
    from .i18n import traduire
    f = _filtres_rapport(request, request.user)
    ctx = _donnees_rapport_v2(f)
    lang = langue_active(request)
    t = lambda cle: traduire(lang, cle)
    gz = json.loads(ctx['graph_zones'])
    gz['labels'] = [t('zone_dangereuse'), t('zone_securisee'), t('zone_indisponible')]
    ctx['graph_zones'] = _json_script(gz)
    if not ctx['intitule_projet']:
        ctx['intitule_projet'] = t('tous_projets')
    ctx['type_choices'] = [(c, t('rapport_' + c)) for c in ('journalier', 'hebdomadaire', 'mensuel', 'annuel', 'personnalise')]
    ctx['section_choices'] = [(c, t('sec_' + c)) for c in TOUTES_SECTIONS_RAPPORT]
    ctx['date_debut_perso'] = f['debut'].isoformat() if f['type'] == 'personnalise' else ''
    ctx['date_fin_perso'] = f['fin'].isoformat() if f['type'] == 'personnalise' else ''
    val_etape = request.GET.get('etape', '1')
    ctx['etape'] = int(val_etape) if val_etape.isdigit() and 1 <= int(val_etape) <= 6 else 1
    ctx['tous_agents'] = User.objects.filter(is_superuser=False, profil__isnull=False).order_by('username')
    agent_selectionne = ''
    if f['agent_id']:
        try:
            u = User.objects.get(pk=f['agent_id'])
            agent_selectionne = u.get_full_name() or u.username
        except User.DoesNotExist:
            pass
    ctx['agent_selectionne'] = agent_selectionne
    return render(request, 'cartographie/rapport.html', ctx)


@login_required
def rapport_telecharger(request, format):
    f = _filtres_rapport(request, request.user)
    ctx = _donnees_rapport_v2(f)
    lang = langue_active(request)
    nom = ctx['nom_fichier']
    if format == 'docx':
        return _generer_docx(ctx, nom, lang)
    if format == 'pdf':
        return _generer_pdf(ctx, nom, lang)
    if format == 'xlsx':
        return _generer_excel(ctx, nom, lang)
    messages.error(request, "Format non supporté.")
    return redirect('rapport_generer')


# ─── PWA : manifest + service worker ──────────────────────────────

def manifest_pwa(request):
    """Manifeste PWA (URL /manifest.webmanifest, MIME correct)."""
    from django.http import JsonResponse
    from django.templatetags.static import static
    from .branding import COULEUR_FOND, COULEUR_THEME, DEVELOPPEUR, NOM, TAGLINE_FR, VERSION
    manifest = {
        'name': 'MUKMAP — Plateforme SIG professionnelle',
        'short_name': NOM,
        'description': 'MUKMAP — ' + TAGLINE_FR + ' : collecte de données géospatiales, cartographie, suivi de terrain, zones de sécurité et rapports pour entreprises, ONG et ingénieurs topographes.',
        'id': '/',
        'start_url': '/',
        'scope': '/',
        'display': 'standalone',
        'orientation': 'any',
        'background_color': COULEUR_FOND,
        'theme_color': COULEUR_THEME,
        'lang': langue_active(request),
        'categories': ['business', 'productivity', 'utilities', 'navigation'],
        'icons': [
            {'src': static('pwa/icon-192.png'), 'sizes': '192x192', 'type': 'image/png', 'purpose': 'any'},
            {'src': static('pwa/icon-512.png'), 'sizes': '512x512', 'type': 'image/png', 'purpose': 'any'},
            {'src': static('pwa/icon-maskable-192.png'), 'sizes': '192x192', 'type': 'image/png', 'purpose': 'maskable'},
            {'src': static('pwa/icon-maskable-512.png'), 'sizes': '512x512', 'type': 'image/png', 'purpose': 'maskable'},
        ],
        'shortcuts': [
            {'name': 'Cartographie', 'short_name': 'Carte', 'url': '/', 'icons': [{'src': static('pwa/icon-192.png'), 'sizes': '192x192'}]},
            {'name': 'Tableau de bord', 'short_name': 'Dashboard', 'url': '/dashboard/', 'icons': [{'src': static('pwa/icon-192.png'), 'sizes': '192x192'}]},
            {'name': 'Rapports', 'short_name': 'Rapports', 'url': '/rapport/', 'icons': [{'src': static('pwa/icon-192.png'), 'sizes': '192x192'}]},
        ],
    }
    response = JsonResponse(manifest, json_dumps_params={'ensure_ascii': False, 'indent': 2})
    response['Content-Type'] = 'application/manifest+json'
    response['Cache-Control'] = 'no-cache'
    return response


def service_worker_pwa(request):
    """Sert le Service Worker à la racine (portée /) avec le bon MIME."""
    from django.http import FileResponse
    from django.views.decorators.cache import never_cache
    chemin = os.path.join(str(settings.STATICFILES_DIRS[0]), 'js', 'sw.js')
    if not os.path.exists(chemin):
        from django.http import HttpResponse
        return HttpResponse(status=404)
    response = FileResponse(open(chemin, 'rb'), content_type='application/javascript')
    response['Service-Worker-Allowed'] = '/'
    response['Cache-Control'] = 'no-cache'
    return response


def _generer_docx(ctx, nom_fichier, lang='fr'):
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from django.utils import timezone
    from .branding import chemin_logo, DEVELOPPEUR, NOM, VERSION
    from .i18n import traduire

    t = lambda cle: traduire(lang, cle)
    maintenant = timezone.now()
    debut, fin = ctx['debut'], ctx['fin']
    activites = ctx['activites']
    total = ctx['total']
    zd = ctx['zones_dangereuses']
    zs = ctx['zones_securisees']
    zi = ctx['zones_indisponibles']
    zones = ctx['zones']
    audits = ctx['audits']
    profils = ctx['profils']
    bene_total = ctx['bene_total']
    sections = set(ctx['sections'])
    doc = Document()

    def nom_agent(u):
        return (u.get_full_name() or u.username) if u else 'N/A'

    meteo_par_activite = ctx.get('meteo_par_activite', {})

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    ACCENT = RGBColor(0x4F, 0x46, 0xE5)
    GRIS = RGBColor(0x6B, 0x72, 0x9C)

    # ── Page de couverture ──────────────────────────────────────
    logo_chemin = chemin_logo()
    if os.path.exists(logo_chemin):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(logo_chemin, width=Inches(1.4))
    titre = doc.add_heading('', 0)
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = titre.add_run(f'{NOM} v{VERSION} — {t("plateforme_sig")}')
    r.font.color.rgb = ACCENT
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(t('couverture'))
    r.font.size = Pt(16)
    r.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(' ').add_break()
    table_info = doc.add_table(rows=6, cols=2)
    table_info.alignment = 1
    table_info.style = 'Light Shading Accent 1'
    intitule = ctx['intitule_projet'] or t('tous_projets')
    for i, (lib, val) in enumerate([(t('periode_rapport'), f"{debut.strftime('%d/%m/%Y')} — {fin.strftime('%d/%m/%Y')}"),
                                    (t('type_rapport'), t('rapport_' + ctx['type'])),
                                    (t('informations_projet'), intitule),
                                    (t('reference'), nom_fichier),
                                    (t('genere_le'), f"{maintenant.strftime('%d/%m/%Y')} {t('heure')} {maintenant.strftime('%H:%M')}"),
                                    (t('developpe_par'), DEVELOPPEUR)]):
        table_info.cell(i, 0).text = lib
        table_info.cell(i, 1).text = val
    doc.add_page_break()

    # ── En-tête de contenu ──────────────────────────────────────
    doc.add_heading(f'{NOM} — {t("plateforme_sig")}', 0)
    doc.add_heading(f"{t('rapport_activites_periode')} — {debut.strftime('%d/%m/%Y')} {t('au')} {fin.strftime('%d/%m/%Y')}", level=1)
    p = doc.add_paragraph()
    r = p.add_run(f"{t('genere_le')} {maintenant.strftime('%d/%m/%Y')} {t('heure')} {maintenant.strftime('%H:%M')} — {NOM} v{VERSION}, {t('developpe_par')} {DEVELOPPEUR}")
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = GRIS

    doc.add_heading(t('resume_periode'), level=2)
    resume = doc.add_table(rows=2, cols=5)
    resume.style = 'Light Shading Accent 1'
    for i, (lib, val) in enumerate([(t('activites'), total), (t('beneficiaires'), bene_total),
                                    (t('zones_dangereuses'), zd), (t('zones_securisees'), zs),
                                    (t('zones_sans_info'), zi)]):
        resume.cell(0, i).text = lib
        resume.cell(1, i).text = str(val)

    if 'stats' in sections:
        doc.add_heading(t('kpis_rapport'), level=3)
        kpis = doc.add_table(rows=2, cols=5)
        kpis.style = 'Light Shading Accent 1'
        for i, (lib, val) in enumerate([(t('kpi_hommes'), ctx['tot_h']), (t('kpi_femmes'), ctx['tot_f']),
                                        (t('kpi_enfants'), ctx['tot_e']), (t('kpi_menages'), ctx['tot_m']),
                                        (t('total_agents'), ctx['total_agents'])]):
            kpis.cell(0, i).text = lib
            kpis.cell(1, i).text = str(val)
        for i, (lib, val) in enumerate([(t('points_visites'), ctx['total_points']),
                                        (t('itineraires_effectues'), ctx['total_itineraires']),
                                        (t('presence_terrain'), ctx['total_sessions']),
                                        (t('photos_label'), ctx['nb_photos']),
                                        (t('total_zones'), ctx['total_zones'])]):
            kpis.cell(0, i).text = lib
            kpis.cell(1, i).text = str(val)

    if 'beneficiaires' in sections:
        doc.add_heading(f"{t('section_beneficiaires')} ({bene_total})", level=2)
        if ctx['par_projet']:
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Light Shading Accent 1'
            hdr = table.rows[0].cells
            for i, k in enumerate(['informations_projet', 'beneficiaires', 'hommes', 'femmes', 'enfants', 'menages']):
                hdr[i].text = t(k)
            for x in ctx['par_projet']:
                sub = [a for a in activites if a.projet_id == x['projet'].pk]
                row = table.add_row().cells
                row[0].text = x['projet'].nom
                row[1].text = str(x['benef'])
                row[2].text = str(sum(a.hommes or 0 for a in sub))
                row[3].text = str(sum(a.femmes or 0 for a in sub))
                row[4].text = str(sum(a.enfants or 0 for a in sub))
                row[5].text = str(sum(a.menages or 0 for a in sub))
        else:
            doc.add_paragraph(t('aucun_activite'))

    doc.add_heading(t('zones_securite'), level=2)
    if zones:
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Shading Accent 1'
        hdr = table.rows[0].cells
        for i, k in enumerate(['zone', 'statut', 'motif', 'declaree_le', 'par']):
            hdr[i].text = t(k)
        for z in zones:
            row = table.add_row().cells
            row[0].text = z.nom
            row[1].text = {
                'dangereuse': t('zone_dangereuse'),
                'securisee': t('zone_securisee'),
                'indisponible': t('zone_indisponible'),
            }.get(z.statut, z.get_statut_display())
            row[2].text = (z.motif or '')[:80]
            row[3].text = z.date_declaration.strftime('%d/%m/%Y %H:%M')
            row[4].text = (z.auteur.get_full_name() or z.auteur.username) if z.auteur else '-'
    else:
        doc.add_paragraph(t('aucune_zone_periode'))

    if 'dangers' in sections:
        dangers = [z for z in zones if z.statut == 'dangereuse']
        doc.add_heading(f"{t('zones_danger')} ({len(dangers)})", level=2)
        if dangers:
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Light Shading Accent 1'
            hdr = table.rows[0].cells
            for i, k in enumerate(['zone', 'statut', 'motif', 'declaree_le', 'par']):
                hdr[i].text = t(k)
            for z in dangers:
                row = table.add_row().cells
                row[0].text = z.nom
                row[1].text = t('zone_dangereuse')
                row[2].text = (z.motif or '')[:80]
                row[3].text = z.date_declaration.strftime('%d/%m/%Y %H:%M')
                row[4].text = nom_agent(z.auteur)
        else:
            doc.add_paragraph(t('aucun_danger_periode'))

    if 'terrain' in sections:
        doc.add_heading(f"{t('presence_terrain')} ({len(ctx['sessions'])})", level=2)
        if ctx['sessions']:
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Light Shading Accent 1'
            hdr = table.rows[0].cells
            for i, k in enumerate(['agent', 'projet_sg', 'activite', 'debut', 'duree']):
                hdr[i].text = t(k)
            for s in ctx['sessions']:
                row = table.add_row().cells
                row[0].text = nom_agent(s.utilisateur)
                row[1].text = s.projet.nom if s.projet else '-'
                row[2].text = (s.activite_nom or '-')[:60]
                row[3].text = s.debut.strftime('%d/%m/%Y %H:%M')
                if s.debut and s.fin:
                    mins = int((s.fin - s.debut).total_seconds() // 60)
                    row[4].text = f"{mins // 60}h{mins % 60:02d}"
                else:
                    row[4].text = '-'
        else:
            doc.add_paragraph(t('aucune_session'))

    if 'itineraire' in sections:
        doc.add_heading(f"{t('itineraires_effectues')} ({len(ctx['itineraires'])})", level=2)
        if ctx['itineraires']:
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Light Shading Accent 1'
            hdr = table.rows[0].cells
            for i, k in enumerate(['nom', 'agent', 'projet_sg', 'date', 'alerte']):
                hdr[i].text = t(k)
            for it in ctx['itineraires']:
                row = table.add_row().cells
                row[0].text = it.nom
                row[1].text = nom_agent(it.utilisateur)
                row[2].text = it.projet.nom if it.projet else '-'
                row[3].text = it.date_creation.strftime('%d/%m/%Y %H:%M')
                row[4].text = (it.alerte or '')[:80]
        else:
            doc.add_paragraph(t('aucun_itineraire'))

    if 'points' in sections:
        doc.add_heading(f"{t('points_visites')} ({len(ctx['points'])})", level=2)
        if ctx['points']:
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Light Shading Accent 1'
            hdr = table.rows[0].cells
            for i, k in enumerate(['nom', 'categorie', 'commune', 'projet_sg', 'agent', 'date']):
                hdr[i].text = t(k)
            for p in ctx['points']:
                row = table.add_row().cells
                row[0].text = p.nom
                row[1].text = p.get_categorie_display()
                row[2].text = p.commune or '-'
                row[3].text = p.projet.nom if p.projet else '-'
                row[4].text = nom_agent(p.auteur)
                row[5].text = p.date_creation.strftime('%d/%m/%Y')
        else:
            doc.add_paragraph(t('aucun_point'))

    doc.add_heading(t('agents'), level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Shading Accent 1'
    hdr = table.rows[0].cells
    for i, k in enumerate(['nom_complet', 'telephone', 'email', 'fonction']):
        hdr[i].text = t(k)
    for p in profils:
        row = table.add_row().cells
        row[0].text = p.utilisateur.get_full_name() or p.utilisateur.username
        row[1].text = p.telephone or '-'
        row[2].text = p.utilisateur.email or '-'
        row[3].text = p.fonction or '-'

    doc.add_heading(t('activites_administrateurs'), level=2)
    if audits:
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Shading Accent 1'
        hdr = table.rows[0].cells
        for i, k in enumerate(['date_heure', 'administrateur', 'action', 'adresse_ip']):
            hdr[i].text = t(k)
        for l in audits[:100]:
            row = table.add_row().cells
            row[0].text = l.date.strftime('%d/%m/%Y %H:%M')
            row[1].text = l.utilisateur.get_full_name() or l.utilisateur.username
            row[2].text = (l.action or '')[:90]
            row[3].text = l.adresse_ip or '-'
    else:
        doc.add_paragraph(t('aucune_admin_periode'))

    if 'observations' in sections or 'recommandations' in sections:
        doc.add_heading(f"{t('observations_recommandations')}", level=2)
        nb = 0
        for a in activites:
            a_obs = a.observations and 'observations' in sections
            a_rec = a.recommandations and 'recommandations' in sections
            if not (a_obs or a_rec):
                continue
            nb += 1
            doc.add_heading(f"{nom_agent(a.agent)} — {a.projet.nom} ({a.date_creation.strftime('%d/%m/%Y')})", level=3)
            if a_obs:
                p = doc.add_paragraph()
                p.add_run(f"{t('observations_du')} : ").bold = True
                p.add_run((a.observations or '')[:300])
            if a_rec:
                p = doc.add_paragraph()
                p.add_run(f"{t('recommandations')} : ").bold = True
                p.add_run((a.recommandations or '')[:300])
        if nb == 0:
            doc.add_paragraph(t('aucun_obs_rec'))

    if 'conditions_meteo' in sections:
        doc.add_heading(t('conditions_meteo_activite'), level=2)
        nb_meteo = 0
        for a in activites:
            m = meteo_par_activite.get(a.pk)
            if not m:
                continue
            nb_meteo += 1
            doc.add_heading(f"{nom_agent(a.agent)} — {a.projet.nom} ({a.date_creation.strftime('%d/%m/%Y')})", level=3)
            if m.donnees_disponibles and m.temperature_c is not None:
                if m.localisation:
                    p = doc.add_paragraph()
                    p.add_run(f"{t('meteo_localisation')} : ").bold = True
                    p.add_run(m.localisation)
                p = doc.add_paragraph()
                p.add_run(f"{t('temperature')} : ").bold = True
                p.add_run(f"{m.temperature_c:.1f} °C")
                if m.conditions:
                    p = doc.add_paragraph()
                    p.add_run(f"{t('conditions_meteo')} : ").bold = True
                    p.add_run(m.conditions)
                if m.humidite is not None:
                    p = doc.add_paragraph()
                    p.add_run(f"{t('humidite')} : ").bold = True
                    p.add_run(f"{m.humidite} %")
                if m.vent_kmh is not None:
                    p = doc.add_paragraph()
                    p.add_run(f"{t('vent')} : ").bold = True
                    p.add_run(f"{m.vent_kmh:.0f} km/h {m.vent_direction or ''}".strip())
                if m.proba_pluie is not None:
                    p = doc.add_paragraph()
                    p.add_run(f"{t('proba_pluie')} : ").bold = True
                    p.add_run(f"{m.proba_pluie} %")
                if m.lever_soleil and m.coucher_soleil:
                    p = doc.add_paragraph()
                    p.add_run(f"{t('soleil')} : ").bold = True
                    p.add_run(f"{m.lever_soleil.strftime('%H:%M')} — {m.coucher_soleil.strftime('%H:%M')}")
            else:
                doc.add_paragraph(t('meteo_indisponible'))
            p = doc.add_paragraph()
            p.add_run(f"{t('meteo_source')} : ").bold = True
            p.add_run(t('source_' + (m.source if m.source in ('temps_reel', 'cache', 'synchronise') else 'temps_reel')))
            p.add_run(f" — {t('meteo_releve_le')} " + (m.horodatage_meteo.strftime('%d/%m/%Y %H:%M') if m.horodatage_meteo else t('meteo_inconnu')))
        if nb_meteo == 0:
            doc.add_paragraph(t('aucune_meteo_periode'))

    doc.add_heading(t('activites_agents'), level=2)
    for a in activites:
        doc.add_heading(f"{a.projet.nom} — {a.date_creation.strftime('%d/%m/%Y %H:%M')}", level=3)
        p = doc.add_paragraph()
        p.add_run(f"{t('agent')} : ").bold = True
        p.add_run(nom_agent(a.agent))
        p = doc.add_paragraph()
        p.add_run(f"{t('zone_visitee')} : ").bold = True
        p.add_run(a.zone_visitee or 'N/A')
        p = doc.add_paragraph()
        p.add_run(f"{t('securite')} : ").bold = True
        p.add_run(a.niveau_securite or 'N/A')
        p = doc.add_paragraph()
        p.add_run(f"{t('coordonnees_gps')} : ").bold = True
        p.add_run(f"{a.latitude}, {a.longitude}")
        p = doc.add_paragraph()
        p.add_run(f"{t('beneficiaires')} : ").bold = True
        p.add_run(str(a.nombre_beneficiaires or 0))
        if 'beneficiaires' in sections:
            p = doc.add_paragraph()
            p.add_run(f"{t('detail_beneficiaires')} : ").bold = True
            p.add_run(f"{t('hommes')} : {a.hommes or 0} | {t('femmes')} : {a.femmes or 0} | {t('enfants')} : {a.enfants or 0} | {t('menages')} : {a.menages or 0}")
        if a.objectif:
            p = doc.add_paragraph()
            p.add_run(f"{t('objectif_activite')} : ").bold = True
            p.add_run((a.objectif or '')[:300])
        p = doc.add_paragraph()
        p.add_run(f"{t('rapport_du')} : ").bold = True
        p.add_run((a.rapport or '')[:300])
        if a.resultats:
            p = doc.add_paragraph()
            p.add_run(f"{t('resultats_activite')} : ").bold = True
            p.add_run((a.resultats or '')[:300])
        if a.difficultes:
            p = doc.add_paragraph()
            p.add_run(f"{t('difficultes_activite')} : ").bold = True
            p.add_run((a.difficultes or '')[:300])
        if a.recommandations:
            p = doc.add_paragraph()
            p.add_run(f"{t('recommandations_activite')} : ").bold = True
            p.add_run((a.recommandations or '')[:300])
        if a.observations:
            p = doc.add_paragraph()
            p.add_run(f"{t('observations_du')} : ").bold = True
            p.add_run((a.observations or '')[:300])

        for photo in a.photos.all():
            if photo.image and os.path.exists(photo.image.path):
                try:
                    doc.add_picture(photo.image.path, width=Inches(3))
                except Exception:
                    pass

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{nom_fichier}.docx"'
    doc.save(response)
    _audit(None, "Téléchargement rapport DOCX", nom_fichier)
    return response


def _generer_pdf(ctx, nom_fichier, lang='fr'):
    from fpdf import FPDF
    from django.utils import timezone
    from .branding import chemin_logo, DEVELOPPEUR, NOM, VERSION
    from .i18n import traduire

    t = lambda cle: traduire(lang, cle)
    maintenant = timezone.now()
    debut, fin = ctx['debut'], ctx['fin']
    activites = ctx['activites']
    total = ctx['total']
    zd = ctx['zones_dangereuses']
    zs = ctx['zones_securisees']
    zi = ctx['zones_indisponibles']
    zones = ctx['zones']
    audits = ctx['audits']
    profils = ctx['profils']
    bene_total = ctx['bene_total']
    sections = set(ctx['sections'])

    def nom_agent(u):
        return (u.get_full_name() or u.username) if u else 'N/A'

    meteo_par_activite = ctx.get('meteo_par_activite', {})

    def net(s):
        if s is None:
            return ''
        if lang == 'zh':
            return str(s)
        return str(s).encode('cp1252', 'replace').decode('cp1252')

    def tronque(s, n):
        s = net(s)
        return s if len(s) <= n else s[:n - 1] + '…'

    POLICES = [
        ('Arial', '', 'C:/Windows/Fonts/arial.ttf'),
        ('Arial', 'B', 'C:/Windows/Fonts/arialbd.ttf'),
        ('Arial', 'I', 'C:/Windows/Fonts/ariali.ttf'),
    ]
    fam = 'Arial'
    if lang == 'zh' and os.path.exists('C:/Windows/Fonts/msyh.ttc'):
        POLICES = [
            ('YaHei', '', 'C:/Windows/Fonts/msyh.ttc'),
            ('YaHei', 'B', 'C:/Windows/Fonts/msyh.ttc'),
            ('YaHei', 'I', 'C:/Windows/Fonts/msyh.ttc'),
        ]
        fam = 'YaHei'
    elif not all(os.path.exists(c) for _, _, c in POLICES):
        DEJAVU = [
            ('DejaVu', '', '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf'),
            ('DejaVu', 'B', '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf'),
            ('DejaVu', 'I', '/usr/share/fonts/truetype/dejavu/DejaVuSans-Oblique.ttf'),
        ]
        if all(os.path.exists(c) for _, _, c in DEJAVU):
            POLICES = DEJAVU
            fam = 'DejaVu'
        else:
            fam = 'Helvetica'

    ACCENT = (77, 67, 246)
    VERT = (34, 197, 94)
    ROUGE = (239, 68, 68)
    JAUNE = (234, 179, 8)
    TEXTE = (31, 36, 60)
    GRIS = (110, 118, 150)
    LIGNE = (225, 227, 245)

    def statut_zone(s):
        return {
            'dangereuse': (t('zone_dangereuse'), ROUGE),
            'securisee': (t('zone_securisee'), VERT),
            'indisponible': (t('zone_indisponible'), JAUNE),
        }.get(s, (s or '-', TEXTE))

    logo_chemin = chemin_logo()
    logo_dispo = os.path.exists(logo_chemin)

    def pdf_propre(s):
        if s is None:
            return ''
        return str(s).translate(str.maketrans({
            '—': '-', '–': '-', '…': '...',
            '‘': "'", '’': "'", '“': '"', '”': '"',
            '€': 'EUR', '•': '*', '·': '*',
        }))

    class RapportPDF(FPDF):
        def _propre(self, txt):
            return pdf_propre(txt) if isinstance(txt, str) else txt

        def cell(self, *args, **kwargs):
            args = list(args)
            if len(args) >= 3:
                args[2] = self._propre(args[2])
            elif 'txt' in kwargs:
                kwargs['txt'] = self._propre(kwargs['txt'])
            return super().cell(*args, **kwargs)

        def multi_cell(self, *args, **kwargs):
            args = list(args)
            if len(args) >= 3:
                args[2] = self._propre(args[2])
            elif 'txt' in kwargs:
                kwargs['txt'] = self._propre(kwargs['txt'])
            return super().multi_cell(*args, **kwargs)

        def header(self):
            if self.page_no() == 1:
                return  # page de couverture, dessinée manuellement
            if logo_dispo:
                try:
                    self.image(logo_chemin, x=8, y=4.5, w=15)
                except Exception:
                    pass
            self.set_fill_color(*ACCENT)
            self.rect(0, 0, 210, 26, 'F')
            self.set_text_color(255, 255, 255)
            self.set_font(fam, 'B', 15)
            self.set_y(4)
            self.set_x(26)
            self.cell(174, 9, f'{NOM} v{VERSION} — {t("plateforme_sig")}', 0, 1, 'C')
            self.set_font(fam, '', 9)
            self.set_x(26)
            self.cell(174, 7, f"{t('rapport_activites_periode')} — {debut.strftime('%d/%m/%Y')} {t('au')} {fin.strftime('%d/%m/%Y')}", 0, 1, 'C')
            self.ln(4)
            self.set_draw_color(255, 255, 255)
            self.line(10, 25.5, 200, 25.5)
            self.set_text_color(*TEXTE)

        def footer(self):
            self.set_y(-14)
            self.set_font(fam, '', 8)
            self.set_text_color(*GRIS)
            self.cell(0, 6, f"{NOM} v{VERSION} — {DEVELOPPEUR}  |  {t('genere_le')} {maintenant.strftime('%d/%m/%Y')} {t('heure')} {maintenant.strftime('%H:%M')}  |  {t('page')} {self.page_no()}/{self.alias_nb_pages()}", 0, 0, 'C')

    pdf = RapportPDF()
    try:
        for nom, st, chemin in POLICES:
            if os.path.exists(chemin):
                pdf.add_font(nom, st, chemin)
    except Exception:
        fam = 'Helvetica'
    pdf.alias_nb_pages()
    pdf.set_auto_page_break(True, margin=18)

    # ── 0. Page de couverture ───────────────────────────────────
    pdf.add_page()
    if logo_dispo:
        try:
            pdf.image(logo_chemin, x=(210 - 55) / 2, y=30, w=55)
        except Exception:
            pass
    pdf.set_y(100)
    pdf.set_font(fam, 'B', 30)
    pdf.set_text_color(*ACCENT)
    pdf.cell(0, 14, f'{NOM} v{VERSION}', 0, 1, 'C')
    pdf.set_font(fam, '', 13)
    pdf.set_text_color(*TEXTE)
    pdf.cell(0, 8, net(t('plateforme_sig')), 0, 1, 'C')
    pdf.ln(12)
    pdf.set_fill_color(*ACCENT)
    pdf.rect(45, pdf.get_y(), 120, 0.8, 'F')
    pdf.ln(12)
    pdf.set_font(fam, 'B', 20)
    pdf.cell(0, 10, net(t('couverture')), 0, 1, 'C')
    pdf.ln(14)
    pdf.set_font(fam, '', 11)
    intitule = ctx['intitule_projet'] or t('tous_projets')
    for lib, val in [(t('periode_rapport'), f"{debut.strftime('%d/%m/%Y')} — {fin.strftime('%d/%m/%Y')}"),
                     (t('type_rapport'), t('rapport_' + ctx['type'])),
                     (t('informations_projet'), intitule),
                     (t('reference'), nom_fichier),
                     (t('genere_le'), f"{maintenant.strftime('%d/%m/%Y')} {t('heure')} {maintenant.strftime('%H:%M')}"),
                     (t('developpe_par'), DEVELOPPEUR)]:
        pdf.set_font(fam, 'B', 11)
        pdf.cell(90, 8, net(lib), 0, 0, 'R')
        pdf.set_font(fam, '', 11)
        pdf.cell(0, 8, '  ' + net(val), 0, 1, 'L')
    pdf.add_page()

    def titre_section(txt):
        if pdf.get_y() > 230:
            pdf.add_page()
        y = pdf.get_y()
        pdf.set_fill_color(*ACCENT)
        pdf.rect(10, y + 1, 4, 8, 'F')
        pdf.set_font(fam, 'B', 12)
        pdf.set_text_color(*TEXTE)
        pdf.cell(0, 9, '  ' + net(txt), 0, 1)
        pdf.set_draw_color(*LIGNE)
        pdf.line(10, pdf.get_y() + 1, 200, pdf.get_y() + 1)
        pdf.ln(6)

    def table_entete(largeurs, cols):
        pdf.set_fill_color(*ACCENT)
        pdf.set_text_color(255, 255, 255)
        pdf.set_font(fam, 'B', 8.5)
        for lib, w in zip(cols, largeurs):
            pdf.cell(w, 7.5, net(lib), 1, 0, 'C', True)
        pdf.ln()
        pdf.set_text_color(*TEXTE)

    def ligne_table(largeurs, vals, hauteur=7, alterner=False):
        pdf.set_font(fam, '', 8.5)
        if alterner:
            pdf.set_fill_color(243, 244, 252)
        for v, w in zip(vals, largeurs):
            pdf.cell(w, hauteur, tronque(v, int(w / 1.7)), 1, 0, 'L', alterner)
        pdf.ln()
        pdf.set_text_color(*TEXTE)

    # ── 1. Résumé (blocs de statistiques)
    titre_section(f"{t('resume_periode')} — {total} {t('activites').lower()} / {bene_total} {t('beneficiaires').lower()}")
    x0 = 10
    bloc = 37
    lignes_stats = [
        [(t('activites'), total, ACCENT), (t('beneficiaires'), bene_total, ACCENT),
         (t('kpi_hommes'), ctx['tot_h'], ACCENT), (t('kpi_femmes'), ctx['tot_f'], ACCENT),
         (t('kpi_enfants'), ctx['tot_e'], ACCENT)],
        [(t('kpi_menages'), ctx['tot_m'], ACCENT), (t('agents'), ctx['total_agents'], ACCENT),
         (t('kpi_zones'), ctx['total_zones'], JAUNE), (t('points_visites'), ctx['total_points'], ACCENT),
         (t('itineraires_effectues'), ctx['total_itineraires'], ACCENT)],
        [(t('kpi_sessions'), ctx['total_sessions'], VERT), (t('photos_label'), ctx['nb_photos'], ACCENT),
         (t('zones_dangereuses'), zd, ROUGE), (t('zones_securisees'), zs, VERT),
         (t('zones_sans_info'), zi, JAUNE)],
    ]
    for ligne in lignes_stats:
        y = pdf.get_y()
        x0 = 10
        for lib, val, col in ligne:
            pdf.set_fill_color(*col)
            pdf.rect(x0, y, bloc, 18, 'F')
            pdf.set_xy(x0, y + 2)
            pdf.set_text_color(255, 255, 255)
            pdf.set_font(fam, 'B', 14)
            pdf.cell(bloc, 7, net(str(val)), 0, 1, 'C')
            pdf.set_font(fam, '', 7)
            pdf.cell(bloc, 7, net(lib), 0, 1, 'C')
            x0 += bloc + 1
        pdf.set_y(y + 22)

    # ── 1bis. Bénéficiaires par projet
    if 'beneficiaires' in sections:
        titre_section(f"{t('section_beneficiaires')} ({bene_total})")
        if ctx['par_projet']:
            largeurs = [70, 26, 26, 26, 26, 26]
            table_entete(largeurs, [t('informations_projet'), t('beneficiaires'), t('hommes'), t('femmes'), t('enfants'), t('menages')])
            for i, x in enumerate(ctx['par_projet']):
                sub = [a for a in activites if a.projet_id == x['projet'].pk]
                ligne_table(largeurs, [tronque(x['projet'].nom, 40), str(x['benef']),
                                       str(sum(a.hommes or 0 for a in sub)),
                                       str(sum(a.femmes or 0 for a in sub)),
                                       str(sum(a.enfants or 0 for a in sub)),
                                       str(sum(a.menages or 0 for a in sub))],
                            alterner=i % 2 == 1)
        else:
            pdf.set_font(fam, 'I', 9)
            pdf.set_text_color(*GRIS)
            pdf.cell(0, 7, net(t('aucun_activite')), 0, 1)

    # ── 2. Zones de sécurité
    titre_section(f"{t('zones_securite')} ({len(zones)})")
    if zones:
        largeurs = [44, 34, 57, 24, 28]
        table_entete(largeurs, [t('zone'), t('statut'), t('motif'), t('declaree_le'), t('par')])
        for i, z in enumerate(zones):
            lib_statut, couleur = statut_zone(z.statut)
            pdf.set_font(fam, '', 8.5)
            if i % 2 == 1:
                pdf.set_fill_color(243, 244, 252)
            pdf.cell(largeurs[0], 7, tronque(z.nom, 24), 1, 0, 'L', i % 2 == 1)
            pdf.set_text_color(*couleur)
            pdf.set_font(fam, 'B', 8.5)
            pdf.cell(largeurs[1], 7, tronque(lib_statut, 18), 1, 0, 'C', i % 2 == 1)
            pdf.set_text_color(*TEXTE)
            pdf.set_font(fam, '', 8.5)
            pdf.cell(largeurs[2], 7, tronque(z.motif, 32), 1, 0, 'L', i % 2 == 1)
            pdf.cell(largeurs[3], 7, tronque(z.date_declaration.strftime('%d/%m/%Y'), 13), 1, 0, 'C', i % 2 == 1)
            pdf.cell(largeurs[4], 7, tronque((z.auteur.get_full_name() or z.auteur.username) if z.auteur else '-', 15), 1, 0, 'L', i % 2 == 1)
            pdf.ln()
    else:
        pdf.set_font(fam, 'I', 9)
        pdf.set_text_color(*GRIS)
        pdf.cell(0, 7, net(t('aucune_zone_periode')), 0, 1)

    # ── 2bis. Zones de danger
    if 'dangers' in sections:
        dangers = [z for z in zones if z.statut == 'dangereuse']
        titre_section(f"{t('zones_danger')} ({len(dangers)})")
        if dangers:
            largeurs = [44, 34, 57, 24, 28]
            table_entete(largeurs, [t('zone'), t('statut'), t('motif'), t('declaree_le'), t('par')])
            for i, z in enumerate(dangers):
                pdf.set_font(fam, '', 8.5)
                if i % 2 == 1:
                    pdf.set_fill_color(243, 244, 252)
                pdf.cell(largeurs[0], 7, tronque(z.nom, 24), 1, 0, 'L', i % 2 == 1)
                pdf.set_text_color(*ROUGE)
                pdf.set_font(fam, 'B', 8.5)
                pdf.cell(largeurs[1], 7, tronque(t('zone_dangereuse'), 18), 1, 0, 'C', i % 2 == 1)
                pdf.set_text_color(*TEXTE)
                pdf.set_font(fam, '', 8.5)
                pdf.cell(largeurs[2], 7, tronque(z.motif, 32), 1, 0, 'L', i % 2 == 1)
                pdf.cell(largeurs[3], 7, tronque(z.date_declaration.strftime('%d/%m/%Y'), 13), 1, 0, 'C', i % 2 == 1)
                pdf.cell(largeurs[4], 7, tronque(nom_agent(z.auteur), 15), 1, 0, 'L', i % 2 == 1)
                pdf.ln()
        else:
            pdf.set_font(fam, 'I', 9)
            pdf.set_text_color(*GRIS)
            pdf.cell(0, 7, net(t('aucun_danger_periode')), 0, 1)

    # ── 2ter. Présence terrain (sessions)
    if 'terrain' in sections:
        titre_section(f"{t('presence_terrain')} ({len(ctx['sessions'])})")
        if ctx['sessions']:
            largeurs = [40, 52, 42, 30, 23]
            table_entete(largeurs, [t('agent'), t('projet_sg'), t('activite'), t('debut'), t('duree')])
            for i, s in enumerate(ctx['sessions']):
                duree = ''
                if s.debut and s.fin:
                    mins = int((s.fin - s.debut).total_seconds() // 60)
                    duree = f"{mins // 60}h{mins % 60:02d}"
                ligne_table(largeurs, [nom_agent(s.utilisateur),
                                       tronque(s.projet.nom if s.projet else '-', 28),
                                       tronque(s.activite_nom or '-', 24),
                                       s.debut.strftime('%d/%m %H:%M'), duree],
                            alterner=i % 2 == 1)
        else:
            pdf.set_font(fam, 'I', 9)
            pdf.set_text_color(*GRIS)
            pdf.cell(0, 7, net(t('aucune_session')), 0, 1)

    # ── 2quater. Itinéraires
    if 'itineraire' in sections:
        titre_section(f"{t('itineraires_effectues')} ({len(ctx['itineraires'])})")
        if ctx['itineraires']:
            largeurs = [50, 42, 30, 32, 33]
            table_entete(largeurs, [t('nom'), t('agent'), t('projet_sg'), t('date'), t('alerte')])
            for i, it in enumerate(ctx['itineraires']):
                ligne_table(largeurs, [tronque(it.nom, 28), nom_agent(it.utilisateur),
                                       tronque(it.projet.nom if it.projet else '-', 16),
                                       it.date_creation.strftime('%d/%m/%Y'),
                                       tronque(it.alerte or '-', 18)],
                            alterner=i % 2 == 1)
        else:
            pdf.set_font(fam, 'I', 9)
            pdf.set_text_color(*GRIS)
            pdf.cell(0, 7, net(t('aucun_itineraire')), 0, 1)

    # ── 2quinquies. Points visités
    if 'points' in sections:
        titre_section(f"{t('points_visites')} ({len(ctx['points'])})")
        if ctx['points']:
            largeurs = [44, 28, 30, 34, 30, 21]
            table_entete(largeurs, [t('nom'), t('categorie'), t('commune'), t('projet_sg'), t('agent'), t('date')])
            for i, p in enumerate(ctx['points']):
                ligne_table(largeurs, [tronque(p.nom, 25), tronque(p.get_categorie_display(), 15),
                                       tronque(p.commune or '-', 16),
                                       tronque(p.projet.nom if p.projet else '-', 18),
                                       nom_agent(p.auteur), p.date_creation.strftime('%d/%m/%Y')],
                            alterner=i % 2 == 1)
        else:
            pdf.set_font(fam, 'I', 9)
            pdf.set_text_color(*GRIS)
            pdf.cell(0, 7, net(t('aucun_point')), 0, 1)

    # ── 3. Agents
    titre_section(f"{t('agents')} ({len(profils)})")
    if profils:
        largeurs = [52, 34, 42, 59]
        table_entete(largeurs, [t('nom_complet'), t('telephone'), t('fonction'), t('email')])
        for i, p in enumerate(profils):
            ligne_table(largeurs, [p.utilisateur.get_full_name() or p.utilisateur.username,
                                   p.telephone or '-', p.fonction or '-', p.utilisateur.email or '-'],
                        alterner=i % 2 == 1)
    else:
        pdf.set_font(fam, 'I', 9)
        pdf.set_text_color(*GRIS)
        pdf.cell(0, 7, net(t('aucun_agent')), 0, 1)

    # ── 4. Activités des administrateurs
    titre_section(f"{t('activites_administrateurs')} ({len(audits)})")
    if audits:
        largeurs = [27, 40, 86, 34]
        table_entete(largeurs, [t('date_heure'), t('administrateur'), t('action'), t('adresse_ip')])
        for i, l in enumerate(audits[:200]):
            ligne_table(largeurs, [l.date.strftime('%d/%m/%Y %H:%M'),
                                   l.utilisateur.get_full_name() or l.utilisateur.username,
                                   (l.action or '')[:90], l.adresse_ip or '-'],
                        alterner=i % 2 == 1)
    else:
        pdf.set_font(fam, 'I', 9)
        pdf.set_text_color(*GRIS)
        pdf.cell(0, 7, net(t('aucune_admin_periode')), 0, 1)

    # ── 4bis. Observations & recommandations
    if 'observations' in sections or 'recommandations' in sections:
        titre_section(f"{t('observations_recommandations')}")
        nb_or = 0
        for a in activites:
            a_obs = a.observations and 'observations' in sections
            a_rec = a.recommandations and 'recommandations' in sections
            if not (a_obs or a_rec):
                continue
            nb_or += 1
            if pdf.get_y() > 235:
                pdf.add_page()
            pdf.set_font(fam, 'B', 9)
            pdf.cell(0, 5.5, f"{nom_agent(a.agent)} — {a.projet.nom} ({a.date_creation.strftime('%d/%m/%Y')})", 0, 1)
            if a_obs:
                pdf.set_x(10)
                pdf.multi_cell(0, 5.5, f"  {t('observations_du')} : " + tronque(a.observations, 200))
            if a_rec:
                pdf.set_x(10)
                pdf.multi_cell(0, 5.5, f"  {t('recommandations')} : " + tronque(a.recommandations, 200))
            pdf.ln(1)
        if nb_or == 0:
            pdf.set_font(fam, 'I', 9)
            pdf.set_text_color(*GRIS)
            pdf.cell(0, 7, net(t('aucun_obs_rec')), 0, 1)

    # ── 4ter. Conditions météorologiques des activités
    if 'conditions_meteo' in sections:
        titre_section(f"{t('conditions_meteo_activite')}")
        nb_meteo = 0
        for a in activites:
            m = meteo_par_activite.get(a.pk)
            if not m:
                continue
            nb_meteo += 1
            if pdf.get_y() > 235:
                pdf.add_page()
            pdf.set_font(fam, 'B', 9)
            pdf.cell(0, 5.5, f"{nom_agent(a.agent)} — {a.projet.nom} ({a.date_creation.strftime('%d/%m/%Y')})", 0, 1)
            pdf.set_font(fam, '', 9)
            pdf.set_text_color(*TEXTE)
            if m.donnees_disponibles and m.temperature_c is not None:
                ligne = f"  {t('temperature')} : {m.temperature_c:.1f} °C"
                if m.conditions:
                    ligne += f"   |   {t('conditions_meteo')} : {tronque(m.conditions, 40)}"
                pdf.set_x(10)
                pdf.multi_cell(0, 5.5, ligne)
                if m.localisation:
                    pdf.set_x(10)
                    pdf.multi_cell(0, 5.5, f"  {t('meteo_localisation')} : " + tronque(m.localisation, 60))
                ligne2 = ''
                if m.humidite is not None:
                    ligne2 += f"{t('humidite')} : {m.humidite} %"
                if m.vent_kmh is not None:
                    ligne2 += f"   |   {t('vent')} : {m.vent_kmh:.0f} km/h {m.vent_direction or ''}".strip()
                if m.proba_pluie is not None:
                    ligne2 += f"   |   {t('proba_pluie')} : {m.proba_pluie} %"
                if ligne2:
                    pdf.set_x(10)
                    pdf.multi_cell(0, 5.5, '  ' + ligne2)
                if m.lever_soleil and m.coucher_soleil:
                    pdf.set_x(10)
                    pdf.multi_cell(0, 5.5, f"  {t('soleil')} : {m.lever_soleil.strftime('%H:%M')} — {m.coucher_soleil.strftime('%H:%M')}")
            else:
                pdf.set_x(10)
                pdf.multi_cell(0, 5.5, '  ' + net(t('meteo_indisponible')))
            src = m.source if m.source in ('temps_reel', 'cache', 'synchronise') else 'temps_reel'
            horo = m.horodatage_meteo.strftime('%d/%m/%Y %H:%M') if m.horodatage_meteo else t('meteo_inconnu')
            pdf.set_font(fam, 'I', 8)
            pdf.set_text_color(*GRIS)
            pdf.set_x(10)
            pdf.cell(0, 5, f"{t('meteo_source')} : {t('source_' + src)}   |   {t('meteo_releve_le')} {horo}", 0, 1)
            pdf.set_text_color(*TEXTE)
            pdf.set_font(fam, '', 9)
            pdf.ln(1)
        if nb_meteo == 0:
            pdf.set_font(fam, 'I', 9)
            pdf.set_text_color(*GRIS)
            pdf.cell(0, 7, net(t('aucune_meteo_periode')), 0, 1)

    # ── 5. Activités des agents
    titre_section(f"{t('activites_agents')} ({total})")
    for i, a in enumerate(activites, 1):
        if pdf.get_y() > 235:
            pdf.add_page()
        y = pdf.get_y()
        pdf.set_fill_color(*ACCENT)
        pdf.rect(10, y, 4, 8, 'F')
        pdf.set_font(fam, 'B', 10.5)
        pdf.set_text_color(*TEXTE)
        pdf.cell(0, 8, f'  {i}. {tronque(a.projet.nom, 60)}', 0, 1)
        pdf.set_draw_color(*LIGNE)
        pdf.line(10, pdf.get_y(), 200, pdf.get_y())
        pdf.ln(2)

        pdf.set_font(fam, '', 9)
        pdf.cell(0, 5.5, f"{t('agent')} : {nom_agent(a.agent)}   |   Date : {a.date_creation.strftime('%d/%m/%Y')} {t('heure')} {a.date_creation.strftime('%H:%M')}", 0, 1)
        pdf.cell(0, 5.5, f"{t('zone_visitee')} : {tronque(a.zone_visitee or 'N/A', 60)}   |   {t('securite')} : {tronque(a.niveau_securite or 'N/A', 18)}", 0, 1)
        pdf.cell(0, 5.5, f"{t('coordonnees_gps')} : {a.latitude}, {a.longitude}   |   {t('beneficiaires')} : {a.nombre_beneficiaires or 0}", 0, 1)
        if 'beneficiaires' in sections:
            pdf.set_x(10)
            pdf.cell(0, 5.5, f"  {t('detail_beneficiaires')} : {t('hommes')} : {a.hommes or 0} | {t('femmes')} : {a.femmes or 0} | {t('enfants')} : {a.enfants or 0} | {t('menages')} : {a.menages or 0}", 0, 1)
        if a.objectif:
            pdf.set_x(10)
            pdf.multi_cell(0, 5.5, f"{t('objectif_activite')} : " + tronque(a.objectif, 200))
        if a.rapport:
            pdf.set_x(10)
            pdf.multi_cell(0, 5.5, f"{t('rapport_du')} : " + tronque(a.rapport, 280))
        if a.resultats:
            pdf.set_x(10)
            pdf.multi_cell(0, 5.5, f"{t('resultats_activite')} : " + tronque(a.resultats, 200))
        if a.difficultes:
            pdf.set_x(10)
            pdf.multi_cell(0, 5.5, f"{t('difficultes_activite')} : " + tronque(a.difficultes, 200))
        if a.recommandations:
            pdf.set_x(10)
            pdf.multi_cell(0, 5.5, f"{t('recommandations_activite')} : " + tronque(a.recommandations, 200))
        if a.observations:
            pdf.set_x(10)
            pdf.multi_cell(0, 5.5, f"{t('observations_du')} : " + tronque(a.observations, 180))

        photos = [p for p in a.photos.all() if p.image and os.path.exists(p.image.path)]
        if photos:
            pdf.ln(2)
            pdf.set_font(fam, '', 8)
            pdf.set_text_color(*GRIS)
            pdf.cell(0, 5, f"{t('photos_label')} ({len(photos)}) :", 0, 1)
            x = 10
            for ph in photos[:4]:
                try:
                    if pdf.get_y() + 48 > 278:
                        pdf.add_page()
                        x = 10
                    pdf.image(ph.image.path, x=x, w=48)
                    x += 48 + 2
                except Exception:
                    continue
            if x > 10:
                pdf.set_y(pdf.get_y() + 2)
        pdf.ln(4)

    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{nom_fichier}.pdf"'
    response.write(bytes(pdf.output()))
    _audit(None, "Téléchargement rapport PDF", nom_fichier)
    return response


def _generer_excel(ctx, nom_fichier, lang='fr'):
    from openpyxl import Workbook
    from openpyxl.utils import get_column_letter
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from django.utils import timezone
    from .branding import NOM, VERSION
    from .i18n import traduire

    t = lambda cle: traduire(lang, cle)
    maintenant = timezone.now()
    debut, fin = ctx['debut'], ctx['fin']
    activites = ctx['activites']

    def nom_agent(u):
        return (u.get_full_name() or u.username) if u else 'N/A'

    meteo_par_activite = ctx.get('meteo_par_activite', {})

    wb = Workbook()
    ENTETE_FILL = PatternFill('solid', fgColor='4F46E5')
    ENTETE_FONT = Font(color='FFFFFF', bold=True)
    BORDURE = Border(*[Side(style='thin', color='D9D9D9')] * 4)

    def preparer(ws, largeurs, entetes):
        for i, w in enumerate(largeurs, 1):
            ws.column_dimensions[get_column_letter(i)].width = w
        for j, lib in enumerate(entetes, 1):
            c = ws.cell(row=1, column=j, value=lib)
            c.fill = ENTETE_FILL
            c.font = ENTETE_FONT
            c.alignment = Alignment(horizontal='center', vertical='center')
        ws.freeze_panes = 'A2'

    def remplir(ws, lignes):
        for i, row in enumerate(lignes, 2):
            for j, val in enumerate(row, 1):
                c = ws.cell(row=i, column=j, value=val)
                c.border = BORDURE
        if lignes:
            ws.auto_filter.ref = ws.dimensions

    # ── Synthèse ─────────────────────────────────────────────────
    ws = wb.active
    ws.title = 'Synthese'
    ws.column_dimensions['A'].width = 34
    ws.column_dimensions['B'].width = 22
    ws['A1'] = NOM + ' v' + VERSION
    ws['A2'] = t('rapport_periode') + f" : {debut.strftime('%d/%m/%Y')} — {fin.strftime('%d/%m/%Y')}"
    ws['A3'] = t('informations_projet') + ' : ' + (ctx['intitule_projet'] or t('tous_projets'))
    ws['A4'] = t('reference') + ' : ' + nom_fichier
    ws['A5'] = t('genere_le') + f" : {maintenant.strftime('%d/%m/%Y %H:%M')}"
    indicateurs = [
        (t('activites'), ctx['total']),
        (t('beneficiaires'), ctx['bene_total']),
        (t('kpi_hommes'), ctx['tot_h']),
        (t('kpi_femmes'), ctx['tot_f']),
        (t('kpi_enfants'), ctx['tot_e']),
        (t('kpi_menages'), ctx['tot_m']),
        (t('agents'), ctx['total_agents']),
        (t('total_zones'), ctx['total_zones']),
        (t('zones_dangereuses'), ctx['zones_dangereuses']),
        (t('zones_securisees'), ctx['zones_securisees']),
        (t('zones_indisponibles'), ctx['zones_indisponibles']),
        (t('points_visites'), ctx['total_points']),
        (t('itineraires_effectues'), ctx['total_itineraires']),
        (t('presence_terrain'), ctx['total_sessions']),
        (t('photos_label'), ctx['nb_photos']),
    ]
    row = 7
    for lib, val in indicateurs:
        ws.cell(row=row, column=1, value=lib).font = Font(bold=True)
        ws.cell(row=row, column=2, value=val)
        row += 1

    # ── Activités ────────────────────────────────────────────────
    ws = wb.create_sheet('Activites')
    cols = [t('date'), t('projet_sg'), t('code_projet'), t('agent'), t('nom_activite'),
            t('zone_visitee'), t('securite'), t('latitude'), t('longitude'),
            t('beneficiaires'), t('hommes'), t('femmes'), t('enfants'), t('menages'),
            t('objectif_activite'), t('rapport_activite'), t('resultats_activite'),
            t('difficultes_activite'), t('recommandations_activite'), t('observations')]
    preparer(ws, [12, 22, 10, 20, 24, 18, 12, 11, 11, 12, 9, 9, 9, 9, 30, 40, 30, 30, 30, 25], cols)
    lignes = []
    for a in activites:
        lignes.append([a.date_creation.strftime('%d/%m/%Y %H:%M'), a.projet.nom, a.projet.code or '',
                       nom_agent(a.agent), a.nom_activite or '', a.zone_visitee or '',
                       a.niveau_securite or '', a.latitude, a.longitude,
                       a.nombre_beneficiaires or 0, a.hommes or 0, a.femmes or 0,
                       a.enfants or 0, a.menages or 0, a.objectif or '', a.rapport or '',
                       a.resultats or '', a.difficultes or '', a.recommandations or '',
                       a.observations or ''])
    remplir(ws, lignes)

    # ── Conditions météo ────────────────────────────────────────
    ws = wb.create_sheet('Meteo')
    cols = [t('date'), t('projet_sg'), t('agent'), t('meteo_localisation'),
            t('temperature'), t('conditions_meteo'), t('humidite'), t('vent'),
            t('proba_pluie'), t('soleil'), t('meteo_source'), t('meteo_releve_le')]
    preparer(ws, [16, 20, 20, 26, 12, 22, 10, 14, 12, 18, 14, 16], cols)
    lignes = []
    for a in activites:
        m = meteo_par_activite.get(a.pk)
        if not m:
            continue
        src = m.source if m.source in ('temps_reel', 'cache', 'synchronise') else 'temps_reel'
        soleil = ''
        if m.lever_soleil and m.coucher_soleil:
            soleil = f"{m.lever_soleil.strftime('%H:%M')} — {m.coucher_soleil.strftime('%H:%M')}"
        vent = ''
        if m.vent_kmh is not None:
            vent = f"{m.vent_kmh:.0f} km/h {m.vent_direction or ''}".strip()
        lignes.append([
            a.date_creation.strftime('%d/%m/%Y %H:%M'), a.projet.nom, nom_agent(a.agent),
            m.localisation or '',
            f"{m.temperature_c:.1f} °C" if m.donnees_disponibles and m.temperature_c is not None else t('meteo_indisponible'),
            m.conditions or '',
            m.humidite if m.humidite is not None else '',
            vent,
            m.proba_pluie if m.proba_pluie is not None else '',
            soleil,
            t('source_' + src),
            m.horodatage_meteo.strftime('%d/%m/%Y %H:%M') if m.horodatage_meteo else t('meteo_inconnu'),
        ])
    remplir(ws, lignes)

    # ── Projets ──────────────────────────────────────────────────
    ws = wb.create_sheet('Projets')
    preparer(ws, [24, 12, 14, 30, 12], [t('informations_projet'), t('code_projet'), t('activites'), t('beneficiaires'), t('statut')])
    lignes = []
    for x in ctx['par_projet']:
        lignes.append([x['projet'].nom, x['projet'].code or '', x['count'], x['benef'],
                       x['projet'].get_statut_display()])
    remplir(ws, lignes)

    # ── Points ───────────────────────────────────────────────────
    ws = wb.create_sheet('Points')
    preparer(ws, [26, 16, 16, 16, 22, 12, 12, 12],
             [t('nom'), t('categorie'), t('province'), t('commune'), t('projet_sg'), t('latitude'), t('longitude'), t('date')])
    lignes = []
    for p in ctx['points']:
        lignes.append([p.nom, p.get_categorie_display(), p.province or '', p.commune or '',
                       p.projet.nom if p.projet else '-', p.latitude, p.longitude,
                       p.date_creation.strftime('%d/%m/%Y')])
    remplir(ws, lignes)

    # ── Itinéraires ──────────────────────────────────────────────
    ws = wb.create_sheet('Itineraires')
    preparer(ws, [26, 22, 22, 16, 30], [t('nom'), t('agent'), t('projet_sg'), t('date'), t('alerte')])
    lignes = []
    for it in ctx['itineraires']:
        lignes.append([it.nom, nom_agent(it.utilisateur), it.projet.nom if it.projet else '-',
                       it.date_creation.strftime('%d/%m/%Y'), it.alerte or ''])
    remplir(ws, lignes)

    # ── Zones ────────────────────────────────────────────────────
    ws = wb.create_sheet('Zones')
    preparer(ws, [26, 16, 34, 22, 16, 22], [t('zone'), t('statut'), t('motif'), t('projet_sg'), t('declaree_le'), t('par')])
    lignes = []
    for z in ctx['zones']:
        lignes.append([z.nom, z.get_statut_display(), z.motif or '',
                       z.projet.nom if z.projet else '-',
                       z.date_declaration.strftime('%d/%m/%Y'), nom_agent(z.auteur)])
    remplir(ws, lignes)

    # ── Agents ───────────────────────────────────────────────────
    ws = wb.create_sheet('Agents')
    preparer(ws, [26, 18, 26, 30], [t('nom_complet'), t('telephone'), t('fonction'), t('email')])
    lignes = []
    for p in ctx['profils']:
        lignes.append([nom_agent(p.utilisateur), p.telephone or '-', p.fonction or '-',
                       p.utilisateur.email or '-'])
    remplir(ws, lignes)

    # ── Présence terrain ─────────────────────────────────────────
    ws = wb.create_sheet('Presence')
    preparer(ws, [26, 24, 24, 18, 12, 30],
             [t('agent'), t('projet_sg'), t('activite'), t('debut'), t('duree'), t('observations')])
    lignes = []
    for s in ctx['sessions']:
        duree = ''
        if s.debut and s.fin:
            mins = int((s.fin - s.debut).total_seconds() // 60)
            duree = f"{mins // 60}h{mins % 60:02d}"
        lignes.append([nom_agent(s.utilisateur), s.projet.nom if s.projet else '-',
                       s.activite_nom or '-', s.debut.strftime('%d/%m/%Y %H:%M'), duree,
                       s.observations or ''])
    remplir(ws, lignes)

    # ── Audits (mode global) ─────────────────────────────────────
    if ctx['audits']:
        ws = wb.create_sheet('Audits')
        preparer(ws, [18, 26, 40, 18], [t('date_heure'), t('administrateur'), t('action'), t('adresse_ip')])
        lignes = []
        for l in ctx['audits']:
            lignes.append([l.date.strftime('%d/%m/%Y %H:%M'), nom_agent(l.utilisateur),
                           l.action or '', l.adresse_ip or '-'])
        remplir(ws, lignes)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="{nom_fichier}.xlsx"'
    wb.save(response)
    _audit(None, "Téléchargement rapport Excel", nom_fichier)
    return response


# ─── IMPORT SIG ──────────────────────────────────────────────────


@login_required
def importer_geometrie(request):
    if request.method != "POST":
        return redirect('index_cartographie')

    est_ajax = request.POST.get('ajax') == '1' or request.GET.get('ajax') == '1'

    fichier = request.FILES.get('fichier_geom')
    nom_couche = request.POST.get('nom_couche', '').strip()
    if not fichier or not nom_couche:
        if est_ajax:
            return JsonResponse({'ok': False, 'erreur': "Nom de couche et fichier requis."}, status=400)
        messages.error(request, "Nom de couche et fichier requis.")
        return redirect('index_cartographie')

    nom_fichier = fichier.name.lower()
    contenu = fichier.read()

    try:
        if nom_fichier.endswith('.kml') or nom_fichier.endswith('.kmz'):
            couche = _importer_kml_kmz(nom_couche, contenu, nom_fichier)
        elif nom_fichier.endswith('.gpx'):
            couche = _importer_gpx(nom_couche, contenu, nom_fichier)
        elif nom_fichier.endswith('.zip') or nom_fichier.endswith('.shp'):
            couche = _importer_shapefile(nom_couche, contenu, nom_fichier)
        elif nom_fichier.endswith('.geojson') or nom_fichier.endswith('.json'):
            couche = _importer_geojson(nom_couche, contenu, nom_fichier)
        elif nom_fichier.endswith('.csv'):
            couche = _importer_csv(nom_couche, contenu, nom_fichier)
        else:
            if est_ajax:
                return JsonResponse({'ok': False, 'erreur': "Format non supporté. Utilisez .geojson/.json/.csv/.kml/.kmz/.gpx/.shp/.zip"}, status=400)
            messages.error(request, "Format non supporté. Utilisez .geojson/.json/.csv/.kml/.kmz/.gpx/.shp/.zip")
            return redirect('index_cartographie')
    except AmbiguiteCoordonnees as e:
        if est_ajax:
            return JsonResponse({'ok': False, 'erreur': "Plusieurs colonnes de coordonnées plausibles : utilisez la page d'import classique pour choisir les colonnes."}, status=400)
        if len(contenu) > 8 * 1024 * 1024:
            messages.error(request, "Fichier trop volumineux pour la sélection interactive des colonnes.")
            return redirect('index_cartographie')
        style_brut = request.POST.get('style_options')
        if style_brut:
            try:
                style_valide = json.loads(style_brut)
            except (ValueError, TypeError):
                style_valide = {}
        else:
            style_valide = {}
        request.session['import_ambig'] = {
            'nom_fichier': nom_fichier,
            'nom_couche': nom_couche,
            'contenu': base64.b64encode(contenu).decode('ascii'),
            'candidats': e.candidats,
            'style_options': style_valide,
        }
        messages.info(request, "Plusieurs colonnes de coordonnées plausibles ont été détectées. Choisissez celles à utiliser.")
        return redirect('import_choix_colonnes')
    except Exception as e:
        if est_ajax:
            return JsonResponse({'ok': False, 'erreur': f"Erreur lors de l'import : {e}"}, status=400)
        messages.error(request, f"Erreur lors de l'import : {e}")
        return redirect('index_cartographie')

    style_brut = request.POST.get('style_options')
    if style_brut:
        try:
            _appliquer_style_couche(couche, json.loads(style_brut))
        except (ValueError, TypeError):
            pass

    if est_ajax:
        nb = couche.geometries.count()
        projet_actif = _projet_actif(request)
        if projet_actif is not None:
            couche.projet = projet_actif
            couche.save(update_fields=['projet'])
        _audit(request, "Import SIG", f"Couche {couche.nom} - {nb} géométries")
        return JsonResponse({
            'ok': True, 'couche_id': couche.pk, 'nom': couche.nom,
            'importes': nb, 'type': couche.type_geometrie,
            'style': couche.style_options or {},
        })

    return _terminer_import(request, couche, nom_fichier)


def _terminer_import(request, couche, nom_fichier):
    """Finalise un import : projet actif, audit, rapport, redirection avec zoom."""
    nb = couche.geometries.count()
    projet_actif = _projet_actif(request)
    if projet_actif is not None:
        couche.projet = projet_actif
        couche.save()
    _audit(request, "Import SIG", f"Couche {couche.nom} - {nb} géométries")
    rapport = getattr(couche, '_rapport', None)
    if not isinstance(rapport, dict):
        rapport = {
            'fichier': nom_fichier,
            'format': nom_fichier.rsplit('.', 1)[-1].upper() if '.' in nom_fichier else 'FICHIER',
            'lignes_analysees': nb,
            'points_detectes': nb,
            'points_crees': nb,
            'sans_coordonnees': 0,
            'invalides': 0,
            'doublons_position': 0,
            'colonnes': None,
            'srid': getattr(couche, 'srid', 4326),
            'crs': 'WGS84 (EPSG:4326)' if getattr(couche, 'srid', 4326) == 4326 else f'EPSG:{couche.srid}',
            'dimension': '2D',
        }
    request.session['rapport_import'] = rapport
    messages.success(request, f"Couche « {couche.nom} » importée ({nb} géométrie(s)).")
    return redirect(f"{reverse('index_cartographie')}?importe={couche.pk}")


def import_choix_colonnes(request):
    """Page de choix interactif des colonnes de coordonnées en cas d'ambiguïté."""
    donnees = request.session.get('import_ambig')
    if not donnees:
        messages.warning(request, "Aucune importation en attente de choix de colonnes.")
        return redirect('index_cartographie')

    if request.method == 'POST':
        col_lon = request.POST.get('col_lon', '').strip()
        col_lat = request.POST.get('col_lat', '').strip()
        col_alt = request.POST.get('col_alt', '').strip()
        if not col_lon or not col_lat:
            messages.error(request, "Sélectionnez une colonne pour la latitude et une pour la longitude.")
            return redirect('import_choix_colonnes')
        try:
            contenu = base64.b64decode(donnees['contenu'])
            couche = _importer_csv(
                donnees['nom_couche'], contenu, donnees['nom_fichier'],
                colonnes_forcees={'col_lon': col_lon, 'col_lat': col_lat, 'col_alt': col_alt or None})
            _appliquer_style_couche(couche, donnees.get('style_options') or {})
        except Exception as e:
            messages.error(request, f"Erreur lors de l'import : {e}")
            return redirect('import_choix_colonnes')
        request.session.pop('import_ambig', None)
        return _terminer_import(request, couche, donnees['nom_fichier'])

    return render(request, 'cartographie/import_choix_colonnes.html', {
        'fichier': donnees['nom_fichier'],
        'couche': donnees['nom_couche'],
        'candidats': donnees['candidats'],
    })


_COULEURS_OK = re.compile(r'^#[0-9a-fA-F]{6}$')
_SYMBOLES_OK = {'cercle', 'carre', 'triangle', 'losange', 'croix', 'etoile', 'point'}


def _appliquer_style_couche(couche, style=None, couleur=None):
    """Applique des options de style validées sur une couche (import ou mise à jour).

    style : dict JSON client {couleur, symbole, taille, opacite, etiquette,
            categories: {champ, classes: [{valeur, couleur, label}]}}
    Retourne le dict nettoyé (utile pour la réponse API)."""
    brut = style if isinstance(style, dict) else {}
    propre = {}

    def _couleur_valide(v, defaut='#3388ff'):
        return v if isinstance(v, str) and _COULEURS_OK.match(v) else defaut

    couleur_finale = _couleur_valide(brut.get('couleur'), _couleur_valide(couleur, '#3388ff'))
    propre['couleur'] = couleur_finale

    symbole = str(brut.get('symbole', 'cercle')).lower()
    propre['symbole'] = symbole if symbole in _SYMBOLES_OK else 'cercle'

    try:
        taille = float(brut.get('taille', 6))
    except (TypeError, ValueError):
        taille = 6.0
    propre['taille'] = round(min(max(taille, 1.0), 20.0), 1)

    try:
        opacite = float(brut.get('opacite', 1))
    except (TypeError, ValueError):
        opacite = 1.0
    propre['opacite'] = round(min(max(opacite, 0.0), 1.0), 2)

    etiquette = str(brut.get('etiquette', '') or '').strip()
    propre['etiquette'] = etiquette[:100]

    categories = brut.get('categories')
    if isinstance(categories, dict) and isinstance(categories.get('champ'), str) and categories['champ'].strip():
        classes = categories.get('classes')
        if isinstance(classes, list) and classes:
            cls_propres = []
            for k in classes[:100]:
                if not isinstance(k, dict):
                    continue
                valeur = k.get('valeur')
                if valeur is None or str(valeur) == '':
                    continue
                cls_propres.append({
                    'valeur': str(valeur)[:200],
                    'couleur': _couleur_valide(k.get('couleur'), couleur_finale),
                    'label': str(k.get('label') or valeur)[:200],
                })
            if cls_propres:
                propre['categories'] = {'champ': categories['champ'].strip()[:100], 'classes': cls_propres}

    couche.style_options = propre
    couche.style_couleur = couleur_finale
    couche.save(update_fields=['style_options', 'style_couleur'])
    return propre


def _importer_kml_kmz(nom_couche, contenu, nom_fichier):
    if nom_fichier.endswith('.kmz'):
        with zipfile.ZipFile(io.BytesIO(contenu)) as z:
            kml_files = [n for n in z.namelist() if n.endswith('.kml')]
            if not kml_files:
                raise ValueError("Aucun fichier KML trouvé dans le KMZ.")
            contenu = z.read(kml_files[0])

    root = ET.fromstring(contenu)
    ns = {'kml': 'http://www.opengis.net/kml/2.2'}
    placemarks = root.findall('.//kml:Placemark', ns) or root.findall('.//Placemark')

    def premier(scope, tag, chemin=False):
        """Premier élément trouvé (namespace d'abord). Pas de `or` : les éléments
        XML vides (0 enfant) sont falsy en Python >= 3.13 (bool = len())."""
        prefix = './/' if chemin else ''
        e = scope.find(prefix + 'kml:' + tag, ns)
        if e is None:
            e = scope.find(prefix + tag)
        return e

    def texte_el(pm, tag):
        e = premier(pm, tag)
        return e.text.strip() if e is not None and e.text else ''

    def proprietes_kml(pm):
        """Tous les attributs du Placemark : Name, Description, ExtendedData, SchemaData, styleUrl…"""
        props = {'nom': texte_el(pm, 'name') or 'Sans nom'}
        desc = texte_el(pm, 'description')
        if desc:
            props['description'] = desc
        for tag in ('Snippet', 'styleUrl'):
            v = texte_el(pm, tag)
            if v:
                props[tag] = v
        ext = premier(pm, 'ExtendedData')
        if ext is not None:
            for data in ext.findall('kml:Data', ns) or ext.findall('Data'):
                nm = data.get('name')
                if nm:
                    val = premier(data, 'value')
                    props[nm] = val.text.strip() if val is not None and val.text else ''
            for schema in ext.findall('kml:SchemaData', ns) or ext.findall('SchemaData'):
                for sd in schema.findall('kml:SimpleData', ns) or schema.findall('SimpleData'):
                    nm = sd.get('name')
                    if nm:
                        props[nm] = sd.text.strip() if sd.text else ''
        for att in ('id', 'targetId'):
            v = pm.get(att)
            if v:
                props[att] = v
        return props

    def coords_kml(coord_el, props):
        if coord_el is None or not coord_el.text:
            return None
        parts_list = [p.strip().split(',') for p in coord_el.text.strip().split() if p.strip()]
        pts = []
        for p in parts_list:
            if len(p) >= 2:
                pt = [float(p[0]), float(p[1])]
                if len(p) >= 3 and p[2]:
                    try:
                        alt = float(p[2])
                        pt.append(alt)
                        props.setdefault('altitude', alt)
                    except ValueError:
                        pass
                pts.append(pt)
        return pts or None

    def sous_geometries(el, props):
        """(type, coordonnees GeoJSON) pour un élément Point/LineString/Polygon."""
        tag = el.tag.rsplit('}', 1)[-1]
        if tag == 'Point':
            pts = coords_kml(premier(el, 'coordinates'), props)
            return [('Point', pts[0])] if pts else []
        if tag == 'LineString':
            pts = coords_kml(premier(el, 'coordinates'), props)
            return [('LineString', pts)] if pts else []
        if tag == 'Polygon':
            outer = premier(el, 'outerBoundaryIs/LinearRing/coordinates', chemin=True)
            pts = coords_kml(outer, props)
            return [('Polygon', [pts])] if pts else []
        return []

    geometries = []
    for pm in placemarks:
        props = proprietes_kml(pm)
        multi = premier(pm, 'MultiGeometry')
        if multi is not None:
            trouves = []
            for child in list(multi):
                trouves.extend(sous_geometries(child, props))
        else:
            trouves = []
            point = premier(pm, 'Point')
            ligne = premier(pm, 'LineString')
            poly = premier(pm, 'Polygon')
            if point is not None:
                trouves.extend(sous_geometries(point, props))
            if ligne is not None:
                trouves.extend(sous_geometries(ligne, props))
            if poly is not None:
                trouves.extend(sous_geometries(poly, props))
        for gtype, gcoords in trouves:
            geometries.append({'type': gtype, 'coords': gcoords, 'proprietes': dict(props)})

    if not geometries:
        raise ValueError("Aucune géométrie trouvée dans le fichier KML/KMZ.")

    type_geo = _type_geometries_vers_couche([g['type'] for g in geometries])
    couche = CoucheGeometrie.objects.create(nom=nom_couche, type_geometrie=type_geo, fichier_source=nom_fichier, srid=4326)
    for g in geometries:
        Geometrie.objects.create(couche=couche, type=g['type'], coordonnees=g['coords'], proprietes=g['proprietes'])
    dim = '3D' if any(any(len(pt) >= 3 for pt in g['coords']) if isinstance(g['coords'][0], list) else len(g['coords']) >= 3 for g in geometries) else '2D'
    couche._rapport = {
        'fichier': nom_fichier,
        'format': 'KMZ' if nom_fichier.endswith('.kmz') else 'KML',
        'lignes_analysees': len(placemarks),
        'points_crees': len(geometries),
        'sans_coordonnees': 0,
        'invalides': 0,
        'doublons_position': 0,
        'colonnes': {'longitude': 'coordinates[0]', 'latitude': 'coordinates[1]', 'altitude': 'coordinates[2]'},
        'srid': 4326,
        'crs': 'WGS84 (EPSG:4326)',
        'dimension': dim,
    }
    return couche


def _importer_gpx(nom_couche, contenu, nom_fichier='fichier.gpx'):
    root = ET.fromstring(contenu)
    ns = {'gpx': 'http://www.topografix.com/GPX/1/1'}
    geometries = []

    def texte(el, tag):
        e = el.find(f'gpx:{tag}', ns)
        return e.text.strip() if e is not None and e.text else ''

    def point_wpt(el):
        lat = float(el.get('lat'))
        lon = float(el.get('lon'))
        props = {'nom': texte(el, 'name') or 'Sans nom'}
        for tag in ('desc', 'ele', 'time', 'sym', 'type', 'cmt', 'fix'):
            v = texte(el, tag)
            if v:
                if tag == 'ele':
                    try:
                        v = float(v)
                    except ValueError:
                        pass
                props[tag] = v
        coords = [lon, lat]
        if isinstance(props.get('ele'), (int, float)):
            coords.append(props['ele'])
        return {'type': 'Point', 'coords': coords, 'proprietes': props}

    for wpt in root.findall('.//gpx:wpt', ns):
        geometries.append(point_wpt(wpt))

    for trk in root.findall('.//gpx:trk', ns):
        props = {'nom': texte(trk, 'name') or 'Track'}
        desc = texte(trk, 'desc')
        if desc:
            props['description'] = desc
        t = texte(trk, 'time')
        if t:
            props['time'] = t
        trkpts = trk.findall('.//gpx:trkpt', ns)
        coords = []
        alts = []
        temps = []
        for pt in trkpts:
            coords.append([float(pt.get('lon')), float(pt.get('lat'))])
            e = texte(pt, 'ele')
            if e:
                try:
                    alts.append(float(e))
                except ValueError:
                    pass
            tm = texte(pt, 'time')
            if tm:
                temps.append(tm)
        if coords:
            if alts:
                props['altitudes'] = alts
            if temps:
                props['times'] = temps
            geometries.append({'type': 'LineString', 'coords': coords, 'proprietes': props})

    for rte in root.findall('.//gpx:rte', ns):
        props = {'nom': texte(rte, 'name') or 'Route'}
        desc = texte(rte, 'desc')
        if desc:
            props['description'] = desc
        rtepts = rte.findall('.//gpx:rtept', ns)
        coords = [[float(pt.get('lon')), float(pt.get('lat'))] for pt in rtepts]
        if coords:
            geometries.append({'type': 'LineString', 'coords': coords, 'proprietes': props})

    if not geometries:
        raise ValueError("Aucune géométrie trouvée dans le fichier GPX.")

    type_geo = _type_geometries_vers_couche([g['type'] for g in geometries])
    couche = CoucheGeometrie.objects.create(nom=nom_couche, type_geometrie=type_geo, fichier_source='fichier.gpx', srid=4326)
    for g in geometries:
        Geometrie.objects.create(couche=couche, type=g['type'], coordonnees=g['coords'], proprietes=g['proprietes'])
    nb_wpts = len(root.findall('.//gpx:wpt', ns))
    couche._rapport = {
        'fichier': nom_fichier,
        'format': 'GPX',
        'lignes_analysees': nb_wpts,
        'points_crees': nb_wpts,
        'sans_coordonnees': 0,
        'invalides': 0,
        'doublons_position': 0,
        'colonnes': {'longitude': 'lon', 'latitude': 'lat', 'altitude': 'ele'},
        'srid': 4326,
        'crs': 'WGS84 (EPSG:4326)',
        'dimension': '3D' if any(isinstance(g.get('proprietes', {}).get('ele'), (int, float)) for g in geometries if g['type'] == 'Point') else '2D',
    }
    return couche


def _importer_shapefile(nom_couche, contenu, nom_fichier):
    import shapefile

    if nom_fichier.endswith('.zip'):
        with zipfile.ZipFile(io.BytesIO(contenu)) as z:
            shp_files = [n for n in z.namelist() if n.endswith('.shp')]
            if not shp_files:
                raise ValueError("Aucun fichier .shp trouvé dans le ZIP.")
            memoire = {}
            for name in z.namelist():
                ext = name.rsplit('.', 1)[-1] if '.' in name else ''
                base = name.rsplit('.', 1)[0] if '.' in name else name
                if ext in ('shp', 'shx', 'dbf', 'prj'):
                    memoire[f'{base}.{ext}'] = z.read(name)
            with shapefile.Reader(shp=memoire) as reader:
                geometries = _lire_shapefile_reader(reader)
    elif nom_fichier.endswith('.shp'):
        raise ValueError("Pour un Shapefile, veuillez fournir un dossier ZIP contenant .shp, .shx, .dbf.")
    else:
        raise ValueError("Format non supporté.")

    if not geometries:
        raise ValueError("Aucune géométrie trouvée dans le Shapefile.")

    type_geo = _type_shapefile_vers_couche(geometries[0]['type'])
    couche = CoucheGeometrie.objects.create(nom=nom_couche, type_geometrie=type_geo, fichier_source=nom_fichier)
    for g in geometries:
        Geometrie.objects.create(couche=couche, type=g['type'], coordonnees=g['coords'], proprietes=g['proprietes'])
    return couche


def _lire_shapefile_reader(reader):
    geometries = []
    field_names = [f[0] for f in reader.fields if f[0] not in ('DeletionFlag',)]
    for shape_rec in reader.shapeRecords():
        shape = shape_rec.shape
        props = dict(zip(field_names, shape_rec.record)) if shape_rec.record else {}
        if shape.shapeType in (1, 11):
            coords = [shape.points[0][0], shape.points[0][1]]
            geometries.append({'type': 'Point', 'coords': coords, 'proprietes': props})
        elif shape.shapeType in (3, 13):
            coords = [[p[0], p[1]] for p in shape.points]
            geometries.append({'type': 'LineString', 'coords': coords, 'proprietes': props})
        elif shape.shapeType in (5, 15):
            parts = list(shape.parts) + [len(shape.points)]
            rings = []
            for i in range(len(parts) - 1):
                ring = [[shape.points[j][0], shape.points[j][1]] for j in range(parts[i], parts[i + 1])]
                rings.append(ring)
            geometries.append({'type': 'Polygon', 'coords': rings, 'proprietes': props})
    return geometries


def _type_shapefile_vers_couche(type_geom):
    return {'Point': 'point', 'LineString': 'ligne', 'Polygon': 'polygone'}.get(type_geom, 'point')


def _type_geometries_vers_couche(types):
    """Type de couche à partir des types GeoJSON des entités (gère les fichiers mixtes)."""
    if not types:
        return 'point'
    if all(t == 'Point' for t in types):
        return 'point'
    if any(t == 'Polygon' for t in types):
        return 'polygone'
    if any(t == 'LineString' for t in types):
        return 'ligne'
    return 'point'


def _exploser_geometrie(gtype, coords):
    """Décompose une géométrie GeoJSON (y compris Multi* et GeometryCollection)
    en une liste de géométries simples (Point, LineString, Polygon)."""
    if gtype == 'Point':
        return [('Point', coords)]
    if gtype == 'MultiPoint':
        return [('Point', c) for c in coords]
    if gtype == 'LineString':
        return [('LineString', coords)]
    if gtype == 'MultiLineString':
        return [('LineString', c) for c in coords]
    if gtype == 'Polygon':
        return [('Polygon', coords)]
    if gtype == 'MultiPolygon':
        return [('Polygon', c) for c in coords]
    if gtype == 'GeometryCollection':
        out = []
        for g in coords or []:
            out.extend(_exploser_geometrie(g.get('type'), g.get('coordinates')))
        return out
    return []


def _importer_geojson(nom_couche, contenu, nom_fichier):
    data = json.loads(contenu.decode('utf-8-sig', errors='replace'))
    dtype = data.get('type') if isinstance(data, dict) else None
    if dtype == 'FeatureCollection':
        features = data.get('features') or []
    elif dtype == 'Feature':
        features = [data]
    elif dtype in ('Point', 'MultiPoint', 'LineString', 'MultiLineString', 'Polygon', 'MultiPolygon', 'GeometryCollection'):
        features = [{'type': 'Feature', 'geometry': data, 'properties': {}}]
    else:
        raise ValueError("GeoJSON non reconnu : type attendu Feature, FeatureCollection ou géométrie.")

    geometries = []
    for f in features:
        if not isinstance(f, dict):
            continue
        g = f.get('geometry')
        if not isinstance(g, dict):
            continue
        gtype = g.get('type')
        gcoords = g.get('coordinates')
        if gtype not in ('Point', 'MultiPoint', 'LineString', 'MultiLineString', 'Polygon', 'MultiPolygon', 'GeometryCollection') or gcoords is None:
            continue
        props = f.get('properties')
        if not isinstance(props, dict):
            props = {'valeur': props} if props is not None else {}
        props = {k: (json.dumps(v, ensure_ascii=False) if isinstance(v, (dict, list)) else v)
                 for k, v in props.items()}
        if f.get('id') is not None:
            props.setdefault('id', f['id'])
        parties = _exploser_geometrie(gtype, gcoords)
        for idx, (ptype, pcoords) in enumerate(parties):
            p = dict(props)
            if len(parties) > 1:
                p['_partie'] = {'type': gtype, 'numero': idx + 1, 'total': len(parties)}
            geometries.append({'type': ptype, 'coords': pcoords, 'proprietes': p})

    if not geometries:
        raise ValueError("Aucune géométrie trouvée dans le fichier GeoJSON.")

    srid = 4326
    crs_nom = 'WGS84 (EPSG:4326)'
    crs_membre = data.get('crs') if isinstance(data, dict) else None
    if isinstance(crs_membre, dict):
        nom_crs = (crs_membre.get('properties') or {}).get('name', '')
        if isinstance(nom_crs, str):
            for code in ('32735', '3857', '2154', '4326', '32635'):
                if code in nom_crs:
                    srid = int(code)
                    break
            if srid != 4326:
                crs_nom = f'CRS déclaré : {nom_crs}'
    dim = '3D' if any(
        (any(len(p) >= 3 for p in g['coords']) if isinstance(g['coords'][0], list) else len(g['coords']) >= 3)
        for g in geometries) else '2D'
    type_geo = _type_geometries_vers_couche([g['type'] for g in geometries])
    couche = CoucheGeometrie.objects.create(nom=nom_couche, type_geometrie=type_geo, fichier_source=nom_fichier, srid=srid)
    for g in geometries:
        Geometrie.objects.create(couche=couche, type=g['type'], coordonnees=g['coords'], proprietes=g['proprietes'])
    nb_features = len(data['features']) if isinstance(data, dict) and isinstance(data.get('features'), list) else len(geometries)
    couche._rapport = {
        'fichier': nom_fichier,
        'format': 'GEOJSON',
        'lignes_analysees': nb_features,
        'points_crees': len(geometries),
        'sans_coordonnees': 0,
        'invalides': 0,
        'doublons_position': 0,
        'colonnes': {'longitude': 'geometry.coordinates[0]', 'latitude': 'geometry.coordinates[1]', 'altitude': 'geometry.coordinates[2]'},
        'srid': srid,
        'crs': crs_nom,
        'dimension': dim,
    }
    return couche


def _cle_colonne(nom):
    """Normalise un nom de colonne : minuscules, sans espaces ni symboles."""
    return ''.join(c for c in str(nom or '').strip().lower() if c.isalnum())


def _coord_decimal(v):
    """Convertit une valeur en degrés décimaux. Supporte le décimal simple
    ('24.456', '-0.123', '29.12345°') et le DMS ('1°14'04"S', '29°07'24"E',
    '12° 30′ 15″ N'). Retourne None si non convertible."""
    s = str(v or '').strip()
    if not s:
        return None
    if ',' in s and '.' not in s:
        s = s.replace(',', '.')
    elif ',' in s and '.' in s:
        s = s.replace(',', '')
    try:
        return float(s)
    except ValueError:
        pass
    m = re.match(
        r'^\s*([+-]?[\d.]+)\s*(?:°|º|d|deg)?'
        r'(?:\s*([\d.]+)\s*(?:[\'’‘′]|m|min))?'
        r'(?:\s*([\d.]+)\s*(?:["”’″]|s|sec))?'
        r'\s*([NSEOWnseow])?\s*$', s)
    if not m or m.group(1) is None:
        return None
    try:
        val = float(m.group(1))
        if m.group(2) is not None:
            val += float(m.group(2)) / 60.0
        if m.group(3) is not None:
            val += float(m.group(3)) / 3600.0
    except ValueError:
        return None
    if m.group(4) in ('S', 's', 'W', 'w'):
        return -abs(val)
    if m.group(4) in ('N', 'n', 'E', 'e'):
        return abs(val)
    return val


class AmbiguiteCoordonnees(Exception):
    """Plusieurs paires de colonnes de coordonnées plausibles — choix utilisateur requis."""

    def __init__(self, candidats, nom_fichier):
        super().__init__("Plusieurs colonnes de coordonnées plausibles détectées.")
        self.candidats = candidats
        self.nom_fichier = nom_fichier


_AXE_LONGITUDE = {
    'exacts': ('longitude', 'longitud'),
    'courts': ('lon', 'lng', 'long', 'coordx', 'xcoord', 'xcoordinate'),
    'racines': ('lon', 'long', 'lng'),
    'borne': 180,
}
_AXE_LATITUDE = {
    'exacts': ('latitude', 'latitud', 'lattitude'),
    'courts': ('lat', 'coordy', 'ycoord', 'ycoordinate'),
    'racines': ('lat',),
    'borne': 90,
}
_AXE_ALTITUDE = {
    'exacts': ('altitude', 'altitud', 'elevation', 'elevationm', 'height', 'heightm', 'hauteur'),
    'courts': ('alt', 'elev', 'coordz', 'zcoord', 'zcoordinate'),
    'racines': ('alt', 'elev', 'height', 'hauteur'),
    'borne': 9000,
}

# Parties DMS (Degrés / Minutes / Secondes) à écarter des candidats par nom.
_DMS_PARTS = (
    'londeg', 'lonmin', 'lonsec', 'latdeg', 'latmin', 'latsec',
    'lond', 'lonm', 'lons', 'latd', 'latm', 'lats',
)


def _fit_plage(vals, borne):
    return sum(1 for v in vals if -borne <= v <= borne) / max(len(vals), 1)


def _est_partie_dms(n):
    return n in _DMS_PARTS or n.endswith(('deg', 'min', 'sec', 'degree', 'minute', 'seconde')) or n in ('d', 'm', 's')


def _score_colonne_axe(n, axe, vals):
    """Score d'une colonne pour un axe : nom (0-60) + adéquation des valeurs (0-40)."""
    nom = 0
    if n in axe['exacts']:
        nom = 60
    elif n in axe['courts']:
        nom = 50
    elif n in ('x', 'y', 'z'):
        if (n == 'x' and axe['borne'] == 180) or (n == 'y' and axe['borne'] == 90) or (n == 'z' and axe['borne'] == 9000):
            nom = 35
    elif n.startswith('coord') or n.endswith('coord') or n.endswith('coordinate'):
        nom = 40
    elif any(r in n for r in axe['racines']):
        nom = 25
    fit = _fit_plage(vals, axe['borne']) if vals else 0
    return nom + int(fit * 40)


def _detecter_colonnes_coord(champs, lignes):
    """Détection intelligente des colonnes Longitude / Latitude / Altitude.

    Priorités :
    1) Alias standards + variantes (insensible à la casse, espaces, _ et -) avec
       scoring nom + valeurs ; plusieurs candidats départagés intelligemment.
    2) Ambiguïté non résolue → ('AMBIGU', candidats, None) pour choix utilisateur.
    3) DMS séparé (Lon_Deg/Lon_Min/Lon_Sec + Lat_Deg/Lat_Min/Lat_Sec).
    4) Analyse des valeurs : X → Longitude, Y → Latitude, Z → Altitude, plages.
    Retourne (col_lon, col_lat, col_alt) ou ('DMS', dms, None) ou ('AMBIGU', candidats, None)."""
    norm = {_cle_colonne(c): c for c in champs}

    def echantillon_vals(col, n=60):
        vals = []
        for l in lignes:
            if len(vals) >= n:
                break
            d = _coord_decimal(l.get(col))
            if d is not None:
                vals.append(d)
        return vals

    # ── 1) Candidats par nom + valeurs, avec scoring ─────────────
    candidats_axes = {}
    for axe, cfg in (('longitude', _AXE_LONGITUDE), ('latitude', _AXE_LATITUDE), ('altitude', _AXE_ALTITUDE)):
        liste = []
        for n, c in norm.items():
            if _est_partie_dms(n):
                continue
            vals = echantillon_vals(c)
            sc = _score_colonne_axe(n, cfg, vals)
            if sc > 0:
                liste.append((sc, c))
        liste.sort(reverse=True)
        candidats_axes[axe] = liste

    lon_liste = candidats_axes['longitude']
    lat_liste = candidats_axes['latitude']

    def ambiguite_axe(liste):
        """Deux candidats nommés quasi ex-aequo sur le même axe → choix requis."""
        if len(liste) < 2:
            return False
        s1, _ = liste[0]
        s2, _ = liste[1]
        return s1 >= 45 and s2 >= 45 and (s1 - s2) <= 5

    if ambiguite_axe(lon_liste) or ambiguite_axe(lat_liste):
        alt_page = [c for sc, c in candidats_axes['altitude'] if sc > 40]
        return 'AMBIGU', {
            'longitude': [c for _, c in lon_liste[:6]],
            'latitude': [c for _, c in lat_liste[:6]],
            'altitude': alt_page[:6],
        }, None

    meilleure = None
    for s1, c1 in lon_liste[:6]:
        for s2, c2 in lat_liste[:6]:
            if c1 == c2:
                continue
            total = s1 + s2
            if meilleure is None or total > meilleure[0]:
                meilleure = (total, c1, c2)

    if meilleure is not None:
        score_pair, col_lon, col_lat = meilleure
        v_lon = echantillon_vals(col_lon)
        v_lat = echantillon_vals(col_lat)
        if score_pair >= 100 and _fit_plage(v_lon, 180) >= 0.5 and _fit_plage(v_lat, 90) >= 0.5:
            # Désambiguïsation X/Y par les valeurs : X prend des latitudes, Y des longitudes.
            xc, yc = norm.get('x'), norm.get('y')
            if xc == col_lon and yc == col_lat:
                fx = _fit_plage(echantillon_vals(xc), 90)
                fy = _fit_plage(echantillon_vals(yc), 90)
                if fx == 1 and fy < 1:
                    col_lon, col_lat = yc, xc

            col_alt = None
            for sc, c in candidats_axes['altitude']:
                if c not in (col_lon, col_lat) and sc > 40:
                    col_alt = c
                    break
            return col_lon, col_lat, col_alt

    # ── 2) DMS séparé : Lon_Deg/Lon_Min/Lon_Sec + Lat_Deg/Lat_Min/Lat_Sec ──
    dms = {
        'lon_deg': norm.get('londeg') or norm.get('lond'),
        'lon_min': norm.get('lonmin') or norm.get('lonm'),
        'lon_sec': norm.get('lonsec') or norm.get('lons'),
        'lat_deg': norm.get('latdeg') or norm.get('latd'),
        'lat_min': norm.get('latmin') or norm.get('latm'),
        'lat_sec': norm.get('latsec') or norm.get('lats'),
    }
    if dms['lon_deg'] and dms['lat_deg']:
        return 'DMS', dms, None

    # ── 3) Analyse intelligente des valeurs : X → Longitude, Y → Latitude, Z → Altitude ──
    numeriques = {}
    for c in champs:
        vals = echantillon_vals(c)
        if len(vals) >= max(2, int(len(lignes) * 0.6)) if lignes else len(vals) >= 2:
            numeriques[c] = vals
    if len(numeriques) >= 2:
        xc = next((norm[k] for k in ('x', 'coordx', 'xcoord') if norm.get(k) in numeriques), None)
        yc = next((norm[k] for k in ('y', 'coordy', 'ycoord') if norm.get(k) in numeriques), None)
        zc = next((norm[k] for k in ('z', 'coordz', 'zcoord', 'altitude', 'alt', 'elevation', 'elev', 'height') if norm.get(k) in numeriques), None)
        if xc and yc and _fit_plage(numeriques[xc], 180) >= 0.8 and _fit_plage(numeriques[yc], 90) >= 0.8:
            col_alt = zc if (zc and _fit_plage(numeriques[zc], 180) < 0.8) else None
            return xc, yc, col_alt

        # Dernier recours : scoring des plages (lat ⊂ lon)
        tri = sorted(numeriques.keys(), key=lambda c: _fit_plage(numeriques[c], 90), reverse=True)
        col_lat = tri[0]
        col_lon = next((c for c in tri[1:] if _fit_plage(numeriques[c], 180) >= 0.8), None)
        if col_lon is not None and _fit_plage(numeriques[col_lat], 90) >= 0.8:
            col_alt = next((c for c in numeriques if c not in (col_lon, col_lat) and _fit_plage(numeriques[c], 180) < 0.8), None)
            return col_lon, col_lat, col_alt

    return None, None, None


def _importer_csv(nom_couche, contenu, nom_fichier, colonnes_forcees=None):
    texte = contenu.decode('utf-8-sig', errors='replace')
    delim = ','
    try:
        snif = csv.Sniffer().sniff(texte[:4096], delimiters=',;\t|')
        if snif.delimiter:
            delim = snif.delimiter
    except csv.Error:
        pass
    reader = csv.DictReader(io.StringIO(texte), delimiter=delim)
    champs = [c for c in (reader.fieldnames or []) if c is not None]
    if not champs:
        raise ValueError("Fichier CSV vide ou illisible.")
    lignes = list(reader)

    if colonnes_forcees:
        col_lon = colonnes_forcees.get('col_lon') or ''
        col_lat = colonnes_forcees.get('col_lat') or ''
        col_alt = colonnes_forcees.get('col_alt') or None
        if col_lon not in champs or col_lat not in champs:
            raise ValueError("Colonnes choisies introuvables dans le fichier.")
        dms_mode = False
    else:
        col_lon, col_lat, col_alt = _detecter_colonnes_coord(champs, lignes)
        if col_lon == 'AMBIGU':
            raise AmbiguiteCoordonnees(col_lat, nom_fichier)
        dms_mode = col_lon == 'DMS'
        if not col_lon or not col_lat or col_lon == col_lat:
            raise ValueError(
                "Aucune paire de coordonnées Latitude / Longitude n'a été détectée dans ce fichier. "
                "Vérifiez qu'il contient des colonnes Latitude/Longitude (Lat, Latitude, Y…) "
                "et Longitude (Long, Lon, Longitude, X…) avec des valeurs numériques.")

    geometries = []
    nb_sans = 0
    nb_invalides = 0
    positions = {}
    dimension = 2
    for ligne in lignes:
        if dms_mode:
            lon = _coord_dms_parts(ligne, col_lat, lat=False)
            lat = _coord_dms_parts(ligne, col_lat, lat=True)
        else:
            lon = _coord_decimal(ligne.get(col_lon))
            lat = _coord_decimal(ligne.get(col_lat))
        if lon is None or lat is None:
            nb_sans += 1
            continue
        if not (-180 <= lon <= 180):
            nb_invalides += 1
            continue
        if not (-90 <= lat <= 90):
            nb_invalides += 1
            continue
        coords = [lon, lat]
        if col_alt:
            alt = _coord_decimal(ligne.get(col_alt))
            if alt is not None:
                coords.append(alt)
                dimension = 3
        props = {k: v for k, v in ligne.items()}
        geometries.append({'type': 'Point', 'coords': coords, 'proprietes': props})
        cle = (round(lon, 6), round(lat, 6))
        positions[cle] = positions.get(cle, 0) + 1

    if not geometries:
        raise ValueError("Aucun point valide trouvé dans le CSV (coordonnées manquantes ou invalides).")

    if dms_mode:
        col_lon_aff, col_lat_aff = col_lat['lon_deg'], col_lat['lat_deg']
    else:
        col_lon_aff, col_lat_aff = col_lon, col_lat
    couche = CoucheGeometrie.objects.create(nom=nom_couche, type_geometrie='point', fichier_source=nom_fichier, srid=4326)
    Geometrie.objects.bulk_create([
        Geometrie(couche=couche, type=g['type'], coordonnees=g['coords'], proprietes=g['proprietes'])
        for g in geometries
    ])
    couche._rapport = {
        'fichier': nom_fichier,
        'format': 'CSV',
        'lignes_analysees': len(lignes),
        'points_detectes': len(lignes) - nb_sans,
        'points_crees': len(geometries),
        'sans_coordonnees': nb_sans,
        'invalides': nb_invalides,
        'doublons_position': sum(n - 1 for n in positions.values() if n > 1),
        'colonnes': {'longitude': col_lon_aff, 'latitude': col_lat_aff, 'altitude': col_alt},
        'srid': 4326,
        'crs': 'WGS84 (EPSG:4326)',
        'dimension': '3D' if dimension == 3 else '2D',
    }
    return couche


def _coord_dms_parts(ligne, colonnes, lat=False):
    """Construit une valeur décimale à partir de colonnes DMS séparées."""
    cle = 'lat' if lat else 'lon'
    deg = _coord_decimal(ligne.get(colonnes.get(cle + '_deg')))
    if deg is None:
        return None
    val = deg
    mn = _coord_decimal(ligne.get(colonnes.get(cle + '_min')))
    if mn is not None:
        val += mn / 60.0
    sc = _coord_decimal(ligne.get(colonnes.get(cle + '_sec')))
    if sc is not None:
        val += sc / 3600.0
    return val


@login_required
def couche_delete(request, pk):
    couche = get_object_or_404(CoucheGeometrie, pk=pk)
    if request.method == "POST":
        if couche.fichier_kml:
            try:
                os.remove(couche.fichier_kml.path)
            except Exception:
                pass
        couche.delete()
        messages.success(request, f"Couche « {couche.nom} » supprimée.")
        return redirect('index_cartographie')
    return render(request, 'cartographie/couche_confirm_delete.html', {'couche': couche})


# ─── ÉDITION DES POINTS + MÉDIAS ───────────────────────────────


@login_required
def point_edit(request, pk):
    point = get_object_or_404(PointGeographique, pk=pk)
    if not (_est_admin(request.user) or point.auteur == request.user):
        messages.error(request, "Vous ne pouvez modifier que vos propres points.")
        return redirect('index_cartographie')
    if request.method == "POST":
        point.nom = request.POST.get('nom', point.nom)[:200]
        point.description = request.POST.get('description', '')
        point.categorie = request.POST.get('categorie', point.categorie)
        point.statut = request.POST.get('statut', point.statut)
        point.province = request.POST.get('province', '')
        point.commune = request.POST.get('commune', '')
        point.quartier = request.POST.get('quartier', '')
        lat, lng = _valider_coordonnees_wgs84(
            request.POST.get('latitude', point.latitude),
            request.POST.get('longitude', point.longitude))
        if lat is None:
            messages.error(request, "Coordonnées invalides (hors plage WGS84).")
            return redirect('point_edit', pk=pk)
        point.latitude = lat
        point.longitude = lng
        if request.FILES.get('photo'):
            point.photo = request.FILES['photo']
        point.save()
        _creer_medias_point(point, request.FILES.getlist('medias'), utilisateur=request.user, commentaire=request.POST.get('commentaire_medias', ''), date_prise_defaut=timezone.now())
        _audit(request, "Modification de point", f"Point #{pk} - {point.nom}")
        messages.success(request, "Point mis à jour.")
        return redirect('index_cartographie')
    return render(request, 'cartographie/point_form.html', {
        'point': point,
        'medias': point.medias.all(),
        'categories': PointGeographique.CATEGORIE_CHOICES,
        'statuts': PointGeographique.STATUT_CHOICES,
    })


@login_required
def point_delete(request, pk):
    point = get_object_or_404(PointGeographique, pk=pk)
    if not (_est_admin(request.user) or point.auteur == request.user):
        messages.error(request, "Vous ne pouvez supprimer que vos propres points.")
        return redirect('index_cartographie')
    if request.method == "POST":
        _audit(request, "Suppression de point", f"Point #{pk} - {point.nom}")
        point.delete()
        messages.success(request, "Point supprimé.")
        return redirect('index_cartographie')
    return render(request, 'cartographie/point_confirm_delete.html', {'point': point})


@login_required
def media_delete(request, pk):
    media = get_object_or_404(MediaPoint, pk=pk)
    point_id = media.point_id
    if request.method == "POST":
        _audit(request, "Suppression de média", f"Média #{pk} du point #{point_id}")
        media.delete()
        messages.success(request, "Média supprimé.")
        return redirect('point_edit', pk=point_id)
    return redirect('point_edit', pk=point_id)


# ─── IMPORT EXCEL INTELLIGENT (v2) ─────────────────────────────


@login_required
def import_page(request):
    return render(request, 'cartographie/import_excel.html')


@login_required
def importer_excel_v2(request):
    if request.method != "POST":
        return redirect('index_cartographie')
    try:
        data = json.loads(request.body)
    except json.JSONDecodeError:
        return JsonResponse({'ok': False, 'erreur': "Données invalides."})

    mapping = data.get('mapping', {})
    lignes = data.get('lignes', [])
    entetes = data.get('entetes') or []
    nom_fichier = (data.get('nom_fichier') or '').strip()[:255]
    format_fichier = (data.get('format') or '').strip()[:20]
    idx = {k: v for k, v in mapping.items() if v is not None and v != ''}

    if not lignes or 'latitude' not in idx or 'longitude' not in idx:
        return JsonResponse({'ok': False, 'erreur': "Colonnes Latitude/Longitude requises et au moins une ligne de données."})

    importes = 0
    invalides = 0
    doublons = 0
    doublons_maj = 0
    projet_doublons = ''
    for row in lignes:
        try:
            lat = float(str(row[idx['latitude']]).replace(',', '.'))
            lng = float(str(row[idx['longitude']]).replace(',', '.'))
        except (ValueError, TypeError, IndexError):
            invalides += 1
            continue
        if not (-90 <= lat <= 90) or not (-180 <= lng <= 180):
            invalides += 1
            continue
        nom = str(row[idx['nom']])[:200] if 'nom' in idx and row[idx['nom']] else f"Point {lat:.5f}, {lng:.5f}"
        categorie = str(row[idx['categorie']]).strip().lower() if 'categorie' in idx and row[idx['categorie']] else 'autre'
        categorie_valide = dict(PointGeographique.CATEGORIE_CHOICES).get(categorie)
        if not categorie_valide:
            categorie = 'autre'
        statut = str(row[idx['statut']]).strip().lower() if 'statut' in idx and row[idx['statut']] else 'actif'
        if statut not in dict(PointGeographique.STATUT_CHOICES):
            statut = 'actif'
        description = str(row[idx['description']]) if 'description' in idx and row[idx['description']] else ''
        province = str(row[idx['province']]) if 'province' in idx and row[idx['province']] else ''
        commune = str(row[idx['commune']]) if 'commune' in idx and row[idx['commune']] else ''

        donnees_pt = {}
        for i, h in enumerate(entetes):
            v = row[i] if i < len(row) else ''
            if h:
                donnees_pt[str(h).strip()] = (str(v) if v is not None else '').strip()

        existe = PointGeographique.objects.filter(
            nom=nom, latitude=lat, longitude=lng
        ).first()
        if existe:
            doublons += 1
            if not projet_doublons and existe.projet is not None:
                projet_doublons = existe.projet.nom
            if existe.donnees != donnees_pt or not existe.source_fichier:
                existe.donnees = donnees_pt
                existe.source_fichier = nom_fichier or existe.source_fichier
                existe.source_format = format_fichier or existe.source_format
                existe.save(update_fields=['donnees', 'source_fichier', 'source_format'])
                doublons_maj += 1
            continue

        PointGeographique.objects.create(
            nom=nom, description=description, latitude=lat, longitude=lng,
            categorie=categorie, statut=statut,
            province=province, commune=commune,
            donnees=donnees_pt,
            source_fichier=nom_fichier, source_format=format_fichier,
            projet=_projet_actif(request),
            auteur=request.user if request.user.is_authenticated else None,
        )
        importes += 1

    _audit(request, "Import Excel intelligent", f"{importes} importés, {doublons} doublons ({doublons_maj} actualisés), {invalides} invalides")
    projet_actif = _projet_actif(request)
    return JsonResponse({
        'ok': True, 'importes': importes, 'ignores': doublons + invalides,
        'doublons': doublons, 'invalides': invalides, 'doublons_maj': doublons_maj,
        'projet': projet_doublons or (projet_actif.nom if projet_actif is not None else ''),
    })


@login_required
def points_liste(request):
    """Tableau de tous les points géographiques, tous projets confondus."""
    qs = PointGeographique.objects.select_related('projet', 'auteur', 'activite').annotate(
        nb_photos=db_models.Count('medias', filter=db_models.Q(medias__type='photo')),
        nb_medias=db_models.Count('medias'),
    )

    q = (request.GET.get('q') or '').strip()
    if q:
        qs = qs.filter(db_models.Q(nom__icontains=q) | db_models.Q(province__icontains=q) |
                       db_models.Q(commune__icontains=q) | db_models.Q(description__icontains=q))
    projet_id = request.GET.get('projet') or ''
    if projet_id:
        qs = qs.filter(projet_id=projet_id)
    categorie = request.GET.get('categorie') or ''
    if categorie:
        qs = qs.filter(categorie=categorie)
    statut = request.GET.get('statut') or ''
    if statut:
        qs = qs.filter(statut=statut)

    points = qs.all()
    return render(request, 'cartographie/points_liste.html', {
        'points': points,
        'total': points.count(),
        'projets': Projet.objects.all().order_by('nom'),
        'categories': PointGeographique.CATEGORIE_CHOICES,
        'statuts': PointGeographique.STATUT_CHOICES,
        'f_q': q, 'f_projet': projet_id, 'f_categorie': categorie, 'f_statut': statut,
    })


@login_required
def table_attributaire(request):
    """Table attributaire professionnelle synchronisée avec la carte."""
    return render(request, 'cartographie/table_attributaire.html', {})


# ─── EXPORTS (GeoJSON / KML / GPX) ─────────────────────────────


@login_required
def export_points(request, format):
    points = PointGeographique.objects.select_related('auteur').filter(supprime=False)
    projet_actif = _projet_actif(request)
    if projet_actif is not None:
        points = points.filter(projet=projet_actif)
    points = points.all()
    date_str = date.today().strftime('%Y%m%d')

    if format == 'geojson':
        features = []
        for p in points:
            features.append({
                "type": "Feature",
                "geometry": {"type": "Point", "coordinates": [p.longitude, p.latitude]},
                "properties": {
                    "nom": p.nom, "description": p.description,
                    "categorie": p.get_categorie_display(), "statut": p.get_statut_display(),
                    "province": p.province, "commune": p.commune, "quartier": p.quartier,
                    "auteur": p.auteur.get_full_name() or p.auteur.username if p.auteur else '',
                    "date": p.date_creation.strftime('%d/%m/%Y %H:%M'),
                },
            })
        content = json.dumps({"type": "FeatureCollection", "features": features}, ensure_ascii=False, indent=2)
        nom_fichier = f"points_{date_str}.geojson"
        content_type = 'application/geo+json'

    elif format == 'kml':
        parts = ['<?xml version="1.0" encoding="UTF-8"?>',
                 '<kml xmlns="http://www.opengis.net/kml/2.2"><Document>',
                 f'<name>Points {date_str}</name>']
        for p in points:
            parts.append('<Placemark>')
            parts.append(f'<name>{_xml_safe(p.nom)}</name>')
            desc = f"{p.get_categorie_display()} - {p.get_statut_display()}"
            if p.description:
                desc += f" - {p.description}"
            parts.append(f'<description>{_xml_safe(desc)}</description>')
            parts.append(f'<Point><coordinates>{p.longitude},{p.latitude},0</coordinates></Point>')
            parts.append('</Placemark>')
        parts.append('</Document></kml>')
        content = '\n'.join(parts)
        nom_fichier = f"points_{date_str}.kml"
        content_type = 'application/vnd.google-earth.kml+xml'

    elif format == 'gpx':
        parts = ['<?xml version="1.0" encoding="UTF-8"?>',
                 '<gpx version="1.1" creator="PIP Cartographie" xmlns="http://www.topografix.com/GPX/1/1">']
        for p in points:
            parts.append(
                f'<wpt lat="{p.latitude}" lon="{p.longitude}">'
                f'<name>{_xml_safe(p.nom)}</name>'
                f'<desc>{_xml_safe(p.get_categorie_display())}</desc>'
                f'</wpt>'
            )
        parts.append('</gpx>')
        content = '\n'.join(parts)
        nom_fichier = f"points_{date_str}.gpx"
        content_type = 'application/gpx+xml'

    else:
        messages.error(request, "Format non supporté.")
        return redirect('index_cartographie')

    response = HttpResponse(content, content_type=content_type)
    response['Content-Disposition'] = f'attachment; filename="{nom_fichier}"'
    return response


# ─── EXPORT DE LA CARTE (PDF / JPEG) ────────────────────────────

FORMATS_PAPIER_MM = {
    'A4': (210, 297),
    'A3': (297, 420),
    'A2': (420, 594),
    'A1': (594, 841),
    'A0': (841, 1189),
}


@login_required
def export_carte_pdf(request):
    """Compose un PDF de mise en page cartographique à partir de l'image
    de la carte rendue par le client (titre, légende, échelle, nord et pied
    de page déjà composés) et des paramètres de page choisis."""
    if request.method != 'POST':
        return redirect('index_cartographie')

    try:
        donnees = json.loads(request.body or b'{}')
    except (ValueError, TypeError):
        return JsonResponse({'ok': False, 'erreur': 'Requête JSON invalide.'}, status=400)

    image_b64 = donnees.get('image', '')
    if not image_b64:
        return JsonResponse({'ok': False, 'erreur': 'Image manquante.'}, status=400)
    try:
        entete, _, b64 = image_b64.partition(',')
        image_bytes = base64.b64decode(b64)
    except Exception:
        return JsonResponse({'ok': False, 'erreur': 'Image illisible.'}, status=400)

    format_page = str(donnees.get('format_page', 'A4')).upper()
    if format_page not in FORMATS_PAPIER_MM:
        format_page = 'A4'
    orientation = str(donnees.get('orientation', 'L')).upper()
    if orientation not in ('P', 'L'):
        orientation = 'L'
    try:
        marge_mm = float(donnees.get('marge_mm', 12))
    except (TypeError, ValueError):
        marge_mm = 12.0
    marge_mm = min(max(marge_mm, 5.0), 40.0)

    from fpdf import FPDF
    from PIL import Image as ImgPil

    try:
        with ImgPil.open(io.BytesIO(image_bytes)) as img_check:
            img_check.verify()
    except Exception:
        return JsonResponse({'ok': False, 'erreur': 'Image corrompue ou non supportée.'}, status=400)

    from .branding import NOM, VERSION

    pdf = FPDF(orientation=orientation, unit='mm', format=format_page)
    pdf.set_auto_page_break(False)
    pdf.add_page()
    largeur, hauteur = FORMATS_PAPIER_MM[format_page]
    if orientation == 'L':
        largeur, hauteur = hauteur, largeur

    pdf.set_fill_color(255, 255, 255)
    pdf.rect(0, 0, largeur, hauteur, 'F')
    pdf.set_draw_color(77, 67, 246)
    pdf.set_line_width(0.8)
    pdf.rect(marge_mm - 2, marge_mm - 2, largeur - 2 * (marge_mm - 2), hauteur - 2 * (marge_mm - 2))

    utile_w = largeur - 2 * marge_mm
    utile_h = hauteur - 2 * marge_mm
    try:
        pdf.image(io.BytesIO(image_bytes), x=marge_mm, y=marge_mm, w=utile_w, h=utile_h)
    except Exception as e:
        return JsonResponse({'ok': False, 'erreur': f"Insertion de l'image impossible : {e}"}, status=400)

    projet_nom = str(donnees.get('projet', '') or 'carte')
    date_str = date.today().strftime('%Y%m%d')
    base_fichier = re.sub(r'[^\w\-\u00C0-\u017F]+', '_', projet_nom)[:60] or 'carte'
    nom_fichier = f"carte_{base_fichier}_{date_str}"
    if orientation == 'P':
        nom_fichier += '_portrait'

    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{nom_fichier}.pdf"'
    response.write(bytes(pdf.output()))
    _audit(request, "Export carte PDF", f"{nom_fichier}.pdf ({format_page} {orientation})")
    return response


# ─── DESSIN (SAUVEGARDE GÉOMÉTRIES) ────────────────────────────


def _couche_dessins(projet=None):
    if projet is not None:
        couche, _ = CoucheGeometrie.objects.get_or_create(
            nom='Dessins', projet=projet, defaults={'type_geometrie': 'point', 'style_couleur': '#8b5cf6'}
        )
        return couche
    couche, _ = CoucheGeometrie.objects.get_or_create(
        nom='Dessins', defaults={'type_geometrie': 'point', 'style_couleur': '#8b5cf6'}
    )
    return couche


@login_required
def dessin_save(request):
    if request.method != "POST":
        return JsonResponse({'ok': False, 'erreur': "Méthode non autorisée."})
    projet_actif = _projet_actif(request)
    if projet_actif is None:
        return JsonResponse({'ok': False, 'erreur': "Veuillez sélectionner un projet avant de dessiner."})
    try:
        data = json.loads(request.body)
    except json.JSONDecodeError:
        return JsonResponse({'ok': False, 'erreur': "Données invalides."})

    nom = str(data.get('nom', 'Dessin'))[:200]
    type_geom = data.get('type', 'LineString')
    coords = data.get('coordonnees')

    if coords is None:
        return JsonResponse({'ok': False, 'erreur': "Coordonnées manquantes."})

    type_dessin = type_geom
    if type_geom in ('Circle', 'Rectangle'):
        type_geom = 'Polygon'

    couche = _couche_dessins(projet_actif)
    if type_geom == 'Point':
        couche.type_geometrie = 'point'
        couche.save()
    elif type_geom == 'LineString':
        couche.type_geometrie = 'ligne'
        couche.save()
    elif type_geom in ('Polygon',):
        couche.type_geometrie = 'polygone'
        couche.save()

    geo = Geometrie.objects.create(
        couche=couche, type=type_geom, coordonnees=coords,
        proprietes={'nom': nom, 'type_dessin': type_dessin, 'auteur': request.user.username if request.user.is_authenticated else ''},
    )
    _audit(request, "Dessin enregistré", f"{type_dessin} - {nom} ({projet_actif.nom})")
    return JsonResponse({'ok': True, 'id': geo.pk, 'couche_id': couche.pk, 'couleur': couche.style_couleur})


@login_required
def geometrie_delete(request, pk):
    if request.method != "POST":
        return JsonResponse({'ok': False, 'erreur': "Méthode non autorisée."})
    geo = get_object_or_404(Geometrie, pk=pk)
    if geo.couche.nom == 'Dessins':
        _audit(request, "Suppression de dessin", f"Géométrie #{pk}")
        geo.delete()
        return JsonResponse({'ok': True})
    return JsonResponse({'ok': False, 'erreur': "Seuls les dessins peuvent être supprimés ici."})


# ─── MODE AVANCÉ / CODES D'ACCÈS ──────────────────────────────


def _etat_mode(request):
    """État courant du mode (classique / avancé) + droit d'accès, sans créer de préférence."""
    user = request.user
    if user.is_authenticated:
        pref = PreferenceUtilisateur.objects.filter(utilisateur=user).first()
        a_preference = pref is not None
        mode_pref = pref.mode if pref else 'classique'
        acces = False
        if _est_admin_principal(user):
            acces = True
        elif request.session.get('mode_avance_autorise'):
            acces = True
        elif pref and pref.code_lie and pref.code_lie.est_valide():
            acces = True
    else:
        a_preference = False
        mode_pref = 'classique'
        acces = bool(request.session.get('mode_avance_autorise'))
    return {
        'mode': mode_pref,
        'acces_avance': bool(acces),
        'a_preference': a_preference,
        'est_admin_principal': _est_admin_principal(user),
        'authentifie': user.is_authenticated,
    }


def api_mode(request):
    """État du mode d'utilisation (lecture)."""
    return JsonResponse(_etat_mode(request))


def _valider_code(request, code, user):
    obj = None
    for candidat in CodeAccesAvance.objects.all():
        if candidat.verifier_code(code):
            obj = candidat
            break
    if obj is None or not obj.est_valide():
        _audit(request, "Échec de validation d'un code Mode Avancé")
        return JsonResponse({'ok': False, 'besoin_code': True, 'erreur': 'Code invalide ou expiré.'}, status=403)
    obj.utiliser()
    request.session['mode_avance_autorise'] = True
    if user.is_authenticated:
        pref, _ = PreferenceUtilisateur.objects.get_or_create(utilisateur=user)
        pref.mode = 'avance'
        if obj.type == 'permanent':
            pref.code_lie = obj
        pref.save()
    _audit(request, "Mode Avancé activé par code", obj.get_type_display())
    return JsonResponse({'ok': True, 'mode': 'avance', 'acces_avance': True})


def api_mode_changer(request):
    """Change le mode d'utilisation (POST JSON {mode, code?})."""
    if request.method != 'POST':
        return JsonResponse({'erreur': 'Méthode non autorisée.'}, status=405)
    try:
        data = json.loads(request.body or '{}')
    except ValueError:
        data = {}
    demande = str(data.get('mode') or 'classique').lower()
    if demande not in ('classique', 'avance'):
        return JsonResponse({'erreur': 'Mode inconnu.'}, status=400)

    user = request.user
    if demande == 'classique':
        if user.is_authenticated:
            pref = PreferenceUtilisateur.objects.filter(utilisateur=user).first()
            if pref:
                pref.mode = 'classique'
                pref.save()
        request.session.pop('mode_avance_autorise', None)
        _audit(request, "Passage en Mode Classique")
        return JsonResponse({'ok': True, 'mode': 'classique', 'acces_avance': False})

    # Demande du Mode Avancé
    if _est_admin_principal(user):
        pref, _ = PreferenceUtilisateur.objects.get_or_create(utilisateur=user)
        pref.mode = 'avance'
        pref.save()
        request.session['mode_avance_autorise'] = True
        _audit(request, "Passage en Mode Avancé (administrateur principal)")
        return JsonResponse({'ok': True, 'mode': 'avance', 'acces_avance': True})

    if user.is_authenticated:
        pref = PreferenceUtilisateur.objects.filter(utilisateur=user).first()
        if pref and pref.code_lie and pref.code_lie.est_valide():
            pref.mode = 'avance'
            pref.save()
            request.session['mode_avance_autorise'] = True
            _audit(request, "Passage en Mode Avancé (code permanent)")
            return JsonResponse({'ok': True, 'mode': 'avance', 'acces_avance': True})

    code = str(data.get('code') or '').strip()
    if not code:
        _audit(request, "Demande de code Mode Avancé (code manquant)")
        return JsonResponse({'ok': False, 'besoin_code': True, 'erreur': 'Un code d\u2019accès est requis pour le Mode Avancé.'}, status=403)
    return _valider_code(request, code, user)


@login_required
def mode_avance_admin(request):
    """Page de gestion des codes d'accès — réservée à l'administrateur principal."""
    if not _est_admin_principal(request.user):
        messages.error(request, "Seul l'administrateur principal peut gérer les codes d'accès au Mode Avancé.")
        return redirect('index_cartographie')
    return render(request, 'cartographie/mode_avance_admin.html', {
        'codes': CodeAccesAvance.objects.all(),
        'est_admin_principal': True,
    })


@user_passes_test(_est_admin)
def guide_utilisation(request):
    """Section « Guide d'utilisation » de l'espace Super Admin :
    consultation, téléchargement (PDF/Word) et régénération du guide officiel MUKESHABA."""
    import datetime
    from django.core.management import call_command

    from cartographie.guide_contenu import CHAPITRES, META

    dossier = os.path.join(settings.BASE_DIR, 'cartographie', 'static', 'docs')
    if request.method == 'POST':
        if not _est_admin_principal(request.user):
            messages.error(request, "Seul l'administrateur principal peut régénérer le guide d'utilisation.")
            return redirect('guide_utilisation')
        try:
            call_command('generer_guide')
            call_command('generer_presentation')
            _audit(request, "Régénération du guide d'utilisation", "PDF + Word + PowerPoint")
            messages.success(request, "Le guide d'utilisation et la présentation ont été régénérés avec succès.")
        except Exception as exc:
            messages.error(request, "Erreur lors de la régénération du guide : %s" % exc)
        return redirect('guide_utilisation')

    fichiers = {}
    for ext, nom in (('pdf', 'MUKMAP_Guide_Complet_Utilisateur.pdf'),
                     ('docx', 'MUKMAP_Guide_Complet_Utilisateur.docx'),
                     ('pptx', 'MUKMAP_Presentation_MUKESHABA.pptx')):
        chemin = os.path.join(dossier, nom)
        if os.path.exists(chemin):
            fichiers[ext] = {
                'nom': nom,
                'taille_ko': round(os.path.getsize(chemin) / 1024, 1),
                'modifie': datetime.datetime.fromtimestamp(
                    os.path.getmtime(chemin)).strftime('%d/%m/%Y %H:%M'),
            }
        else:
            fichiers[ext] = None

    return render(request, 'cartographie/guide_utilisation.html', {
        'meta': META,
        'chapitres': CHAPITRES,
        'fichiers': fichiers,
        'est_admin_principal': _est_admin_principal(request.user),
    })


@login_required
def api_codes_mode(request):
    """Génération / liste des codes d'accès (POST/GET) — réservée à l'administrateur principal."""
    if not _est_admin_principal(request.user):
        return JsonResponse({'erreur': 'Accès refusé.'}, status=403)

    if request.method == 'GET':
        return JsonResponse({'codes': [
            {
                'id': c.pk, 'libelle': c.libelle, 'type': c.type,
                'expire_le': c.expire_le.strftime('%d/%m/%Y %H:%M') if c.expire_le else '',
                'max_utilisations': c.max_utilisations,
                'utilisations': c.utilisations,
                'actif': c.actif,
                'date_creation': c.date_creation.strftime('%d/%m/%Y %H:%M'),
            }
            for c in CodeAccesAvance.objects.all()
        ]})

    if request.method == 'POST':
        try:
            data = json.loads(request.body or '{}')
        except ValueError:
            data = {}
        type_code = str(data.get('type') or 'permanent').lower()
        if type_code not in ('temporaire', 'permanent'):
            return JsonResponse({'erreur': 'Type de code invalide.'}, status=400)

        expire_le = None
        max_util = None
        if type_code == 'temporaire':
            heures = data.get('duree_heures')
            if heures is not None:
                try:
                    expire_le = timezone.now() + timedelta(hours=float(heures))
                except (ValueError, TypeError):
                    return JsonResponse({'erreur': 'Durée de validité invalide.'}, status=400)
        max_util_data = data.get('max_utilisations')
        if max_util_data is not None and max_util_data != '':
            try:
                max_util = int(max_util_data)
                if max_util < 1:
                    raise ValueError
            except (ValueError, TypeError):
                return JsonResponse({'erreur': "Nombre d'utilisations invalide."}, status=400)

        code_clair = CodeAccesAvance.generer()
        obj = CodeAccesAvance.objects.create(
            libelle=str(data.get('libelle') or '')[:200],
            code_hash=CodeAccesAvance.hacher(code_clair),
            type=type_code,
            expire_le=expire_le,
            max_utilisations=max_util,
            cree_par=request.user,
        )
        _audit(request, "Génération de code Mode Avancé", f"Code #{obj.pk} ({obj.get_type_display()})")
        return JsonResponse({'ok': True, 'id': obj.pk, 'code': code_clair})

    return JsonResponse({'erreur': 'Méthode non autorisée.'}, status=405)


@login_required
def api_code_revoquer(request, pk):
    """Révocation d'un code — réservée à l'administrateur principal."""
    if not _est_admin_principal(request.user):
        return JsonResponse({'erreur': 'Accès refusé.'}, status=403)
    if request.method != 'POST':
        return JsonResponse({'erreur': 'Méthode non autorisée.'}, status=405)
    obj = get_object_or_404(CodeAccesAvance, pk=pk)
    obj.actif = False
    obj.save(update_fields=['actif'])
    _audit(request, "Révocation de code Mode Avancé", f"Code #{obj.pk}")
    return JsonResponse({'ok': True})


def _fonds_carte_json(obj, user=None):
    d = {
        'id': f'ext-{obj.pk}',
        'pk': obj.pk,
        'nom': obj.nom,
        'type': obj.type_fond,
        'url': obj.url,
        'attribution': obj.attribution,
        'categorie': obj.categorie,
        'crs': obj.crs,
        'layers': obj.layers,
        'projet': obj.projet_id,
        'portee': 'projet' if obj.projet_id else 'personnel',
    }
    if user is not None and user.is_authenticated and (obj.auteur == user or _est_admin_principal(user)):
        d['cle_api'] = obj.cle_api
    return d


@login_required
def api_fonds_personnalises(request):
    """Liste (GET) / création (POST) de fonds de carte personnalisés — Mode Avancé requis en écriture."""
    if request.method == 'GET':
        qs = FondCartePersonnalise.objects.filter(visible=True)
        if not (_est_admin_principal(request.user) or _etat_mode(request)['acces_avance']):
            qs = qs.filter(Q(auteur=request.user) | Q(projet__statut='actif'))
        return JsonResponse({'fonds': [_fonds_carte_json(o, request.user) for o in qs]})

    if request.method == 'POST':
        if not _etat_mode(request)['acces_avance']:
            return JsonResponse({'erreur': 'Mode Avancé requis pour créer un fond de carte.'}, status=403)
        try:
            data = json.loads(request.body or '{}')
        except ValueError:
            data = {}
        nom = str(data.get('nom') or '').strip()
        url = str(data.get('url') or '').strip()
        if not nom:
            return JsonResponse({'erreur': 'Le nom est requis.'}, status=400)
        if not url.startswith('http://') and not url.startswith('https://'):
            return JsonResponse({'erreur': 'URL invalide (http/https requis).'}, status=400)
        type_fond = str(data.get('type') or 'xyz').lower()
        if type_fond not in dict(FondCartePersonnalise.TYPE_CHOICES):
            return JsonResponse({'erreur': 'Type de fond invalide.'}, status=400)
        if type_fond in ('xyz', 'wms', 'wmts', 'vector', 'arcgis'):
            if '{' not in url or '}' not in url:
                return JsonResponse({'erreur': "L'URL doit contenir un modèle de tuiles, ex. {z}/{x}/{y}, {TileMatrix}/{TileCol}/{TileRow} ou {bbox-epsg-3857}."}, status=400)
        if type_fond == 'wms' and '{bbox-epsg-3857}' not in url:
            return JsonResponse({'erreur': "L'URL WMS doit contenir le paramètre BBOX={bbox-epsg-3857}."}, status=400)
        categorie = str(data.get('categorie') or 'geologie').lower()
        if categorie not in dict(FondCartePersonnalise.CATEGORIE_CHOICES):
            categorie = 'geologie'
        projet = None
        try:
            projet_id = int(data.get('projet') or 0)
        except (TypeError, ValueError):
            projet_id = 0
        if projet_id:
            projet = Projet.objects.filter(pk=projet_id, statut='actif').first()
        elif str(data.get('portee') or '') == 'projet':
            session_pid = request.session.get('projet_actif_id')
            if session_pid:
                projet = Projet.objects.filter(pk=session_pid, statut='actif').first()
        obj = FondCartePersonnalise.objects.create(
            nom=nom[:120],
            type_fond=type_fond,
            url=url,
            attribution=str(data.get('attribution') or '')[:255],
            categorie=categorie,
            cle_api=str(data.get('cle_api') or '')[:200],
            crs=str(data.get('crs') or 'EPSG:3857')[:30],
            layers=str(data.get('layers') or '')[:300],
            projet=projet,
            auteur=request.user,
        )
        _audit(request, "Création d'un fond de carte personnalisé", f"#{obj.pk} {obj.nom} ({type_fond})")
        return JsonResponse({'ok': True, 'fond': _fonds_carte_json(obj, request.user)})

    return JsonResponse({'erreur': 'Méthode non autorisée.'}, status=405)


@login_required
def api_fond_personnalise_detail(request, pk):
    """Mise à jour (PATCH) / suppression (DELETE) d'un fond personnalisé — auteur ou administrateur principal."""
    obj = get_object_or_404(FondCartePersonnalise, pk=pk)
    if obj.auteur != request.user and not _est_admin_principal(request.user):
        return JsonResponse({'erreur': 'Accès refusé.'}, status=403)

    if request.method == 'PATCH':
        try:
            data = json.loads(request.body or '{}')
        except ValueError:
            data = {}
        for champ, max_len in (('nom', 120), ('attribution', 255), ('cle_api', 200), ('crs', 30), ('layers', 300)):
            if champ in data and isinstance(data[champ], str):
                setattr(obj, champ, data[champ][:max_len])
        if 'type' in data and str(data['type']) in dict(FondCartePersonnalise.TYPE_CHOICES):
            obj.type_fond = str(data['type'])
        if 'url' in data and isinstance(data['url'], str) and data['url'].strip():
            obj.url = data['url'].strip()
        if 'categorie' in data and str(data['categorie']) in dict(FondCartePersonnalise.CATEGORIE_CHOICES):
            obj.categorie = str(data['categorie'])
        if 'visible' in data:
            obj.visible = bool(data['visible'])
        if 'projet' in data:
            try:
                pid = int(data['projet']) or None
            except (TypeError, ValueError):
                pid = None
            obj.projet = Projet.objects.filter(pk=pid, statut='actif').first() if pid else None
        obj.save()
        _audit(request, "Modification d'un fond de carte personnalisé", f"#{obj.pk} {obj.nom}")
        return JsonResponse({'ok': True, 'fond': _fonds_carte_json(obj, request.user)})

    if request.method == 'DELETE':
        _audit(request, "Suppression d'un fond de carte personnalisé", f"#{obj.pk} {obj.nom}")
        obj.delete()
        return JsonResponse({'ok': True})

    return JsonResponse({'erreur': 'Méthode non autorisée.'}, status=405)


# ─── COUCHES WMS (superposables, opacité réglable) ─────


def _couche_wms_json(obj):
    return {
        'id': obj.pk,
        'nom': obj.nom,
        'url': obj.url,
        'layers': obj.layers,
        'version': obj.version,
        'attribution': obj.attribution,
        'opacite': obj.opacite,
        'visible': obj.visibilite,
        'ordre': obj.ordre,
        'auteur': obj.auteur.username if obj.auteur else '',
    }


@login_required
def api_couches_wms(request):
    """Liste (GET) / création (POST) de couches WMS — Mode Avancé requis en écriture."""
    if request.method == 'GET':
        qs = CoucheWMS.objects.all()
        if not (_est_admin_principal(request.user) or _etat_mode(request)['acces_avance']):
            qs = qs.filter(auteur=request.user)
        return JsonResponse({'couches': [_couche_wms_json(o) for o in qs]})

    if request.method == 'POST':
        if not _etat_mode(request)['acces_avance']:
            return JsonResponse({'erreur': 'Mode Avancé requis pour créer une couche WMS.'}, status=403)
        try:
            data = json.loads(request.body or '{}')
        except ValueError:
            data = {}
        nom = str(data.get('nom') or '').strip()
        url = str(data.get('url') or '').strip()
        if not nom:
            return JsonResponse({'erreur': 'Le nom est requis.'}, status=400)
        if not url.startswith('http://') and not url.startswith('https://'):
            return JsonResponse({'erreur': 'URL invalide (http/https requis).'}, status=400)
        if '{bbox-epsg-3857}' not in url:
            return JsonResponse({'erreur': "L'URL WMS doit contenir le paramètre BBOX={bbox-epsg-3857}."}, status=400)
        try:
            opacite = float(data.get('opacite') or 0.7)
            opacite = max(0.0, min(1.0, opacite))
        except (TypeError, ValueError):
            opacite = 0.7
        obj = CoucheWMS.objects.create(
            nom=nom[:150],
            url=url,
            layers=str(data.get('layers') or '')[:500],
            version=str(data.get('version') or '1.1.1')[:20],
            attribution=str(data.get('attribution') or '')[:255],
            opacite=opacite,
            visibilite=bool(data.get('visible', True)),
            ordre=int(data.get('ordre') or 0),
            projet=None,
            auteur=request.user,
        )
        _audit(request, "Création d'une couche WMS", f"#{obj.pk} {obj.nom}")
        return JsonResponse({'ok': True, 'couche': _couche_wms_json(obj)}, status=201)

    return JsonResponse({'erreur': 'Méthode non autorisée.'}, status=405)


@login_required
def api_couche_wms_detail(request, pk):
    """Mise à jour (PATCH) / suppression (DELETE) — auteur ou administrateur principal."""
    obj = get_object_or_404(CoucheWMS, pk=pk)
    if obj.auteur != request.user and not _est_admin_principal(request.user):
        return JsonResponse({'erreur': 'Accès refusé.'}, status=403)

    if request.method == 'PATCH':
        try:
            data = json.loads(request.body or '{}')
        except ValueError:
            data = {}
        if 'nom' in data:
            obj.nom = str(data['nom'] or '').strip()[:150] or obj.nom
        if 'attribution' in data:
            obj.attribution = str(data['attribution'] or '')[:255]
        if 'opacite' in data:
            try:
                obj.opacite = max(0.0, min(1.0, float(data['opacite'])))
            except (TypeError, ValueError):
                pass
        if 'visible' in data:
            obj.visibilite = bool(data['visible'])
        if 'ordre' in data:
            try:
                obj.ordre = max(0, int(data['ordre']))
            except (TypeError, ValueError):
                pass
        if 'layers' in data:
            obj.layers = str(data['layers'] or '')[:500]
        obj.save()
        _audit(request, "Mise à jour d'une couche WMS", f"#{obj.pk} {obj.nom}")
        return JsonResponse({'ok': True, 'couche': _couche_wms_json(obj)})

    if request.method == 'DELETE':
        _audit(request, "Suppression d'une couche WMS", f"#{obj.pk} {obj.nom}")
        obj.delete()
        return JsonResponse({'ok': True})

    return JsonResponse({'erreur': 'Méthode non autorisée.'}, status=405)


# ─── IMAGERIE (orthophotos / images drone géoréférencées) ───────


def _coordonnees_exif(fichier):
    """Extrait la position GPS des métadonnées EXIF — (lon, lat) ou None."""
    try:
        from PIL import ExifTags, Image as PILImage
        fichier.seek(0)
        with PILImage.open(fichier) as img:
            exif = img.getexif()
            if not exif:
                return None
            gps = exif.get_ifd(ExifTags.IFD.GPSInfo)
            if not gps or 2 not in gps or 4 not in gps:
                return None

            def degres(valeurs, ref):
                d, m, s = [float(x) for x in valeurs]
                v = d + m / 60.0 + s / 3600.0
                return -v if ref in ('S', 'W') else v

            return (degres(gps[4], gps.get(3, 'E')), degres(gps[2], gps.get(1, 'N')))
    except Exception:
        return None


def _coordonnees_worldfile(worldfile, dimensions):
    """Applique les 6 coefficients du WorldFile aux dimensions (w, h) de l'image."""
    try:
        valeurs = []
        for ligne in worldfile.read().decode('utf-8', errors='replace').splitlines():
            ligne = ligne.strip()
            if not ligne or ligne.startswith('#'):
                continue
            valeurs.append(float(ligne))
            if len(valeurs) == 6:
                break
        if len(valeurs) < 6:
            return None
        A, D, B, E, C, F = valeurs  # X = C + A*col + B*row ; Y = F + D*col + E*row
        w, h = dimensions
        coins = [
            [C, F],                                   # (0,0) haut-gauche
            [C + A * w, F + D * w],                   # (w,0) haut-droite
            [C + A * w + B * h, F + D * w + E * h],   # (w,h) bas-droite
            [C + B * h, F + E * h],                   # (0,h) bas-gauche
        ]
        lons = [c[0] for c in coins]
        lats = [c[1] for c in coins]
        return {
            'coords': coins,
            'min_lon': min(lons), 'max_lon': max(lons),
            'min_lat': min(lats), 'max_lat': max(lats),
        }
    except (ValueError, UnicodeDecodeError, TypeError):
        return None


def _image_aerienne_json(obj):
    return {
        'id': obj.pk,
        'nom': obj.nom,
        'url': obj.fichier.url if obj.fichier else '',
        'type': obj.type_imagerie,
        'mode_geo': obj.mode_geo,
        'min_lon': obj.min_lon, 'min_lat': obj.min_lat,
        'max_lon': obj.max_lon, 'max_lat': obj.max_lat,
        'coords': obj.coords,
        'visible': obj.visibilite,
        'date_prise': obj.date_prise.isoformat() if obj.date_prise else '',
        'description': obj.description,
    }


@login_required
def api_imagerie(request):
    """Liste (GET) / ajout (POST multipart) d'images aériennes — Mode Avancé requis en écriture."""
    if request.method == 'GET':
        return JsonResponse({'images': [_image_aerienne_json(o) for o in ImageAerienne.objects.all()]})

    if request.method == 'POST':
        if not _etat_mode(request)['acces_avance']:
            return JsonResponse({'erreur': 'Mode Avancé requis pour ajouter une image.'}, status=403)
        fichier = request.FILES.get('fichier')
        if not fichier:
            return JsonResponse({'erreur': 'Fichier image requis.'}, status=400)
        nom = str(request.POST.get('nom') or '').strip() or (getattr(fichier, 'name', '') or 'Image')
        type_imagerie = str(request.POST.get('type') or 'ortho').lower()
        if type_imagerie not in dict(ImageAerienne.TYPE_CHOICES):
            type_imagerie = 'ortho'

        bbox = {}
        try:
            min_lon = float(request.POST.get('min_lon') or '')
            min_lat = float(request.POST.get('min_lat') or '')
            max_lon = float(request.POST.get('max_lon') or '')
            max_lat = float(request.POST.get('max_lat') or '')
            bbox = {'min_lon': min_lon, 'min_lat': min_lat, 'max_lon': max_lon, 'max_lat': max_lat}
        except ValueError:
            pass

        try:
            from PIL import Image as PILImage
            fichier.seek(0)
            with PILImage.open(fichier) as img:
                dimensions = img.size
        except Exception:
            dimensions = None
        try:
            fichier.seek(0)
        except Exception:
            pass

        mode_geo = 'bbox'
        coords = None
        worldfile = request.FILES.get('worldfile')
        if worldfile and dimensions:
            geo = _coordonnees_worldfile(worldfile, dimensions)
            if geo:
                bbox = geo
                coords = geo['coords']
                mode_geo = 'worldfile'
            else:
                return JsonResponse({'erreur': 'WorldFile illisible (6 coefficients numériques attendus).'}, status=400)
        elif not bbox:
            geo = _coordonnees_exif(fichier)
            if geo:
                lon, lat = geo
                demi = 0.0015
                bbox = {'min_lon': lon - demi, 'min_lat': lat - demi,
                        'max_lon': lon + demi, 'max_lat': lat + demi}
                mode_geo = 'exif'
        if not bbox:
            return JsonResponse({'erreur': 'Coordonnées manquantes : fournissez une emprise (bbox), un WorldFile ou des métadonnées EXIF GPS.'}, status=400)

        altitude_val = None
        altitude = request.POST.get('altitude') or ''
        if altitude:
            try:
                altitude_val = float(altitude)
            except ValueError:
                altitude_val = None
        date_val = None
        date_prise = request.POST.get('date_prise') or ''
        if date_prise:
            try:
                from datetime import datetime
                date_val = datetime.strptime(date_prise, '%Y-%m-%d').date()
            except ValueError:
                date_val = None

        obj = ImageAerienne.objects.create(
            nom=nom[:150],
            fichier=fichier,
            type_imagerie=type_imagerie,
            mode_geo=mode_geo,
            min_lon=bbox.get('min_lon'), min_lat=bbox.get('min_lat'),
            max_lon=bbox.get('max_lon'), max_lat=bbox.get('max_lat'),
            coords=coords,
            altitude_m=altitude_val,
            date_prise=date_val,
            description=str(request.POST.get('description') or '')[:2000],
            projet_id=request.session.get('projet_actif_id') or None,
            auteur=request.user,
        )
        _audit(request, "Ajout d'une image aérienne", f"#{obj.pk} {obj.nom} ({mode_geo})")
        return JsonResponse({'ok': True, 'image': _image_aerienne_json(obj)})

    return JsonResponse({'erreur': 'Méthode non autorisée.'}, status=405)


@login_required
def api_imagerie_detail(request, pk):
    """Suppression (DELETE) d'une image aérienne — auteur ou administrateur principal."""
    if request.method != 'DELETE':
        return JsonResponse({'erreur': 'Méthode non autorisée.'}, status=405)
    obj = get_object_or_404(ImageAerienne, pk=pk)
    if obj.auteur != request.user and not _est_admin_principal(request.user):
        return JsonResponse({'erreur': 'Accès refusé.'}, status=403)
    _audit(request, "Suppression d'une image aérienne", f"#{obj.pk} {obj.nom}")
    if obj.fichier:
        try:
            obj.fichier.delete(save=False)
        except Exception:
            pass
    obj.delete()
    return JsonResponse({'ok': True})


@login_required
def api_imagerie_visibilite(request, pk):
    """Bascule la visibilité d'une image aérienne (POST {visible: bool})."""
    if request.method != 'POST':
        return JsonResponse({'erreur': 'Méthode non autorisée.'}, status=405)
    if not _etat_mode(request)['acces_avance']:
        return JsonResponse({'erreur': 'Mode Avancé requis.'}, status=403)
    obj = get_object_or_404(ImageAerienne, pk=pk)
    if obj.auteur != request.user and not _est_admin_principal(request.user):
        return JsonResponse({'erreur': 'Accès refusé.'}, status=403)
    try:
        data = json.loads(request.body or '{}')
    except ValueError:
        data = {}
    obj.visibilite = bool(data.get('visible', not obj.visibilite))
    obj.save(update_fields=['visibilite'])
    return JsonResponse({'ok': True, 'visible': obj.visibilite})
